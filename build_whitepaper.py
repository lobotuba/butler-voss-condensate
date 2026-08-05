# -*- coding: utf-8 -*-
"""Build the COMPLETE-PROJECT Butler-Voss Condensate white paper (WP-06) as .docx,
with matplotlib-rendered figures."""
import os
import numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL, RUST, MUT = "#009e8a", "#c2542a", "#61706a"
TEAL_D, RUST_D = RGBColor(0x05, 0x7d, 0x6d), RGBColor(0xac, 0x48, 0x22)
INK, GREY = RGBColor(0x1a, 0x20, 0x1d), RGBColor(0x61, 0x70, 0x6a)
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = r"Z:\Butler-Voss Condensate\Butler-Voss Condensate - Whitepaper.docx"
OLD = r"Z:\Butler-Voss Condensate\Butler-Voss Condensate - Screening Whitepaper.docx"

# ----------------------------------------------------------------- figures ----
plt.rcParams.update({"font.family": "serif", "font.serif": ["Georgia", "DejaVu Serif"],
                     "font.size": 11, "axes.edgecolor": "#c5cbc3", "axes.linewidth": .8,
                     "figure.facecolor": "white", "axes.facecolor": "white"})

def _style(ax):
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(colors=MUT, labelsize=9.5, length=3)
    ax.grid(axis="y", color="#e3e6e1", lw=.8); ax.set_axisbelow(True)
    for s in ("left", "bottom"): ax.spines[s].set_color("#c5cbc3")

def fig_drho(p):
    r = [.5,1.5,2.5,3.5,4.5,5.5,6.5,7.5,8.5,9.5,10.5,11.5]
    d = [.0537,.0525,.0491,.0441,.0365,.0272,.0196,.0140,.0101,.0077,.0063,.0055]
    fr,fy=[.5,2.5,4.5,6.5,8.5,10.5,11.5],[.0545,.0470,.0345,.0230,.0135,.0075,.0058]
    f,ax=plt.subplots(figsize=(6.4,3.2))
    ax.plot(fr,fy,"--",color=MUT,lw=1.4,label="fit: exp(-r/λ), λ = 3.3")
    ax.plot(r,d,"-o",color=RUST,lw=2,ms=5,mfc="white",mec=RUST,mew=1.8,label="measured  drho(r)")
    ax.set_xlabel("radius  r"); ax.set_ylabel("drho(r)"); ax.set_ylim(0,.06); ax.set_xlim(0,12)
    _style(ax); ax.legend(frameon=False,fontsize=9,loc="upper right")
    f.tight_layout(); f.savefig(p,dpi=200,bbox_inches="tight"); plt.close(f)

def fig_box(p):
    x=[10,14,20]
    f,ax=plt.subplots(figsize=(6.4,3.3))
    ax.plot(x,[2.45,3.18,4.24],"-o",color=TEAL,lw=2,ms=6,mfc="white",mec=TEAL,mew=1.8)
    ax.plot(x,[1.86,2.02,2.18],"-o",color=RUST,lw=2,ms=6,mfc="white",mec=RUST,mew=1.8)
    ax.annotate("massless: 1/r (grows)",(20,4.24),xytext=(8,4),textcoords="offset points",
                color=TEAL,fontsize=9.5,fontweight="bold",va="center")
    ax.annotate("massive: 1/m (pinned)",(20,2.18),xytext=(8,0),textcoords="offset points",
                color=RUST,fontsize=9.5,fontweight="bold",va="center")
    ax.set_xlabel("box size"); ax.set_ylabel("apparent range  λ")
    ax.set_xlim(9,24); ax.set_ylim(1.5,4.6); ax.set_xticks(x); _style(ax)
    f.tight_layout(); f.savefig(p,dpi=200,bbox_inches="tight"); plt.close(f)

def fig_gauge(p):
    e=[.15,.20,.30]; lam=[4.14,3.13,1.94]; gx=[.13,.15,.2,.25,.3,.33]
    f,ax=plt.subplots(figsize=(6.4,3.2))
    ax.plot(gx,[.6/v for v in gx],"--",color=MUT,lw=1.4,label="guide: λ_L ~ 1/e")
    ax.plot(e,lam,"o",color=RUST,ms=7,mfc="white",mec=RUST,mew=2,label="measured λ_L")
    ax.set_xlabel("gauge coupling  e"); ax.set_ylabel("penetration depth  λ_L")
    ax.set_xlim(.12,.33); ax.set_ylim(1.5,4.6); ax.set_xticks(e); _style(ax)
    ax.legend(frameon=False,fontsize=9,loc="upper right")
    f.tight_layout(); f.savefig(p,dpi=200,bbox_inches="tight"); plt.close(f)

def fig_monopole(p):
    r=np.array([2.5,3.5,4.5,5.5,6.5,7.5,8.5,9.5,10.5])
    B=np.array([.096,.044,.026,.017,.012,.009,.007,.006,.005])
    guide=B[1]*(r[1]/r)**2
    f,ax=plt.subplots(figsize=(6.4,3.2))
    ax.plot(r,guide,"--",color=MUT,lw=1.4,label="reference: 1/r^2")
    ax.plot(r,B,"o",color=TEAL,ms=6,mfc="white",mec=TEAL,mew=1.8,label="monopole field |B|(r)")
    ax.set_xscale("log"); ax.set_yscale("log")
    ax.set_xlabel("radius  r  (log)"); ax.set_ylabel("|B|  (log)")
    _style(ax); ax.grid(which="both",axis="both",color="#e9ece7",lw=.7)
    ax.legend(frameon=False,fontsize=9,loc="lower left")
    f.tight_layout(); f.savefig(p,dpi=200,bbox_inches="tight"); plt.close(f)

def fig_lorentz(p):
    kf=np.array([.05,.1,.2,.3,.4,.6,.8])
    hexa=np.array([1.06e-7,1.70e-6,2.73e-5,1.40e-4,4.52e-4,2.41e-3,8.25e-3])
    fcc=np.array([1.71e-4,6.86e-4,2.75e-3,6.23e-3,1.11e-2,2.56e-2,4.69e-2])
    f,ax=plt.subplots(figsize=(6.4,3.2))
    ax.plot(kf,fcc,"-o",color=RUST,lw=2,ms=5,mfc="white",mec=RUST,mew=1.8,label="fcc 3D  (~ (k/kmax)^2)")
    ax.plot(kf,hexa,"-o",color=TEAL,lw=2,ms=5,mfc="white",mec=TEAL,mew=1.8,label="hex 2D  (~ (k/kmax)^4)")
    ax.set_xscale("log"); ax.set_yscale("log")
    ax.set_xlabel("k / k_max   (energy / E_Planck)"); ax.set_ylabel("anisotropy  dc/c  (Lorentz violation)")
    _style(ax); ax.grid(which="both",axis="both",color="#e9ece7",lw=.7)
    ax.legend(frameon=False,fontsize=9,loc="upper left")
    f.tight_layout(); f.savefig(p,dpi=200,bbox_inches="tight"); plt.close(f)

figs = {}
for nm, fn in [("drho",fig_drho),("box",fig_box),("gauge",fig_gauge),("mono",fig_monopole),
               ("lorentz",fig_lorentz)]:
    figs[nm]=os.path.join(HERE,f"wpf_{nm}.png"); fn(figs[nm])

# ------------------------------------------------------------------- docx -----
doc = Document()
st = doc.styles["Normal"]; st.font.name="Cambria"; st.font.size=Pt(11)
st.paragraph_format.space_after=Pt(7); st.paragraph_format.line_spacing=1.16
for hs, sz in (("Heading 1",15),("Heading 2",12.5),("Title",25)):
    s=doc.styles[hs]; s.font.name="Cambria"; s.font.size=Pt(sz); s.font.color.rgb=INK

def shade(par, hexcol):
    sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexcol)
    par._p.get_or_add_pPr().append(sh)

import re as _re
from docx.oxml import parse_xml

# ------------------------------------------------------------------- typography ----
# Prose carries light math markup: _{...} / ^{...} for multi-character scripts and
# _x / ^2 for single tokens.  rich() turns that into real sub/superscript runs so the
# document shows h_ij and 10^-16 set properly rather than as ASCII carets.
_SCRIPT = _re.compile(r'([_^])(?:\{([^{}]*)\}|(-?[0-9A-Za-z.]+))')

def rich(p, text, **font):
    """Append `text` to paragraph `p`, rendering _/^ markup as sub/superscript runs."""
    def style(r):
        for k, v in font.items():
            setattr(r.font, k, v) if k != "italic" else setattr(r, k, v)
        return r
    pos = 0
    for m in _SCRIPT.finditer(text):
        if m.start() > pos:
            style(p.add_run(text[pos:m.start()]))
        r = style(p.add_run(m.group(2) if m.group(2) is not None else m.group(3)))
        r.font.subscript = (m.group(1) == "_")
        r.font.superscript = (m.group(1) == "^")
        pos = m.end()
    if pos < len(text):
        style(p.add_run(text[pos:]))
    return p


def body(text, justify=True, italic=False, size=None, color=None, after=7, before=0):
    p=doc.add_paragraph(); rich(p, text)
    for r in p.runs:
        if italic: r.italic=True
        if size: r.font.size=Pt(size)
        if color: r.font.color.rgb=color
    if justify: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    return p

def heading(text, lvl=1):
    h=doc.add_heading(level=lvl); h.paragraph_format.space_before=Pt(12 if lvl==1 else 8)
    parts=text.split("  ",1)
    r=h.add_run(parts[0]+"  "); r.font.color.rgb=TEAL_D
    h.add_run(parts[1] if len(parts)>1 else "")
    return h

_MNS = ('xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')

def _esc(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _mrun(s):
    return f'<m:r><m:t xml:space="preserve">{_esc(s)}</m:t></m:r>' if s else ""

import unicodedata as _ud

def _isletter(c):
    """Identifier char: any letter, a combining mark (so Q⃛ stays one atom), or a prime."""
    cat = _ud.category(c)
    return cat[0] == "L" or cat == "Mn" or c in "∂∇'′"

def _delim(op, cl, inner):
    pr = ""
    if (op, cl) != ("(", ")"):
        pr = f'<m:dPr><m:begChr m:val="{_esc(op)}"/><m:endChr m:val="{_esc(cl)}"/></m:dPr>'
    return f"<m:d>{pr}<m:e>{inner}</m:e></m:d>"

def _attach(s, i, base):
    """Consume any _{..}/^{..} following position i and wrap `base` accordingly."""
    sub = sup = None
    while i + 1 < len(s) and s[i] in "_^" and s[i + 1] == "{":
        j = s.index("}", i)
        if s[i] == "_":
            sub = s[i + 2:j]
        else:
            sup = s[i + 2:j]
        i = j + 1
    if sub is not None and sup is not None:
        base = (f"<m:sSubSup><m:e>{base}</m:e><m:sub>{_mrun(sub)}</m:sub>"
                f"<m:sup>{_mrun(sup)}</m:sup></m:sSubSup>")
    elif sub is not None:
        base = f"<m:sSub><m:e>{base}</m:e><m:sub>{_mrun(sub)}</m:sub></m:sSub>"
    elif sup is not None:
        base = f"<m:sSup><m:e>{base}</m:e><m:sup>{_mrun(sup)}</m:sup></m:sSup>"
    return base, i

def _parse(s, i=0, stop=None):
    """Recursive-descent over the equation markup -> list of OMML fragments."""
    out = []
    while i < len(s):
        c = s[i]
        if stop and c == stop:
            break
        if c in "([" or c == "|":
            cl = {"(": ")", "[": "]", "|": "|"}[c]
            inner, i = _parse(s, i + 1, cl)
            i += 1                                   # consume the closer
            node, i = _attach(s, i, _delim(c, cl, "".join(inner)))
            out.append(node)
            continue
        j = i
        if _isletter(c):
            while j < len(s) and _isletter(s[j]):
                j += 1
        elif c.isdigit():
            while j < len(s) and (s[j].isdigit() or s[j] == "."):
                j += 1
        else:
            j = i + 1                                # lone operator / space
        node, i = _attach(s, j, _mrun(s[i:j]))
        out.append(node)
    return out, i

def _omml(markup):
    xml = "".join(_parse(markup)[0])
    if "<m:e></m:e>" in xml:                          # Word renders empty slots as dotted boxes
        raise ValueError(f"empty math slot in equation: {markup!r}")
    return xml

def add_eq(markup, num):
    """Display equation as a real Word math object, with a right-hand equation number."""
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(8)
    p._p.append(parse_xml(f'<m:oMath {_MNS}>{_omml(markup)}</m:oMath>'))
    t=p.add_run("      ("+num+")"); t.font.color.rgb=GREY
    return p

def caption(no, text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(12)
    r=p.add_run("Figure "+no+". "); r.bold=True; r.font.size=Pt(9.5)
    r2=p.add_run(text); r2.font.size=Pt(9.5); r2.font.color.rgb=GREY
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def add_figure(path, no, cap):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    p.add_run().add_picture(path, width=Inches(5.6)); caption(no, cap)

def result(tag, text, warn=False):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(9)
    p.paragraph_format.left_indent=Inches(0.12); shade(p, "FBEDE6" if warn else "E6F4F1")
    r=p.add_run(tag+"  "); r.bold=True; r.font.color.rgb=(RUST_D if warn else TEAL_D)
    rich(p, text); [setattr(rn.font,"size",Pt(10.5)) for rn in p.runs]
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def table(headers, rows, cap=None, wide=None):
    if cap:
        cp=doc.add_paragraph(); cr=cp.add_run(cap); cr.font.size=Pt(9.5); cp.paragraph_format.space_after=Pt(3)
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Light Grid Accent 1"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        rr=t.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold=True; rr.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cp=cells[i].paragraphs[0]; rich(cp, str(v))
            for rn in cp.runs: rn.font.size=Pt(9)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# --- masthead ---
rh=doc.add_paragraph(); rh.paragraph_format.space_after=Pt(2)
r=rh.add_run("BUTLER–VOSS CONDENSATE PROJECT   ·   COMPLETE WORKING REPORT   ·   WP-46")
r.font.size=Pt(8.5); r.font.color.rgb=GREY; r.font.name="Consolas"
tp=doc.add_paragraph(style="Title"); tp.add_run("The Butler–Voss Condensate")
sub=doc.add_paragraph(); sr=sub.add_run("Emergent Particles, Charges, and Forces from an Active Spatial Medium")
sr.font.size=Pt(13); sr.italic=True; sr.font.color.rgb=GREY; sub.paragraph_format.space_after=Pt(8)
bl=doc.add_paragraph(); br=bl.add_run("Robert Voss"); br.bold=True; br.font.size=Pt(10.5)
bl.add_run("   ·   Independent research, Butler–Voss Condensate Project").font.size=Pt(10.5)
dl=doc.add_paragraph(); dr=dl.add_run("Draft — 21 July 2026   ·   toy-model study, computational")
dr.font.size=Pt(9); dr.font.color.rgb=GREY; dl.paragraph_format.space_after=Pt(10)

# --- abstract ---
ab=doc.add_paragraph(); ab.paragraph_format.space_after=Pt(3)
abh=ab.add_run("ABSTRACT"); abh.bold=True; abh.font.size=Pt(9.5); abh.font.color.rgb=GREY
A=[
"The Butler–Voss condensate is a physics-inspired model in which space is treated as an active "
"medium — a condensate — sampled by a network of mobile nodes carrying a complex order-parameter "
"field. This report collects the project end to end: the emergence of persistent particles and their "
"quantum numbers, the behaviour of those charges on a self-organizing medium, the emergence of gravity, "
"and a systematic account of what sets the range of a force.",
"Nonlinear focusing produces persistent localized excitations (oscillons). Promoting the field to a "
"complex order parameter, charge appears as an integer topological winding and spin as a conserved "
"Noether charge, and the stable species reduce to a handful (n = 0, ±1, plus the Noether family). "
"Opposite charges bind into composites, and a cloud of mutually attracting nodes self-assembles into an "
"isotropic close-packed lattice that sets its own spacing — on which the topological charge survives the "
"medium moving and reconnecting beneath it. A variational density coupling then makes field energy "
"compress the medium and slow its waves, so excitations refract toward mass: two masses attract, energy "
"conserved. This emergent gravity is, however, short-ranged (Yukawa, λ ≈ 3) in both two and three dimensions.",
"A sequence of controlled experiments identifies the cause — screening is a mass term on the mediating "
"field, supplied by the medium's pinning to its rest spacing — and the escape: a source protected by a "
"conservation law. The model's conserved topological charge, carried by a massless Goldstone mode, gives a "
"genuinely long-range force; gauging the same symmetry restores screening (Meissner), penetration depth "
"λ_L ∝ 1/e, confirmed both statically and from vortex motion. In three dimensions the topological defect "
"becomes a line (long-range interaction per unit length) or, in a larger target space, a point hedgehog (a "
"genuine 3D point charge, but linearly self-energetic and short-range); gauging the point charge yields a "
"magnetic monopole whose field is a true 1/r² Coulomb — electromagnetism-in-3D. From one medium and one "
"field promoted step by step, the model reproduces the qualitative menu of physical forces, with the reach "
"of each set by whether a symmetry protects its mediator from a mass term. Every claim rests on "
"energy-conserving dynamics with explicit numerical correctness gates; the model is a toy, and each idea is "
"reduced to a measurement.",
"Finally, the same discipline is turned on whether the model could be fundamental. The obstacles usually "
"fatal to a 'space is a medium' theory are met one at a time: the full Lorentz group emerges at long "
"wavelength (isotropy, a single universal cone, and boosts, with violations suppressed as (E/E_Planck)^2); "
"relativistic Dirac fermions emerge on a bipartite medium and a single chiral fermion binds to a domain wall "
"(evading Nielsen-Ninomiya); and the linear sector quantizes to a proper relativistic quantum field theory. "
"Underlying several of these is a single principle: anything bolted on beside the structure "
"brings its own light cone, while anything made of the structure inherits it. Taken to its conclusion, the "
"medium's own bond fluctuations, read off the fermion dispersion, are an emergent photon (the Dirac-node "
"position) and an emergent graviton (the cone's tetrad) — so fermions, electromagnetism and gravity share one "
"cone by construction. Even quantum mechanics proves largely emergent — the Schrödinger wave, ℏ as a material "
"property, de Broglie's λ = h/p, and the Born rule as a stochastic attractor (given the wave, at the medium's own ℏ) "
"all follow from the condensate; decoherence in the medium's own phonon bath then einselects a pointer basis and "
"Born-weighted branches at a rate ∝ (separation)², leaving only the selection of a single definite outcome — the "
"measurement problem's hard core — as a postulate. Cast against experiment, the model makes one specific, falsifiable "
"prediction: a Planck-suppressed, species-universal, quadratic and subluminal Lorentz violation whose falsifiable "
"content is robust across the microscopic lattice rather than a tuned coefficient. Its universality half passes real "
"data twice — the GW170817 and TXS 0506+056 multi-messenger coincidences — and its coefficient is brought, for the "
"first time, into genuine contact with data: a proper dimension-six threshold analysis puts the bare model in "
"roughly one order-of-magnitude tension with the observed GZK cutoff, and then computing the composite-proton "
"suppression the coefficient must pass through pulls the effective value down onto the bound — ξ_eff ≈ 0.5–1.5 times "
"the GZK limit once real parton moments and their uncertainties are folded in, straddling the exclusion boundary "
"rather than clearing it. The model's one empirical claim thus survives every test it has faced and sits exactly at "
"the edge of ultra-high-energy sensitivity — its first real brush with falsification, and a concrete target for the "
"next round of cosmic-ray data. Gravity required one retraction before it yielded. Every elastic route fails on a measurement — the "
"topological/curvature sector is unshieldable but its like charges repel with a force that grows with distance, and "
"the tetrad graviton has a long-range 1/r^2 field yet is shieldable and exerts no force at all (Eshelby-Crum). The "
"resolution came from applying the paper's own range principle to gravity's mediator for the first time: it is the "
"condensate's amplitude mode, which is unprotected and therefore gapped — precisely why gravity always appeared "
"screened — and which couples monopolarly to positive-definite energy, so that scalar exchange between like charges "
"attracts. Its range is measured to be the inverse gap (λ m_A = 1.00), and the potential is exactly a screening "
"exponential times a 1/r Newtonian core; at criticality the exponential goes to unity and Newton's law survives "
"alone, universally attractive. This is scalar (Nordström) gravity, and an exhaustive arc of measurements then shows "
"it is not a way-station to Einstein: a static mass sources the Newtonian time potential but never the spatial "
"curvature that bends light, in every channel tested — smooth loop, elastic body force, compression, and the "
"topological disclination channel, in the continuum and directly on the lattice — so the light-bending parameter γ "
"is zero at every scale, and general relativity would require a background-independent construction rather than a "
"refinement of this one. The barriers usually fatal to a 'space is a medium' theory are surmountable one at a time; "
"the honest ceiling is that most results are internal consistency or the reproduction of known physics, that the "
"single-outcome problem, the specific Standard-Model group and the observed cosmological constant remain open, and "
"that the reach beyond reproduction is the one structurally-robust, still-unfalsified Lorentz-violation prediction.",
]
for para in A:
    p=body(para, size=10); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.right_indent=Inches(0.25)
kw=doc.add_paragraph(); kr=kw.add_run("Keywords — emergent gravity · topological charge · Noether charge · "
"self-assembly · Goldstone mode · Abelian Higgs · magnetic monopole · emergent Lorentz invariance · "
"Dirac/domain-wall fermions · elasticity-fracton duality · lattice field theory")
kr.font.size=Pt(9); kr.italic=True; kr.font.color.rgb=GREY
kw.paragraph_format.left_indent=Inches(0.25); kw.paragraph_format.space_after=Pt(10)

# ===== 1 Introduction =====
heading("1  Introduction", 1)
body("The Butler–Voss condensate is a deliberately simple, deliberately falsifiable model of an active "
"spatial medium. Space is represented not as an empty backdrop but as a substance — a condensate — sampled "
"by a network of nodes. A node is a physical locus that carries field state, not a bare graph vertex; the "
"relations between nodes are real, tension-bearing, and dynamical. Disturbances of the field propagate as "
"ripples; nonlinearity can self-trap them into persistent localized structures; and the medium's own "
"organization is the candidate source of the quantum numbers and forces one would like a fundamental theory "
"to produce rather than postulate.")
body("The guiding discipline is that every conceptual claim is turned into a measurement, and each result is "
"reported as confirmed, refuted, or qualified by the numbers. The report proceeds in five movements: the "
"particle sector and its charges (§3); those charges on a self-organizing medium (§4); the emergence of "
"gravity and the discovery that it is short-ranged (§5); a systematic account of what sets the range of a "
"force (§6); and the extension of that account to three dimensions and to genuine electromagnetism-like "
"behaviour (§7). The organizing result of the whole program is that the range of a force is governed by "
"whether a symmetry keeps its mediator massless.")

# ===== 2 Model & methods =====
heading("2  The model and methods", 1)
body("The medium is a set of nodes at positions Xi carrying a scalar or complex field. Node interactions and "
"self-assembly are governed by a Lennard-Jones potential with equilibrium spacing R0; a cooled relaxation "
"from a disordered cloud yields a close-packed lattice (hexagonal in 2D, face-centred cubic in 3D) that sets "
"its own spacing. A symmetric, smooth pair weight defines a local density and a meshfree field operator,")
add_eq("W_{ij} = exp(-r_{ij}^{2}/2h^{2}),     ρ_{i} = Σ_{j} W_{ij},", "1")
body("The gravity work descends from a single energy functional whose two couplings are exact gradients, so "
"energy is conserved and the field operator is automatically symmetric (stable):")
add_eq("E = 1/2 Σ π_{i}^{2} + 1/4 Σ_{ij}(γ_{i}+γ_{j}) W_{ij}(u_{i}-u_{j})^{2} + 1/2 m^{2} Σ u_{i}^{2} + 1/2 Σ |X'_{i}|^{2} + U_{LJ},", "2")
body("with γ = g(ρ) a bounded, decreasing function of local density (denser = slower waves). A single "
"knob β sets its sharpness; β = 0 switches the density response off and provides the control. The "
"engine is dimension-agnostic — the same pairwise operators run unchanged in 2D and 3D.")
body("Methodologically, weak forces are read from static energy or response rather than motion (dynamic "
"drift probes are swamped by dispersion); every new operator is checked by an explicit correctness gate "
"appropriate to it (gauge invariance, analytic-vs-finite-difference gradients, sparse-vs-dense agreement, "
"topological-charge quantization); and every range claim is decided by system-size scaling, because a "
"screened and a long-range quantity are indistinguishable over a single finite window.")

# ===== 3 Particles & charges =====
heading("3  Particles and charges (H1–H10)", 1)
body("The base engine is a displacement field on fixed lattices. Nonlinear focusing produces a persistent "
"oscillon that survives a full run while an unfocused control disperses (H1); a first tension-advection "
"attempt at gravity mostly drains energy and is abandoned (H2); and isotropic lattices are stable while an "
"anisotropic cubic lattice is artifact-prone (H4). Promoting the field to a complex order parameter with a "
"Mexican-hat vacuum lets the medium carry two kinds of charge: an integer topological winding (H6) and a "
"continuous, conserved Noether charge from the U(1) symmetry (H7). A census of candidate species finds the "
"stable set reduces to a handful — n = 0, ±1 plus the Noether family (H8). Opposite windings bind into "
"composites while like windings repel, an n = 2 defect splits, and a scale gap of order 6.6× separates the "
"layers (H9). Finally, a cloud of mutually attracting nodes self-selects an isotropic close-packed lattice "
"and sets its own spacing (H10) — which is why a self-organizing medium avoids the cubic-lattice artifact "
"entirely.")
table(["id","Claim","Result"],
 [["H1","focusing makes persistent particles","supported (oscillon lives; control disperses)"],
  ["H2","gravity via tension-advection","not supported (leaky / disruptive)"],
  ["H3","a particle is a moving pattern","weak (moves in the push direction)"],
  ["H4","results independent of lattice","artifact-sensitive (cubic anisotropy)"],
  ["H5","2D and 3D differ","supported (both host oscillons)"],
  ["H6","charge = topological winding","confirmed (integer, conserved)"],
  ["H7","spin = Noether charge","confirmed (Q-ball)"],
  ["H8","only ~3–5 fundamentals","derived: n = 0, ±1 + Noether"],
  ["H9","higher layers by confluence","opposite bind; n=2 splits; scale gap ~6.6x"],
  ["H10","self-cohering medium","self-assembles to isotropic close packing"]],
 cap="Table 1.  The particle-sector hypotheses and their measured outcomes.")

# ===== 4 Field on the medium =====
heading("4  Charge on a self-organizing medium", 1)
body("Carrying the complex field on the mobile, self-assembled medium tests whether the quantum numbers "
"survive a substrate that is itself dynamical. On an irregular mesh the naive graph Laplacian degrades "
"badly, so an accurate least-squares meshfree operator is required (P0–P1); with it, a vortex's topological "
"charge is conserved on the self-assembled medium. It then survives the medium moving and reconnecting "
"beneath it — hundreds of reconnection events — and fails only when the medium melts (P2). Topology, in "
"other words, is robust to the substrate rearranging, exactly as a genuine conserved quantum number should "
"be.")

# ===== 5 Gravity =====
heading("5  Emergent gravity — and its range", 1)
body("The non-leaky replacement for the abandoned tension-advection gravity is gravity-by-density: field "
"energy compresses the medium, a denser medium slows its own waves, and other excitations refract toward "
"the compressed region. Built from the variational functional of §2 (energy-conserving by construction, "
"with a bounded g(ρ) that cures a blow-up), the mechanism passes the existence test — two field lumps "
"drift together (separation 9.3 → 7.5), with no attraction at β = 0 and stronger attraction at larger "
"β, energy conserved throughout. A single lump does not self-bind, however: the medium is nearly "
"incompressible, so the force is real but weak.")
body("The character of that force is measured statically, from the medium's density response to a frozen "
"source. Scaling to a large ordered fcc medium (N = 9213) to reach far from both source and free surface, "
"the response is unambiguously exponential:")
add_figure(figs["drho"], "1",
 "Static density response of a large ordered fcc medium (N = 9213). An exponential fit (Σ-of-squares 0.011) "
 "beats the best power law (0.10) by an order of magnitude: gravity-by-density is exponentially screened, "
 "λ ≈ 3.3, comparable to the 2D value (5.7). Adding a dimension does not lengthen it.")
result("Result 5.", "Emergent gravity is short-ranged (Yukawa, λ ≈ 3) in both 2D and 3D — real but "
"contact-like, not 1/r². The behaviour matches the Bitter–Crum theorem: two centres of dilatation in an "
"isotropic elastic medium have no long-range interaction. An earlier apparent long range (λ ≈ 14) was a "
"finite-size illusion, corrected by system-size scaling.", warn=True)

# ===== 6 Force range from symmetry =====
heading("6  Force range from symmetry", 1)
body("Why is gravity screened, and is the screening escapable? Four experiments answer this.")
body("First, coupling strength is not the lever (Screen-0): sweeping β lengthens the range slightly but "
"collapses the amplitude and the total compression together — a trade, never a strong long-range force.")
table(["β","range λ","peak amplitude","integrated compression"],
 [["20","3.2","0.073","103"],["40","3.7","0.059","96"],["60","4.3","0.036","63"],["100","5.8","0.016","32"]],
 cap="Table 2.  Screen-0: density response vs coupling sharpness. Range rises while amplitude falls — a trade.")
body("Second, screening is a mass term (Screen-1). Solving a discrete Poisson equation for a point source, "
"massless vs massive, and growing the box: a massless (1/r) field has no intrinsic length so its apparent "
"range climbs without bound, while a massive field pins at 1/m.")
add_figure(figs["box"], "2",
 "Screen-1: apparent range vs system size for the massless and massive Poisson solves. The massless field "
 "never settles — the fingerprint of a true 1/r tail — while the massive field locks to a fixed length. "
 "Screening is exactly the mass term; the medium's pinning to R0 supplies it.")
body("Third, a conserved charge evades the mass term (Screen-2). The model's topological winding is carried "
"by the field's phase, a Goldstone mode that is massless by symmetry. In 2D a neutral vortex pair's energy "
"grows as a Coulomb logarithm with a box-independent slope while any fitted screening length merely tracks "
"the box — a genuinely long-range force. Fourth, gauging that symmetry (the Abelian Higgs model) hands the "
"mediator a mass by the Meissner effect, and the vortex interaction becomes screened at an intrinsic "
"penetration depth λ_L that scales as 1/e:")
add_figure(figs["gauge"], "3",
 "Gauged U(1): Meissner penetration depth vs gauge coupling. e*λ_L is constant (λ_L ~ 1/e), and the "
 "length is intrinsic — box-independent — unlike Screen-2's box-growing log. The same length is recovered "
 "from moving vortices in an overdamped dynamical measurement, confirming the screened force from motion.")
result("Result 6 — a menu of forces.", "Nothing changes between these outcomes but the symmetry structure "
"of the source: a conserved (topological) charge with a massless Goldstone mediator gives a long-range force "
"(EM-like); energy density, whose strain is pinned to R0 (a mass), gives a short-range screened force "
"(gravity/nuclear-like); gauging supplies a Meissner mass on demand for a tunable short-range force. Reach is "
"set by what protects the mediator from a mass term.")
table(["Source → mediator → range","Behaviour","Analog"],
 [["conserved topological charge → massless Goldstone → 1/r (log in 2D)","unscreened, long range","EM-like"],
  ["energy density → strain pinned to R0 (massive) → exp(-r/λ)","screened, contact-like (λ ~ 3)","gravity / nuclear-like"],
  ["gauged charge → Meissner-massive photon → exp(-r/λ_L)","screened, tunable (λ_L ~ 1/e)","superconductor-like"]],
 cap="Table 3.  One field, three interaction archetypes, separated only by the symmetry of the source.")

# ===== 7 Three-dimensional program =====
heading("7  The three-dimensional program", 1)
body("The 2D long-range demonstration used a point vortex — the U(1) defect in 2D. Its 3D fate depends on "
"the field's target space, and three routes complete the picture.")
body("Route A keeps the single complex field, whose 3D defect is a line. A straight vortex line is a genuine "
"3D object (energy proportional to its length, winding threading every transverse slab), two antiparallel "
"lines interact by a box-independent Coulomb logarithm per unit length (the long-range force survives into "
"3D), and a vortex ring is a closed 3D line whose energy is tension times circumference. Route B enlarges "
"the target space to a three-component unit vector (the O(3) model), where a point defect — the hedgehog — "
"carries an integer charge: a genuine 3D point 'particle'. But the global hedgehog is a global monopole, its "
"self-energy diverging linearly with system size, and a neutral pair's interaction is finite, localized, and "
"short-ranged — not confining and not 1/r. The clean long-range point charge requires gauging.")
body("Route C gauges the charge: a compact U(1) gauge field in its Coulomb (deconfined) phase, where the "
"topological charge is a magnetic monopole (quantized flux out of a cube) and minimizing the Maxwell energy "
"is the massless Gauss law. The gauged monopole has a finite, box-independent self-energy (deconfined, "
"unlike Route B), and its field is a genuine 1/r² Coulomb:")
add_figure(figs["mono"], "4",
 "Route C: the relaxed magnetic monopole's radial field on log axes. In the clean interior the field follows "
 "|B| ~ 1/r^{2.03} — a textbook Coulomb law (the flattening at large r is the finite-box floor). A genuine "
 "1/r² force between quantized topological point charges: electromagnetism-in-3D.")
result("Result 7.", "The 3D program mirrors the 2D story exactly. A conserved charge gives a long-range force "
"(Route A line; Coulomb log per length); a bare point charge is self-energetic and short-range (Route B "
"global hedgehog); and gauging deconfines it into a genuine 1/r² Coulomb (Route C monopole). Broken-phase "
"gauging gives Meissner screening (§6); Coulomb-phase gauging gives unscreened 1/r² — the two gauge phases, "
"both realized.")
table(["Route","Defect","Outcome"],
 [["A","vortex line (S¹, 3D)","long-range Coulomb log per unit length; ring = tension × circumference"],
  ["B","hedgehog point (S², 3D)","genuine integer point charge; global monopole — divergent, short-range"],
  ["C","gauged monopole (U(1))","deconfined: finite self-energy, genuine 1/r² Coulomb (EM-in-3D)"]],
 cap="Table 4.  The three-dimensional routes to a point charge and their force laws.")

# ===== 8 Toward fundamental physics =====
heading("8  Toward fundamental physics", 1)
body("The program so far asks whether the medium can host the phenomena of physics. A sharper question is "
"whether it could be fundamental. A medium has a preferred frame (the nodes' rest frame), so the model is "
"viable as fundamental physics only if the symmetries of nature emerge at long wavelength and the known "
"obstacles are met rather than assumed away. Each barrier is turned into a measurement.")

heading("8.1  Emergent Lorentz invariance", 2)
body("Lorentz symmetry is the first-order threat: it is tested to extraordinary precision, and a medium "
"generically violates it. Three facets, all read from the exact lattice dispersion. isotropy: on the "
"self-assembled isotropic lattices the field wave speed becomes direction-independent as k → 0, the "
"anisotropy (the Lorentz violation) vanishing as a power of k/kmax, with the lattice spacing as the model's "
"Planck scale.")
add_figure(figs["lorentz"], "5",
 "Emergent Lorentz invariance: the directional anisotropy of the wave speed (the Lorentz violation) vanishes "
 "at low energy as (k/kmax)^2 (fcc 3D) or (k/kmax)^4 (hex 2D) — i.e. as (E/E_Planck)^{2..4}. Rotational "
 "invariance is emergent; the lattice spacing plays the role of the Planck length.")
body("universality: a stable elastic medium has c_L > c_T (two cones) and the field adds a third; but governing "
"medium and field by one isotropic (vector-Hooke) operator on the self-assembled geometry collapses them to a "
"single universal cone (c_L = c_T = c_field exactly). boosts: applying a Lorentz boost with the emergent c to "
"the dispersion, the massless cone and the massive mass-shell map to themselves up to a residual ~ (k/kmax)^2, "
"ω^2 - c^2 k^2 is a genuine Lorentz invariant, and the front velocity never exceeds c. Isotropy, "
"universality and boosts together: the full Lorentz group emerges at long wavelength, violations suppressed as "
"(E/E_Planck)^2 — the standard emergent-relativity story, here measured rather than assumed. This is, however, "
"a within-sector statement; adding fermions introduces a further cone, and Section 8.5 shows what closes it.")

heading("8.2  Emergent relativistic fermions", 2)
body("Matter is fermions — spin-1/2 and chiral — while the model's fields are bosonic. Relativistic fermions can "
"emerge near a band-touching point: a tight-binding model on a bipartite (honeycomb) medium has a linear, "
"isotropic Dirac cone (Fermi velocity v_F = 3/2), whereas the plain close-packed (triangular) medium has an "
"ordinary quadratic band and no cone. The two Dirac points carry opposite chirality and cancel (Nielsen-"
"Ninomiya doubling), so a single chiral fermion seems forbidden — until the standard escape: a Wilson-Dirac "
"(Chern) strip binds a single chiral fermion to each edge (a domain wall to the trivial vacuum), its opposite-"
"chirality partner spatially separated to the far edge. So a lattice that is vector-like overall carries a "
"single chiral fermion on a wall — the mechanism a fundamental version would use for Standard-Model chirality.")

heading("8.3  Quantization", 2)
body("The unified linear sector is coupled harmonic oscillators — a free field. Canonical quantization gives "
"bosonic quanta whose energies lie on the relativistic mass-shell, and whose quantum vacuum correlator is the "
"relativistic form: a power law for a massless field, a Yukawa exponential for a massive one — the same "
"massless/massive, long/short-range dichotomy of the forces program, now at the level of the vacuum. This "
"shows the model quantizes to a proper relativistic quantum field theory. Honestly, though, canonical "
"quantization imposes the commutation relations; it does not derive quantum mechanics from the sub-quantum "
"medium. That deeper question — whether the quantum wave, ℏ, the Born rule and the guidance of a particle "
"by its wave arise from the condensate's own mechanics — is taken up directly in Section 8.6.")

heading("8.4  Gravity — the sharpest barrier (and a route that did not survive)", 2)
result("Retraction (2026-07-13).", "This section originally concluded that the elasticity-fracton duality "
"provided a concrete Route past the gravity screening. That conclusion was never backed by a measurement of the "
"force between two curvature charges, and when the force was finally measured it failed. Section 8.9 reports the "
"measurements and the corrected position: long-range emergent gravity in this model is open, not solved. The "
"multipole analysis below stands as stated; the inference drawn from it did not.", warn=True)
body("Real gravity is long-range, universally attractive, and spin-2. The model's gravity-by-density is none of "
"these at long range: it couples to energy density, and the medium's displacement field is a vector whose "
"phonons are spin-0 + spin-1, with no spin-2 mode. But the two failures share one cause and one cure. In 2D "
"linear elasticity every defect sources the same biharmonic equation, differing only by the multipole order of "
"the source, which fixes its range: a dilatation (energy density, the gravity-by-density coupling) is the "
"most-screened multipole — a contact term, Bitter-Crum-screened; a dislocation is logarithmic; a disclination "
"(curvature) is the least-screened, genuinely long-range. By the elasticity-fracton duality these defects are "
"the charges of a rank-2 symmetric-tensor gauge theory — the structure of linearized gravity — and in 2D a "
"point mass is precisely a conical deficit, i.e. a disclination. So the medium already contains long-range "
"'masses' that curve space around them; one must couple gravity to curvature, not energy density.")
body("That tensor sector is a genuine graviton: the transverse-traceless field has exactly two polarizations, "
"and they carry helicity plus/minus 2 (they rotate by twice the angle of a rotation about k, versus a photon's "
"helicity plus/minus 1) — the 'plus' and 'cross' of a gravitational wave, with a universal 1/r^2 attraction "
"and the light-bending factor of two that distinguishes spin-2 from scalar gravity. Both halves of long-range "
"spin-2 gravity are thus present in the medium's structure. Section 8.7 makes this graviton dynamical — a "
"propagating, luminal gravitational wave whose kinetic term is induced from the fermion loop — leaving the fully "
"self-consistent 3D back-reaction, and clearing the Weinberg-Witten theorem (whose loophole the model's "
"cutoff-scale Lorentz violation already opens), as the open work.")
heading("8.5  One structure: cone universality, the emergent photon and the graviton", 2)
body("The Lorentz result of Section 8.1 is a within-sector statement — one round universal cone for the medium "
"and the bosonic field. Adding fermions exposes the gap. On the same honeycomb both cones are computable: the "
"fermion cone from the slope of the Dirac bands, v_F = (3/2) t a, and an independent spring-boson cone from the "
"acoustic branch, c_B = (sqrt3/2) √(K/m) a. Their ratio, v_F / c_B = √(3) t / √(K/m), is an arbitrary, "
"tunable number fixed by independent couplings. Equality is a fine-tuning, not a symmetry — so generically there "
"are two light cones, i.e. Lorentz violation between statistics. (Real graphene is the cautionary case: only its "
"fermion sector is even approximately relativistic.)")
body("The cure is structural, and it is the principle the whole program keeps rediscovering. Do not put the boson "
"in by hand; let it be a collective mode of the fermions, and it has no cone of its own. Every composite "
"(particle-hole) boson of momentum q costs at least the lower edge of the interband continuum, "
"ω_min(q) = min_k [E_+(k+q) - E_-(k)] = v_F |q| — measured on the honeycomb bands as 0.995 v_F|q| at "
"|q| = 0.02 and tending to v_F|q| as q → 0, isotropically. The effective-action version is Sakharov's: a boson "
"induced by integrating the fermions out inherits their Lorentz invariance. Its one-loop polarization must then "
"depend on frequency and momentum only through the invariant s = Ω^2 + v_F^2 q^2 — and it does: the quantity "
"(Π / q^2) √(s) is the same to 0.10% whether the invariant is carried by momentum or by frequency, the "
"residual shrinking with energy exactly as (E/E_Planck)^2.")
body("The construction that realizes this is Volovik's, and the model's own medium supplies it. Near a Dirac node "
"the fermion Hamiltonian is H = e^a_i σ_a (k_i - A_i): the position of the node, A, is an emergent U(1) gauge "
"field — a photon, spin-1 — and the shape of the cone, the tetrad e, is an emergent metric — a graviton, spin-2. "
"Perturbing the medium's own three nearest-neighbour bonds and reading the fermion bands, a uniform bond stretch "
"leaves the node fixed and the cone round (a pure conformal rescaling: no photon, no graviton), while the doublet "
"bond fluctuations shift the node and deform the cone anisotropically, sourcing both fields at once, each a clean "
"linear response. The photon and the graviton are therefore not added to the medium: they are the medium's bond "
"fluctuations, seen by its fermions. Because both are read off the fermion dispersion, neither carries a light "
"cone of its own, and cross-statistics Lorentz invariance holds by construction.")
result("Result 8.5 — one structure.", "Anything bolted on beside the structure brings its own light cone; anything "
"made of the structure inherits it. This single principle resolves the field-versus-medium cone mismatch, the "
"boson-versus-fermion mismatch, and — in the same stroke — produces electromagnetism and gravity as the medium's "
"own bond fluctuations. The emergent graviton is the tetrad: the spin-2 sector that Section 8.4 located in the "
"medium's defects.")

heading("8.6  The origin of quantum mechanics", 2)
body("Section 8.3 quantized the medium but imposed the quantum rules. The deeper question is whether quantum "
"mechanics itself — the wave, ℏ, probability, and the link between a particle and its wave — is mechanics of "
"the condensate. Quantum mechanics has two halves, and they are not equally hard: the wave half (the Schrödinger "
"equation, superposition, interference and ℏ) and the probabilistic half (the Born rule and measurement). The "
"medium is a condensate, so its natural language is hydrodynamics, and Madelung's theorem is that Schrödinger's "
"equation is the hydrodynamics of a fluid with one special gradient energy. The emergent massive field the "
"project already has, written χ = √(ρ) exp(i S / ℏ), is that fluid.")
body("wave half. The massive field's dispersion ω = √(c^2 k^2 + Ω^2) has, in the slow (non-relativistic) "
"limit, ω ~ Ω + (c^2 / 2 Ω) k^2 — the free-Schrödinger dispersion, with the diffusion constant fixed by "
"the medium's own gap and speed:")
add_eq("ℏ / 2m = c^{2} / 2 Ω", "8.6a")
body("So a slow wave packet of the emergent field must spread at the exact Schrödinger rate, and ℏ is a material "
"property of the medium — its gap-to-dispersion-curvature ratio — not a postulate. Evolving the complex field for a "
"Gaussian packet, the measured spreading gives a diffusion constant of 0.830 against the 0.833 predicted from the "
"gap alone (0.4%). Superposition and interference are automatic, since it is a linear wave. Half of quantum "
"mechanics — the wave mechanics and ℏ — is free.")
body("probabilistic half. That |ψ|^2 is a probability is the genuinely hard part, and linear wave dynamics do not "
"supply it. Nelson's construction supplies the mechanism, and the condensate supplies its one scale: |ψ|^2 is the "
"equilibrium of a diffusion whose noise is set by the medium's own fluctuation strength (the same ℏ/2m derived in "
"8.6a) and whose osmotic drift, u = (ℏ/2m) d/dx ln|ψ|^2, is the ordinary entropic force down a density gradient. The "
"decisive, non-trivial claim (Valentini's quantum relaxation) is that |ψ|^2 is not merely consistent but an "
"attractor. Started in a wrong (uniform, non-Born) ensemble, the density relaxes to |ψ|^2 — the Kullback-Leibler "
"divergence falling from 4.39 to 10^-4 for a ground state and from 2.82 to 2×10^-4 for a structured interference "
"density. So |ψ|^2 is the unique equilibrium probability, an attractor rather than a postulate — with one boundary "
"stated the same way in the underlying test: the osmotic drift is read from |ψ| itself (the guiding wave, as in "
"Nelson/Bohm), so the wavefunction is still needed and it is only its modulus-squared that is derived. Deriving the "
"drift from the sub-quantum medium, rather than reading it off |ψ|, is not done here; what the medium genuinely "
"fixes is the diffusion scale ℏ/2m, and the Born statistics then follow as that diffusion's equilibrium.")
body("the guidance, and the honest boundary. What still linked the two halves by hand was the guidance rule — that a "
"particle moves with its wave, v = ∇(S)/m. de Broglie's double solution offers to derive it from one field: far "
"from the particle the smooth pilot wave, at the particle a localized soliton core of the same field. Realized in "
"the medium's non-linear Schrödinger limit, whose bright soliton is a legitimate particle, the test splits sharply. "
"A soliton carrying its own phase exp(i k x) drifts at exactly v = k across every wavenumber (slope 1.000): "
"de Broglie's λ = h/p is therefore a theorem of the medium, the envelope being Galilean-covariant, not a "
"postulate. But a resting soliton is not steered by a separate pilot wave (slope ~ 0): at the particle's core the "
"total phase is dominated by the particle's own flat phase, so nothing guides it, and the non-linearity only "
"scatters it weakly. de Broglie's full double solution — a soliton phase-locked to and steered by a distinct pilot "
"wave — is not realized by naive non-linear superposition, the very step he never rigorously closed.")
result("Result 8.6 — most of quantum mechanics is condensate mechanics.", "Emergent from the medium alone: the "
"Schrödinger wave and ℏ as a material property (8.6a), and de Broglie's v = ∇(S)/m for a particle's own wave. "
"Emergent given the wave: the Born rule as the unique equilibrium of a Nelson diffusion whose scale is the medium's "
"ℏ/2m — a Valentini attractor, not a postulate, though its osmotic drift is read from |ψ| rather than derived from "
"the sub-quantum medium (the wavefunction is still needed; only its modulus-squared is derived). not emergent: the "
"guidance of a particle by a separate pilot wave and the "
"selection of a single definite outcome — the hard core of the measurement problem, which stays a postulate here as "
"it does everywhere. The honest boundary is clean: most of quantum mechanics is mechanics of the condensate; the "
"residue is exactly the piece that is unsolved for everyone.")

heading("8.7  A dynamical graviton, and its induced kinetic term", 2)
body("Section 8.4 located the long-range spin-2 sector in the medium's curvature defects and Section 8.5 identified "
"the graviton with the tetrad — the shape of the fermion cone. Here that graviton is made dynamical, and its "
"equation of motion is shown not to be imposed. Evolving the transverse-traceless field as a wave packet, both "
"polarizations — the 'plus' and 'cross' of a gravitational wave — propagate identically at a group velocity of "
"0.970 c: a massless, luminal graviton, the ~3% deficit being the same (E/E_Planck)^2 lattice dispersion measured "
"everywhere else. Its Lorentz-violation coefficient, ζ_graviton = 0.250, has the same (E/E_Planck)^2 form as the "
"boson, fermion and photon: on the common lattice every excitation shares the one operator, so the graviton rides "
"the single universal cone. In the static limit the same field gives a 1/r^2 attraction — the potential measured as "
"1/r^1.13, the small excess a finite-box artefact tending to 1/r in the continuum. That last statement, however, "
"must be read with the caveat that Section 8.9 makes unavoidable: the inverse-square law here follows from a Poisson "
"equation that was imposed, with a mass-density source put in by hand. It is a consistency check on the tensor "
"sector's static limit, not a derivation of an attractive long-range force from the medium — and when that force was "
"measured directly, it was not there.")
body("The decisive step is that the wave operator just evolved is not put in by hand: it is induced, exactly as the "
"photon's Maxwell term was in Section 8.5. Integrating out the fermions, the interband stress-tensor correlator "
"<T T> on the low-energy Dirac cone must, if it is to give a Lorentz-invariant Einstein-Hilbert term, depend on "
"frequency and momentum only through the invariant")
add_eq("s = Ω^{2} + v_{F}^{2} q^{2}", "8.7a")
body("and it does: holding s fixed while shifting its frequency/momentum mix, the polarization Π_+ varies by only "
"5.3%, 1.4%, 0.37% as the energy √(s) is halved from 0.24 to 0.06 — the residual Lorentz violation vanishing "
"linearly in s. So the graviton's Einstein-Hilbert kinetic term is generated by the fermion loop (Sakharov's "
"induced gravity), carrying the fermion cone, not imposed. One caveat is honest and expected: the correlator also "
"has a large s-independent piece, Π_+(0,0) = 0.134 — an induced cosmological term, the cutoff-dependent vacuum "
"energy that is the well-known burden of every induced-gravity scenario, and the model's own cosmological-constant "
"problem.")
result("Result 8.7 — the graviton is dynamical and its dynamics are induced.", "The emergent graviton propagates as "
"a massless, luminal, spin-2 gravitational wave on the one universal cone, and mediates a universal 1/r^2 "
"attraction; its Einstein-Hilbert kinetic term is not imposed but induced from the fermion loop a la Sakharov, "
"carrying the fermion cone — so fermion, photon and graviton and all of their dynamics descend from the single "
"fermion structure. What remains open is sourcing the propagating tetrad from matter energy self-consistently in "
"full 3D (the non-linear back-reaction), and taming the induced cosmological term.")

heading("8.8  The model's first empirical prediction", 2)
body("Everything to this point shows the model can host known physics. A theory earns its keep, though, by "
"predicting something nature could use to kill it. Read as a real crystal, the emergent medium violates Lorentz "
"invariance at its cutoff in a specific, computable way, and that signature is the project's first empirical claim. "
"From the fcc dispersion the two leading coefficients are a boost term, 1 - v/c = 0.245 (k/k_max)^2, and a "
"crystallographic rotation term, dc/c = 0.068 (k/k_max)^2 — both quadratic in energy (mass-dimension-6, n = 2), "
"the rotation part carrying the lattice's cubic angular pattern. Cast as a modified dispersion,")
add_eq("v(E)/c = 1 - ζ (E/E_{Planck})^{2}", "8.8a")
body("with ζ of order unity, this is an effective quantum-gravity scale E_QG,2 = E_Planck / √(ζ) ~ "
"2.5 × 10^19 GeV. Confronted with data it is safe by many orders of magnitude — the strongest current n = 2 bounds, "
"from Fermi-LAT γ-ray bursts and from ultra-high-energy cosmic rays, sit near 10^10-10^11 GeV, against the "
"model's 2.5 × 10^19 GeV — the closest frontier being UHECR, where the predicted speed shift is only |dv/c| ~ "
"1.6 × 10^-17. Crucially, it makes three qualitative predictions that do not need Planck-energy access and that "
"would falsify it: (1) the violation is quadratic (n = 2), so a confirmed linear photon dispersion kills it; "
"(2) the rotation-violating part is anisotropic with the emergent lattice's crystallographic pattern, correlated "
"between the boost and rotation sectors; and (3) there is one universal cone — no leading-order species-dependent "
"maximal speed — so a confirmed c_photon ≠ c_electron, or c_gravity ≠ c_light in the spirit of GW170817, at "
"order E/E_Planck would kill it.")
result("Result 8.8 — a falsifiable signature.", "The model's Lorentz violation is quadratic, Planck-suppressed, "
"cross-species-universal, and crystallographically anisotropic. It survives every current bound by many orders of "
"magnitude, yet it is falsifiable in structure — by a linear photon dispersion, by species-dependent cones, or by "
"a crystallographic anisotropy pattern — none of which requires reaching the Planck scale. This is the project's "
"first prediction, as opposed to a reproduction: a specific signature nature could rule out. tempered by Section "
"8.16: put against real numbers, the quadratic suppression puts this effect about 10^-16 below current sensitivity, "
"so while it is safe against every bound it is also not presently falsifiable — reaching it needs ~8 orders of "
"improvement, and today's experiments probe the linear signal the model does not predict. The project's genuinely "
"reachable prediction turns out to be gravitational, not photonic: the short-range γ of 8.11/8.16.")

heading("8.9  Shielding, and the retraction of the gravity claim", 2)
body("Gravity cannot be shielded. A slab of lead between you and the Sun changes nothing, and gravitational waves "
"cross the universe unattenuated. The tempting explanation -- that gravity is lossless, so it cannot be absorbed -- "
"is wrong on the mechanism, and the error is instructive. Screening is not dissipation: a superconductor screens a "
"magnetic field over the London depth with exactly zero loss, and a Yukawa field conserves energy exactly. Screening "
"is evanescent, not absorptive. The model's own screening never dissipated anything either, and it is still fatal. "
"The real principle is unneutralizability: every screening mechanism in physics -- Debye, the Faraday cage, the "
"Meissner effect -- works by the medium rearranging opposite-sign charge into a cancelling cloud, and mass is "
"unipolar. There is no negative mass to build the cloud from. You cannot screen what you cannot cancel.")
body("That converts into a sharp, box-independent test: place a source, let the medium do anything it likes in a "
"surrounding shell, and ask by a Gauss-law integral how much charge is still visible outside. A dilatation (energy "
"density) charge is cancelled exactly -- the medium's response is strictly local (div u = C s(x) to 7.5×10^-16, the "
"Bitter-Crum contact term), leaving only the source's own tail, 2×10^-10 of peak. It is neutralizable, therefore "
"shieldable, therefore short-range. A topological (curvature) charge is not: the charge seen outside stays exactly "
"+1.000000000 under a violent smooth deformation of the shell, a fifty-fold stiffer shell with full relaxation, a "
"fifty-fold softer one, and all of them at once -- because any single-valued response of the medium carries zero "
"winding, identically. Nucleating a genuine anti-defect does flip it to zero, so the probe is sensitive and the "
"invariance is physics; but that charge is quantized, so no infinitesimal screening cloud exists. topological "
"quantization does in the model exactly what 'there is no negative mass' does in nature.")
body("This looked like a decisive vindication of Route 1. It was not, because it tests only the screening half. The "
"force law between two curvature charges had never been measured -- Section 8.4's route was an inference from the "
"multipole hierarchy, and Section 8.7's inverse-square law came from an imposed Poisson equation with a mass-density "
"source, not from a disclination. Measured properly, in real space on a clamped disc (removing the periodic zero mode "
"that had previously saturated the calculation), and gated against box size, the result is fatal: two like "
"disclinations repel, with an interaction growing as R^1.97. The sign is wrong -- gravity is universally attractive, "
"and a disclination behaves like a charge, not like a mass -- and the range is wrong in direction, the force growing "
"with distance rather than falling as 1/r^2. In hindsight the earlier infrared saturation was this R^2 growth: the "
"calculation was not failing to find a force law, it was correctly refusing to converge on an unbounded one. Nor is "
"this a surprise once stated properly: in 2+1-dimensional general relativity a point mass is a conical deficit, and "
"2+1D gravity is topological -- no local propagating modes, no Newtonian attraction between static masses. The "
"curvature sector is reproducing 2+1D gravity faithfully, which is precisely why it cannot yield a Newtonian force.")
body("The tetrad graviton of Section 8.5 fares no better, and the manner of its failure corrects a "
"common misunderstanding. It does have one genuine success: it is long-range from ordinary energy density. Bitter-Crum "
"screens the trace of the medium's response to a mass, but not the shear -- the deviatoric strain falls as r^-2.01 "
"(box-gated) -- and the tetrad is precisely the traceless part of the cone deformation, so it reads the one sector "
"the screening theorem never touched. Gravity-by-density failed because it coupled to the trace. But a field falloff "
"is not A force. Two objections settle it. First, an intervening shell attenuates the tetrad at a probe outside it "
"by up to fourfold; any impedance mismatch scatters it, and real gravity shows no such attenuation. Second, and "
"decisively, the classical Eshelby-Crum theorem states that in an infinite isotropic elastic medium the interaction "
"energy of two centres of dilatation vanishes -- and measurement confirms it: the force between two masses collapses "
"by a factor of 69 between r = 10 and r = 22, saturating to nothing. What short-range attraction remains is merely "
"the contact term, i.e. the screened gravity-by-density of Section 5. Breaking the medium's isotropy makes a genuine "
"long-range force appear -- 46 times larger, and repulsive -- which proves the probe is sensitive exactly where the "
"isotropic case reads zero.")
result("Result 8.9 — the gravity claim, retracted.", "Every long-range gravitational candidate has now failed on a "
"measurement. The topological/curvature sector is unshieldable, but its like charges repel with a force that grows "
"with distance. The tetrad graviton has a genuine long-range 1/r^2 field, but it is shieldable and exerts no "
"long-range force (Crum). Gravity-by-density is genuinely attractive, but screened. The model has produced exactly "
"one real attraction between two masses -- the nonlinear gravity-by-density of Section 5, where two lumps drift "
"together -- and that one is short-ranged. The linear/elastic sector is provably the wrong place to look, since Crum "
"forbids a long-range force there; the attraction that does exist is nonlinear, and that is the lead worth "
"following. Long-range emergent gravity in this model is open, not 'a route found'. It is the project's outstanding "
"failure, and it is stated here as such.", warn=True)

heading("8.10  Gravity, solved — the amplitude mode", 2)
body("Four failures share one property, and naming it gives the answer. The elastic/displacement sector is "
"structurally dead for gravity, and now provably so: a mass is a force dipole, so Bitter-Crum makes its density "
"response a contact term for any choice of moduli, and Eshelby-Crum makes the force between two masses vanish "
"outright. No amount of tuning rescues it. So stop looking there — and instead apply the principle this project "
"established for every other force, but never once applied to gravity's: A force'S range is set by whether A symmetry "
"protects its mediator from A mass term.")
body("What, then, mediates gravity? Write the condensate as χ = (φ0 + η) exp(i θ) in the Mexican-hat "
"potential V = -(a/2) φ^2 + (b/4) φ^4, so that φ0 = √(a/b) and the amplitude gap is m_A^2 = V''(φ0) = 2a. "
"Two facts then settle everything. The phase is a Goldstone: its shift symmetry permits only derivative couplings, so "
"it can never mediate a monopole force — it is protected, and useless for gravity. The amplitude is not protected — a "
"radial mode never is — so it carries a mass; and it does couple monopolarly, since matter's energy density enters "
"honestly as g ρ |χ|^2, which contains a term linear in η. Therefore gravity in this medium is a Yukawa force "
"of range 1/m_A. That single sentence explains every 'gravity is screened' result the project ever obtained: the "
"amplitude mode was gapped. And the sign comes for free — energy density is positive-definite, and scalar exchange "
"between like charges attracts.")
add_eq("E_{int}(R) = - C exp(-R/λ) / R,     λ = 1 / m_{A}", "8.10a")
body("Measured in three dimensions (an inverse-square law requires them), on the full non-linear field, with the same "
"interaction-energy probe that had correctly returned 'no force' three times. Sweeping the gap, the fitted range "
"satisfies λ times m_A = 1.010, 1.005, 1.003, 1.004, 1.023 — with m_A read off the potential, not fitted to the "
"force — and the interaction is attractive at every point. The decisive step is the form check: divide the "
"exponential out, and what remains is flat in R to 0.1-0.5%, a pure R^-1.00 power law, for every gap. So (8.10a) "
"holds exactly: a screening exponential multiplying a 1/R Newtonian core. This is measured, not extrapolated. Send "
"the medium toward its critical point, m_A → 0, the exponential goes to unity, and Newton's law survives alone: "
"potential 1/r, force 1/r^2, universally attractive. A box gate confirms the range is physics and not periodic "
"wrap-around (λ m_A = 1.023 at N = 64 becomes 1.000 at N = 96).")
result("Result 8.10 — a working long-range gravity.", "Gravity is mediated by the condensate's amplitude mode. Being "
"unprotected it is gapped, which is precisely why every earlier measurement found gravity screened; being a scalar "
"coupled to positive-definite energy, its exchange between like charges is universally attractive, with no sign put "
"in by hand. Near criticality the mediator becomes massless and the force becomes exactly Newtonian. Honest ceiling: "
"this is scalar (Nordström) gravity — it gives Newton's law but not light bending, and not two-polarization "
"gravitational waves. Full general relativity requires a massless spin-2 field protected by diffeomorphism "
"invariance (Weinberg's uniqueness theorem), and a fixed-background medium has no diffeomorphism invariance — so GR "
"itself remains out of reach. It also requires the medium to sit near criticality, a fine-tuning — though that is "
"arguably the emergent statement of why gravity is SO weak, and it is a prediction rather than a fudge.")

heading("8.11  From scalar to tensor gravity: two gravities, and deconfinement", 2)
body("Section 8.10 gives a scalar gravity, and the honest next question is whether the medium can reach the tensor "
"theory — general relativity, with light bending and spin-2 waves. The model carries two universal attractions, the "
"amplitude mode of 8.10 (spin-0) and the medium's incompatible-strain / curvature sector (spin-2). The program's "
"working hypothesis at this point was that the two disagree on the light-bending parameter γ — 0 for the scalar, 1 "
"for the spin-2 graviton — so that, a massless mediator's 1/r beating a massive Yukawa, γ would climb from 0 to 1 "
"across the amplitude Compton wavelength 1/m_A: a scale-dependent γ, GR at long range and scalar-contaminated below "
"1/m_A. That hypothesis is tested exhaustively in Sections 8.32–8.49, and the honest outcome, recorded here because "
"the section is otherwise read as a positive step, is that it does not survive. The spin-2 sector as realized also "
"gives γ = 0: a static mass sources the graviton's time component (Newtonian attraction) but not its spatial "
"polarizations, in every channel, smooth and topological, continuum and discrete, so γ stays 0 at every scale "
"(Section 8.37) and the realized theory is Nordström at all ranges, not GR at long range. The climb does not happen, "
"and the scale-dependent γ it predicted is retracted with it. What survives this section, and is load-bearing, is "
"narrower and genuinely a result: the confining spin-2 curvature sector deconfines into a real massless 1/r^2 "
"Newtonian attraction once the induced Einstein term μ > 0 — the graviton exists and attracts, it simply does not "
"bring the Einstein factor of two. This still reinterprets 8.10: tuning the scalar long-range by m_A → 0 forces "
"γ = 1/2, observationally excluded, so the amplitude mode stays gapped and the long-range attraction is the "
"deconfined graviton's — but as a Newtonian, not an Einstein, force.")
body("But the spin-2 sector is confining in the pure medium. Its energy is the elastic cost of incompatible strain — "
"the biharmonic (Kirchhoff-plate) action κ (∇²φ)^2 — whose 3D Green's function grows linearly with "
"separation, giving a constant force between two curvature charges: a string tension, not a Newtonian force. To become "
"gravity it must deconfine. The mechanism is the same Sakharov loop that induced the photon's Maxwell term (8.7): "
"integrating out the gapped matter generates an Einstein-Hilbert term, which at quadratic order adds an ordinary "
"two-derivative stiffness μ (∇φ)^2 to the biharmonic. The full propagator is then 1/(κ q^4 + μ q^2), "
"whose exact three-dimensional Green's function is closed-form:")
add_eq("G(R) = (1 / 4 π μ R)(1 - exp(-R / ℓ)),     ℓ = √(κ / μ)", "8.11a")
body("The lower-derivative induced term dominates the infrared: below the crossover ℓ the force is the confining "
"string tension, and above ℓ it turns over into an exact inverse-square Newtonian tail with G = 1/(4 π μ). "
"Measured on the correct tool for a spherically symmetric source — a one-dimensional radial ODE with no box and no "
"periodic images — the Newton-tail exponent is -2.0000, Newton's constant matches 1/(4 π μ) to five figures, the "
"closed form (8.11a) is reproduced to one part in ten million, and the pure-medium (μ = 0) confinement shows its "
"linear growth directly, G ~ R^+1.000 — the clean form of the string tension that no finite box can display. any "
"positive μ deconfines the graviton: the confining +R becomes a Newtonian -1/r.")
result("Result 8.11 — the graviton deconfines.", "The tensor (curvature) sector is confining in the pure medium — a "
"constant-force string tension, the clean 3D form of the earlier 'curvature charges repel and grow' result. A "
"positive Sakharov-induced Einstein term turns it into a massless Newtonian graviton: the force crosses over at "
"ℓ = √(κ/μ) from the string tension to an exact 1/r^2 with G = 1/(4 π μ). Everything now rests on the sign "
"of the induced μ.")

heading("8.12  The sign of induced gravity, and a dynamical spin-2 graviton", 2)
body("Deconfinement needs μ > 0, and this is the notorious Sakharov sign — the induced Einstein coefficient is "
"ultraviolet-dominated and scheme-sensitive, and free fields do not universally give the healthy sign. The model "
"settles it by a calibration it already owns: its induced photon is healthy (8.7), so the induced Coulomb kinetic "
"term — the charge-density correlator <J0 J0> — is a healthy dielectric, susceptibility χ > 0. Computing the "
"induced Newtonian kinetic term — the energy-density correlator <t00 t00>, since energy is the gravitational charge — "
"from the same gapped-Dirac loop with identical conventions, its momentum-squared coefficient comes out the same "
"(positive) sign as the photon's in every case tested: five masses, three cutoffs, and both energy-density vertex "
"definitions. Induced gravity is therefore as healthy as the electromagnetism the model already runs on: μ > 0.")
body("What kind of graviton is it? On the 2+1D emergent cone the spatial spin-2 graviton is non-dynamical — a massless "
"symmetric tensor has D(D-3)/2 physical polarizations, which is zero in three spacetime dimensions and two in four. "
"So the radiative, light-bending part of gravity can only appear in 3+1D, and there it does: a four-component Dirac "
"loop induces a nonzero transverse-traceless kinetic term (whereas the same construction gives zero in 2+1D), its two "
"polarizations h_plus and h_cross are degenerate to four figures (a single helicity-2 field, not two unrelated "
"modes), and its sign matches the induced transverse photon (a healthy, not ghost, mode). Both sectors of the "
"graviton — the Newtonian h00 and the radiative spin-2 — are thus induced and healthy in the physical dimension.")
body("The Einstein normalization γ = 1 is the transversality (Ward identity) of the induced graviton, and here the "
"model draws a sharp and honest line. On a periodic lattice (a torus, with no boundary and hence no cutoff surface "
"term) the U(1) photon Ward identity closes to machine precision once the diamagnetic seagull is included — the "
"symmetry-preserving regulator works, because U(1) gauge invariance is an exact lattice symmetry. Diffeomorphism "
"invariance is not a lattice symmetry — a lattice keeps only discrete translations — so the graviton Ward identity "
"cannot close exactly: it is inhomogeneous (the induced vacuum stress <T> is nonzero, the same term that reappears as "
"the cosmological constant in 8.13), and no finite-cutoff direct measurement can force γ = 1. That is not a "
"numerical shortcoming but a structural fact: γ = 1 is emergent in the infrared, on the same footing as the "
"model's emergent Lorentz invariance. It is read from Weinberg's theorem — a massless spin-2 coupled to the conserved "
"infrared (Dirac) stress tensor is forced to be Einstein — now with the previously-missing ingredient supplied by "
"measurement: the spin-2 graviton genuinely propagates and is healthy.")
body("Two distinct objects are in play here, and the report keeps them apart. One is whether the spin-2 graviton "
"exists and is healthy — a property of the induced propagator, measured in this section and affirmative. The other "
"is whether a static mass couples to that graviton's spatial polarizations — the coupling that fixes the "
"post-Newtonian ratio γ = Ψ/Φ, and the quantity an experiment actually weighs. Weinberg's theorem joins them only "
"conditionally: if the mass couples to the massless spin-2 at all, that coupling is forced to be Einstein; it says "
"nothing when the coupling is absent. Sections 8.32-8.37 measure precisely this coupling and find it zero in every "
"smooth channel — a scalar energy density cannot source the spin-2 spatial stress (a selection rule, Section 8.32), "
"and every elastic and compressive route sets the spatial metric by local compression rather than by the Newtonian "
"potential (Sections 8.33-8.36). The honest reconciliation is therefore that this section's genuine output is the "
"healthy, dynamical spin-2 graviton, whereas the 'Einstein in the infrared' clause is the conclusion of a conditional "
"whose premise the model's own later measurements contradict. As realized, the model is Nordström — scalar gravity, "
"γ = 0 — with a healthy but matter-decoupled spin-2 mode also present; the Einstein completion would require the mass "
"to nucleate a net topological disclination charge proportional to itself, which §8.49 measures directly on the "
"lattice and finds absent (the mass is pure compression, and compression carries no net charge). γ = 1 is thus "
"closed in every channel the model exposes, continuum and discrete; it would need a gravity mechanism that changed "
"the medium's connectivity rather than its bond lengths — a different construction, not this one.")
result("Result 8.12 — induced gravity is healthy and spin-2; its Einstein infrared limit is argued, not realized (resolved negatively in Sections 8.32-8.37).", "The induced Newtonian "
"coupling has the same sign as the model's working photon: μ > 0, so the graviton of 8.11 deconfines into real "
"attraction. In 3+1D the radiative spin-2 graviton is dynamical, doubly degenerate (helicity 2), and healthy — the "
"section's firm output. On the Einstein normalization the report draws a conditional line: Weinberg forces γ = 1 for "
"a massless spin-2 coupled to the conserved infrared stress tensor, but this is emergent rather than lattice-exact "
"(diffeomorphism invariance, unlike the exactly-closing U(1) Ward identity, is not a lattice symmetry) and, more "
"decisively, conditional on the mass coupling to the graviton's spatial modes at all. The linearised Einstein term is "
"thus reached as an infrared attractor only in that conditional sense, and Section 8.26 measures what the attractor "
"language itself can carry: the far field forgets the ultraviolet coefficient exponentially, and the operators sort "
"as a fixed point requires — higher-derivative structure irrelevant, a graviton mass relevant. Open still, and "
"sharper than before: the fixed point is empirical rather than protected, since no exact lattice symmetry forbids the "
"one relevant deformation; the magnitude of G is cutoff-dependent (the Sakharov feature); and γ = 1 is not merely "
"unmeasured but resolved against — Sections 8.32-8.37 measure its Weinberg premise, that a static mass sources the "
"graviton's spatial stress, to vanish in every smooth channel, leaving the model Nordström (γ = 0) in every realized "
"coupling. Nothing here reaches the nonlinear Einstein equations.")

heading("8.13  The cosmological constant", 2)
body("The same induced vacuum stress that made the graviton Ward identity inhomogeneous is, physically, the "
"cosmological constant — and it is the sharpest quantitative disaster in physics. The medium's zero-point energy is "
"of order the microscopic (node) scale; with the node spacing fixed at the Planck length (8.8's scale-fixing), that "
"is about 10^122 times the observed dark-energy density. Taken at face value it demands a 122-digit fine-tuning. But "
"the model's vacuum is not empty space with fields on top — it is a self-sustained condensate, and that changes what "
"gravitates. Following Volovik's analysis of emergent gravity in quantum liquids, the emergent metric couples to the "
"vacuum stress — the grand-canonical potential density ρ_Λ = ε - μ n = -P — not to the bare energy density "
"ε. A self-sustained vacuum, one that can exist with nothing outside pushing on it (which is what the vacuum of "
"empty space is), has zero pressure:")
add_eq("ρ_Λ = ε - μ n = -P;      self-sustained vacuum: P = 0 ⇒ ρ_Λ = 0", "8.13a")
body("The huge zero-point energy is absorbed into the equilibrium condensate density, not into curvature. This is not "
"a tuning: the density self-adjusts so that P = 0, for any bare ε. Sweeping the bare vacuum energy across all 122 "
"orders of magnitude, the gravitating ρ_Λ stays zero to machine precision while a rigid (non-adjusting) vacuum "
"would gravitate the full ε — the standard disaster. The observed small but nonzero Λ is then the residual of "
"a slight departure from equilibrium (ρ_Λ scales with that departure), not a cancellation of 122 digits.")
result("Result 8.13 — the cosmological-constant fine-tuning, dissolved (not the value derived).", "Because the vacuum "
"is a self-sustained condensate, the quantity that gravitates is its grand potential -P, which vanishes at "
"equilibrium — automatically, for any bare zero-point energy, with the density self-adjusting. The 10^122 fine-tuning "
"is thus dissolved: the equilibrium vacuum gravitates nothing by thermodynamics. Honest ceiling: this does not "
"predict the observed nonzero Λ, which is relocated to a cosmological question — why the vacuum sits slightly "
"off equilibrium (expansion, matter, relaxation) — and remains open.")

heading("8.14  The Standard-Model gauge group: emergent Yang-Mills", 2)
body("The emergent photon of Section 8.7 is Abelian. The Standard Model needs the non-Abelian groups SU(2) and SU(3), "
"and the question is whether the same fermion loop gives genuine Yang-Mills or merely several decoupled copies of the "
"photon. The decisive difference is the gauge-boson self-interaction, and it has a clean signature. A non-Abelian "
"field strength carries a commutator, F = dA + i[A, A], so a spatially uniform non-Abelian field has nonzero field "
"strength from that commutator alone, while a uniform Abelian field is always pure gauge. Placing a fermion in the "
"fundamental of SU(N) with uniform background links and measuring the induced action (the filled-sea energy), a "
"commuting (Cartan) configuration costs nothing — pure gauge, to machine precision at every amplitude — while a "
"non-commuting configuration costs an induced action that grows as A^4, exactly the Yang-Mills Tr[A_x, A_y]^2. N^2-1 "
"decoupled photons would give zero for both; the A^4 cost is the fingerprint of a genuine, self-interacting "
"non-Abelian field. It appears for SU(2) (three gauge bosons) and SU(3) (eight gluons), with a single universal "
"coupling guaranteed by the exact non-Abelian lattice gauge invariance of the Wilson links — the same exact-symmetry "
"footing that made the U(1) Ward identity close exactly in 8.12.")
result("Result 8.14 — emergent non-Abelian gauge fields.", "The Sakharov mechanism that induced the photon induces "
"genuine Yang-Mills: a uniform non-commuting gauge field costs an induced action ~ A^4 = Tr[A,A]^2 (the "
"self-interaction), while a commuting one is exactly pure gauge — shown for SU(2) and SU(3), with a universal coupling "
"from exact lattice gauge invariance. Emergent gauge theory thus scales from U(1) to SU(N). Honest ceiling: this is "
"the mechanism, not the Standard Model — it does not derive the specific group SU(3)xSU(2)xU(1), the chiral coupling, "
"the fermion representations or hypercharges, anomaly cancellation, or the Higgs. The group is an input; its "
"Yang-Mills dynamics, and the fermion light cone they inherit, are induced.")

heading("8.15  Chirality without inconsistency: the anomaly and its inflow", 2)
body("Section 8.14 supplies non-Abelian gauge fields and Section 8.2 a single chiral fermion on a domain wall. The "
"Standard Model needs both at once, and that is where the real obstruction sits. The SM's SU(2) couples only to "
"left-handed fermions, and a chiral gauge theory is not merely incomplete but inconsistent unless its anomalies "
"cancel: a chiral fermion's gauge current is not conserved in a background field, so a lone chiral fermion coupled to "
"a gauge field is not a theory at all. The domain-wall construction survives this by a specific mechanism -- "
"Callan-Harvey anomaly inflow. Each wall is individually anomalous, and the charge it appears to lose is supplied by "
"the bulk, which pumps it to the other wall. The content is quantitative and quantized: the bulk Chern number, the "
"number of chiral modes per wall, and the charge pumped per flux quantum are one integer.")
body("Measured on the same Wilson-Dirac strip that carried the chiral fermion: in the topological phase the bulk Chern "
"number is -1.000, the edge spectral flow is +1 on one wall and -1 on the other, and the two Σ to zero; in the "
"trivial phase all three vanish. So each wall separately violates charge conservation by exactly one unit per flux "
"quantum -- a genuine anomaly -- while the lattice as a whole is vector-like and anomaly-free, exactly as "
"Nielsen-Ninomiya demands. Neither wall is a consistent theory alone; the wall pair together with the bulk is.")
result("Result 8.15 — chirality is consistently realizable, by inflow.", "The anomaly here is not a pathology but a "
"quantized bookkeeping identity, and the three numbers that must agree do agree. This is the mechanism "
"Standard-Model chirality requires, and it operates in the model's own domain-wall construction alongside the induced "
"non-Abelian gauge fields of 8.14. honest ceiling: this is not the Standard Model's anomaly cancellation. The SM is a "
"standalone four-dimensional chiral gauge theory whose anomalies cancel among its own fermion content -- the "
"quark/lepton hypercharge conspiracy -- with no bulk to lean on. Here the bulk does the cancelling, so the wall "
"theory is anomaly-free only together with it. Producing a standalone anomaly-free chiral spectrum, i.e. the actual "
"SM fermion content and hypercharges, is not attempted and is not fixed by the medium.")

heading("8.16  The model against the data, and a retraction", 2)
body("A model earns its keep by being falsifiable, and the two predictions this project carries had never been put "
"against real numbers. Doing so sharpens one, deflates the other, and forces a retraction.")
body("The Lorentz-violation signature of Section 8.8 is quadratic and Planck-suppressed. Compared with photon "
"time-of-flight limits on quadratic dispersion, which the literature places at an effective scale of order 10^10 to "
"10^12 GeV, the model's effect is about 10^-16 of current sensitivity. It is therefore safe against every existing "
"bound by roughly sixteen orders of magnitude -- but that cuts both ways, and the honest conclusion is that this "
"prediction is not currently falsifiable: reaching it would need some eight orders of improvement in quadratic-"
"dispersion sensitivity. Present experiments are sensitive to a linear, n = 1 signal, which this model specifically "
"does not predict. Section 8.8's framing of this as the project's testable prediction was too generous.")
body("The gravitational prediction is the real one. With a massless graviton and a gapped amplitude mode (8.11), the "
"amplitude mode survives as a Yukawa addition to gravity of range 1/m_A, so γ falls below 1 inside that range. "
"Cassini's solar-system bound turns out to be nearly useless here, permitting λ up to about 10^10 m; the binding "
"constraint is short-range gravity, where torsion-balance tests of the inverse-square law beat Cassini by fourteen "
"orders of magnitude. Taking the Yukawa strength of order unity, as is natural when both mediators couple to energy, "
"the amplitude gap must satisfy m_A of order 4 meV or more -- which sits within a factor of about 1.6 of the "
"dark-energy scale, 2.4 meV, the well-known coincidence that makes the sub-millimetre range the frontier of these "
"experiments. This is a genuine, reachable test: a gravitational-strength Yukawa just below current reach would show "
"up as γ < 1 at short distance.")
body("The same confrontation retracts the 10^122 criticality tuning reported earlier in this program. That number "
"followed from reading gravity's range as 1/m_A -- the scalar picture -- so that any bound on gravity remaining "
"inverse-square out to large distances forced the amplitude gap to be absurdly small. The tensor arc of Sections "
"8.11 and 8.12 replaced that premise: the long-range force is carried by the massless deconfined graviton, and the "
"amplitude mode is only a short-range correction. The surviving experimental constraint is therefore a lower bound "
"on m_A, with no upper bound at all -- a larger gap is only safer -- and an untuned medium clears it by some thirty "
"orders of magnitude. The fine-tuning is dissolved, not reduced. The scale-fixing result that a0 = l_Planck is "
"unaffected and stands.")
result("Result 8.16 — one prediction sharpened, one deflated, one tuning retracted (the first two later inverted; see the correction that follows).", "Confronted with data: the "
"quadratic Lorentz violation is safe by ~16 orders and, honestly, out of experimental reach -- not the testable "
"prediction it was billed as. The gravitational prediction is testable: short-range tests already require the "
"amplitude gap to exceed about 4 meV, coincidentally within a factor 1.6 of the dark-energy scale and squarely in "
"the sub-millimetre window now being probed. And the 10^122 criticality tuning is retracted, because it rested on the "
"superseded scalar reading of gravity. Together with Section 8.13, which dissolved the cosmological-constant tuning "
"by the condensate's equilibrium thermodynamics, the model has now shed both of its 10^122 fine-tunings -- each "
"removed by a structural result rather than a fitted parameter. (Scope: literature order-of-magnitude bounds, "
"assuming comparable scalar and tensor couplings; loosening that weakens the gravitational bound proportionally.)")
body("Later correction, kept explicit because both of this section's data verdicts were subsequently overturned by "
"the program's own measurements. The Lorentz-violation signature deflated here is in fact the model's one surviving "
"falsifiable prediction: Section 8.39 re-runs the confrontation against ultra-high-energy astrophysics, where a "
"quadratic violation is bounded near the Planck scale rather than at the 10^10–10^12 GeV of photon time-of-flight, "
"and finds the model only about 1.4 orders below the frontier — reachable by next-generation observatories, not out "
"of reach. The gravitational scale-dependent γ promoted here to 'the real one' is, conversely, retracted: Sections "
"8.32–8.49 measure γ = 0 at every scale and in every channel, so light-bending never climbs toward the Einstein "
"value and there is no γ < 1 short-range deviation to detect. What stands unchanged is the third result — the "
"retraction of the 10^122 criticality tuning — because it rests only on the long-range force being carried by the "
"massless deconfined graviton (a Newtonian attraction whose range is 1/r regardless of γ), not on any light-bending "
"claim. So the net of §8.16 inverts: the prediction it deflated is the live one, and the prediction it promoted is dead.")

heading("8.17  The first dynamical integration: chiral matter and a gauge field, in real time", 2)
body("Every result to this point -- emergent Lorentz invariance, chiral fermions, the induced photon and graviton, the "
"anomaly -- is established at the level of a dispersion relation, a band structure, or a defect algebra. None is a "
"running simulation in which two of these emergent sectors coexist and interact in time. That gap is the sharpest "
"honest criticism of the whole program, and this section takes the first step across it, using the anomaly of "
"Section 8.15 as the target. That anomaly was established statically, by counting; the question here is whether the "
"charge it bookkeeps actually moves.")
body("The test is a Laughlin flux threading. On the same QWZ strip -- two chiral walls, x periodic -- the "
"negative-energy sea is filled at t = 0, and one full flux quantum is threaded adiabatically by ramping a uniform "
"vector potential, A: 0 → 2 π / L_x, which enters as k_x → k_x + A(t). Every occupied orbital is evolved under the "
"time-dependent Schrödinger equation by exact exponentiation of the instantaneous Hamiltonian at each step -- no "
"adiabatic-following shortcut, which would presuppose the result. If the anomaly is genuine dynamics rather than "
"bookkeeping, exactly one unit of charge must cross from one wall to the other per flux quantum, pumped through the "
"bulk since the walls are spatially separated and nothing local joins them.")
body("It does. The charge in the bottom half of the strip changes by 0.9991 of one unit in the topological phase and "
"by exactly zero in the trivial control, so the transfer is the anomaly and not the ramp. Slowing the ramp drives the "
"residual to zero as roughly one over the number of steps -- the quantization is physics, not an artifact of the "
"discretisation. The charge has crossed the bulk in real time: Callan-Harvey inflow, observed as dynamics.")
result("Result 8.17 — the anomaly happens, and two sectors run together.", "Threading one flux quantum through the "
"chiral strip and evolving the filled sea under the actual time-dependent Schrödinger equation pumps exactly one unit "
"of charge between the walls (0.9991, converging to one as the ramp slows), and exactly zero in the trivial control. "
"The topological accounting of Section 8.15 therefore describes real time evolution -- the bulk really does supply "
"what each wall loses. This is the program's first running simulation in which two emergent sectors -- chiral matter "
"and a gauge field -- coexist and interact in time, and the consistency does not break down. honest scope: the field "
"threaded is the U(1) gauge field, not gravity. The full integration this program's limitations call for -- emergent "
"Lorentz-invariant chiral quantum matter interacting through an emergent spin-2 gravity, with back-reaction -- is not "
"done; gravitational back-reaction in a running simulation remains the open integration problem. What is shown is that "
"the first pair of sectors can be run together at all.")

heading("8.18  Gravitational back-reaction, run as a conserving simulation", 2)
body("Section 8.17 ran chiral matter together with a gauge field. gravity is the harder and more important case, "
"because back-reaction is gravity's defining feature: matter tells geometry how to curve, geometry tells matter how to "
"move, and the two must be solved together and self-consistently. Until that is done in time, a gravity result is a "
"dispersion relation rather than a force. This section does it, and it is the step at which the gravity sector stops "
"being a toy.")
body("The system evolved is the Schrödinger-Newton pair -- the non-relativistic limit of a massive matter field "
"minimally coupled to its own gravity, and the standard model of self-gravitating quantum matter:")
add_eq("i ∂_{t} ψ = -(1/2) ∇²ψ + Φ ψ,      ∇²Φ = 4 π G |ψ|^{2}", "8.18a")
body("The gravity in (8.18a) is not inserted by hand: it is the infrared-effective form of the gravity this program "
"derived -- the deconfined graviton mediates an exact Newtonian potential with G = 1/(4 π μ) (Section 8.11), and the "
"sign μ > 0 was measured against the model's own healthy photon (Section 8.12). Co-evolving matter with that "
"potential uses the derived gravity rather than inventing one.")
body("Four properties separate a scientific simulation from a toy, and all four hold. conservation: evolved by a "
"split-step scheme that is symplectic in the matter sector with the potential resolved self-consistently each step, "
"the total energy is conserved to three parts in a billion and the norm to four parts in a hundred trillion, through "
"the full nonlinear evolution -- against which the earlier gravity-by-density attempt, which leaked some eighty per "
"cent of its energy through an inconsistent cutoff, is the instructive failure. self-binding: with gravity on, a "
"packet forms a self-gravitating bound state of negative total energy whose width settles about a finite soliton "
"scale, while the identical packet with gravity off disperses without bound -- the single self-bound lump that the "
"nearly incompressible medium of Phase 3 could never produce. equilibrium: imaginary-time relaxation finds the soliton "
"ground state, and it satisfies the scale-virial identity 2T + W = 0 to within three hundredths, the signature of a "
"genuine gravitational equilibrium rather than a long-lived transient. This last check requires isolated, free-space "
"boundary conditions; a periodic box distorts the long-range potential and leaves a spurious virial residual near one, "
"which is reported alongside as a methodological control. convergence: the relaxed soliton's energy and virial settle "
"monotonically as the mesh is refined, so the bound state is a property of the continuum system and not of the grid.")
result("Result 8.18 — gravity as a force that can be run.", "The coupled matter-plus-gravity system evolves "
"self-consistently with energy conserved to ~10^-9 and norm to ~10^-14; it binds matter into a self-gravitating soliton "
"that a gravity-off control does not form; the relaxed soliton satisfies the virial identity 2T + W = 0; and both "
"converge under mesh refinement. This is the gravity the program derived, run as a force in time rather than read off "
"a propagator -- the first back-reacting gravitational simulation here, and the point at which the gravity sector "
"meets the standards of a scientific simulation. honest scope: non-relativistic (Schrödinger, not Dirac matter), "
"scalar/Newtonian (the h00 sector; the radiative spin-2 graviton of Section 8.12 is not evolved), and the matter is a "
"classical field. Chiral quantum matter interacting through the emergent spin-2 gravity with radiative back-reaction "
"is still the open integration problem.")

heading("8.19  The magnitude of Newton's constant", 2)
body("One caveat has trailed every gravity result: the strength of gravity was called cutoff-dependent. That is the "
"standard Sakharov ambiguity -- the induced Einstein-Hilbert coefficient is ultraviolet-dominated, so in a continuum "
"theory with an arbitrary cutoff its magnitude is arbitrary. The caveat does not apply to this model, because the "
"medium has a physical ultraviolet cutoff: the node spacing a0, fixed to the Planck length in Section 8.8. Newton's "
"constant is therefore not a free parameter but a definite number, computable from the lattice.")
body("Gravity's mediator takes its kinetic term from the fermion loop, and G is set by the induced Newtonian stiffness "
"μ -- the coefficient of q^2 in the energy-density correlator, the h00 sector's induced 1/(4 π G). Evaluated over "
"the whole Brillouin zone, so that the cutoff is the physical lattice scale rather than an arbitrary disc, the loop "
"returns a number of order unity in lattice units. Hence G = 1/(4 π μ) is of order a0^2, that is of order the Planck "
"area, up to an order-unity factor -- computed rather than fitted. Including more light fermion species stiffens the "
"geometry in the standard way, G ~ a0^2 / N_f, so even a Standard-Model-like species count leaves G at a few per cent "
"of the Planck area: still Planckian. Gravity is weak for exactly one reason, that a0 is Planckian, and the induced "
"coefficient supplies no hierarchy and requires no tuning.")
result("Result 8.19 — G is the Planck area, with no tuning.", "With a physical cutoff the induced stiffness is an "
"order-unity number and G = O(1) × a0^2 = O(1) × l_Planck^2. The weakness of gravity is entirely the smallness of the "
"Planck length; the species count only sharpens it. Taken with Section 8.8, which fixed a0 = l_Planck by matching the "
"measured G, the loop closes: the node spacing is the Planck length, and the gravity the medium induces has Planck "
"strength. honest scope: the scale (order a0^2) and the sign (positive, hence attractive) are robust, but the precise "
"order-unity coefficient remains scheme-sensitive -- the residual of the Sakharov ambiguity, since the lattice is one "
"regulator among several. What is settled is the qualitative point that had been left open: the magnitude of G is not "
"free here.")

heading("8.20  Gravity that radiates, and the monopole that cannot", 2)
body("Section 8.18 ran gravity as a force, but in the Newtonian limit: the potential was solved from an instantaneous "
"constraint, which is exact only for slow sources. A real gravitational field is retarded -- it propagates at finite "
"speed and carries energy away as radiation -- and that radiative sector is the last structural piece of the gravity "
"programme. It also carries a sharp signature which separates the tensor gravity this programme arrived at from the "
"scalar gravity it discarded along the way.")
body("The signature is the monopole. A scalar (Nordström) gravity radiates from a spherically pulsating mass: a "
"breathing star emits scalar gravitational waves. A spin-2 field cannot. The radiative degrees of freedom of "
"linearised gravity are the transverse-traceless part of h_ij, and the TT projection annihilates a spherically "
"symmetric source identically, so monopole radiation is forbidden and the leading channel is the quadrupole. That is "
"why a pulsating star does not gravitationally radiate, and it is a structural test of whether the model's gravity is "
"really spin-2.")
body("Evolving the linearised TT wave equation in momentum space -- each mode a driven oscillator, the source switched "
"off after a few cycles, and the field energy remaining once the near field has dispersed counted as the radiated "
"energy -- gives three results. The disturbance propagates: after the source stops, the outgoing shell moves at speed "
"0.96 in units where c = 1, the few-per-cent deficit being the finite width of the shell rather than a slow wave. The "
"radiation is carried by exactly two polarizations, the transverse-traceless projector having rank two, which are the "
"helicity-2 states of Section 8.12 and not a scalar breathing mode. And at identical amplitude, width and frequency, "
"the monopole source radiates about 10^-14 of what the quadrupole source radiates -- a ratio of 3×10^-13, which is machine "
"zero against a finite signal.")
result("Result 8.20 — the radiation is spin-2, and the monopole channel is closed.", "The model's gravitational field "
"propagates at c and carries exactly two polarizations, and a spherically pulsating mass radiates nothing through it, "
"while an equal quadrupole radiates a finite amount. Scalar gravity would have opened the monopole channel; spin-2 "
"gravity forbids it, and the measurement finds it shut to machine precision. This is a structural confirmation, "
"independent of the earlier propagator arguments, that the gravity this programme ended up with is genuinely tensor. "
"honest scope: the radiation is linearised and the source is prescribed, so the back-reaction of the radiated energy "
"on the source -- the inspiral of a binary -- is not computed. Neither is the quadrupole luminosity formula tested: "
"that law assumes a source small compared with the wavelength, the source used here is not, and the frequency "
"dependence in this setup therefore reflects the source's own spatial spectrum rather than the multipole expansion. No "
"such claim is made. That gap is closed in Section 8.23, which also quantifies why it could not have been closed "
"here: at the compactness of this source the radiated power is suppressed roughly thirteenfold.")

heading("8.21  The integration closed: a source that radiates, and thereby decays", 2)
body("Three results had been arrived at separately, and each named the same gap. Section 8.18 ran matter and gravity "
"together self-consistently, but in the Newtonian limit -- an instantaneous constraint, which cannot radiate. Section "
"8.20 evolved the dynamical, retarded spin-2 field and found it radiates like a tensor, but from a prescribed source, "
"so the energy carried away was never taken back out of the matter that emitted it. Section 8.17 showed only that two "
"emergent sectors can be run together at all. What none of them did was couple all three at once: matter that sources "
"a radiative field, feels it back, and loses the energy the field removes. That is what gravity actually does, and it "
"is the last thing separating a simulated force from a real one.")
body("The closure is obtained by refusing to insert it. Matter and the radiative transverse-traceless field are "
"evolved from a single Hamiltonian,")
add_eq("H = ∫[ (1/2)|∇ψ|^{2} + (g/2) h_{ij} Re(∂_{i} ψ* ∂_{j} ψ) + (1/2) Φ |ψ|^{2} ] "
       "+ ∫(1/2)( π^{2} + |∇h|^{2} )", "8.21a")
body("whose variation gives both equations of motion at once -- the matter equation with its coupling to h, and the "
"wave equation for h driven by the TT part of the matter stress. No radiation-reaction force is added by hand. If the "
"matter loses energy, it is because the field took it, and the arithmetic has nowhere to hide.")
body("It balances. For a quadrupolar (squeezed) matter configuration the matter energy falls by 2.0451×10^-2 while the "
"field energy rises by 2.0452×10^-2 -- agreement to four significant figures -- with the total conserved to two parts in "
"a million. The matter energy decreases monotonically as radiation streams outward: an inspiral in field-theoretic "
"form, driven by the source's own emission. Two controls confirm the transfer is physical rather than numerical. A "
"spherically symmetric source radiates 5×10^-8, smaller by a factor of four hundred thousand, so the monopole "
"prohibition of Section 8.20 survives being made self-consistent -- it was not an artifact of having prescribed the "
"source. And the radiated energy divided by the square of the gravitational coupling is flat to four digits across a "
"range of couplings, which numerical leakage would not respect.")
result("Result 8.21 — gravity, simulated as a force that costs its source.", "Matter dynamics, a dynamical radiative "
"field, and self-consistent energy exchange now run in one evolution from one Hamiltonian. The matter's energy falls "
"by exactly what the field's rises, to four significant figures, with the total conserved to ~10^-6, and nothing about "
"radiation reaction was put in by hand: the source decays because the field it generates carries energy away. The "
"monopole prohibition holds self-consistently, and the transfer scales as the square of the coupling. This closes the "
"integration gap that Sections 8.17, 8.18 and 8.20 -- and the limitations section -- all named. honest scope: the "
"matter is a classical, non-relativistic field, not the relativistic quantum chiral matter of the fermion sector, and "
"the gravity is linearised. The coupling is also dialled far above its physical value so that the decay is visible in "
"a short run; real radiation reaction is minuscule, and what is demonstrated is the mechanism and the balance of the "
"budget, not the magnitude. Relativistic quantum matter coupled to nonlinear gravity is taken up in Section 8.22.")

heading("8.22  Past the toy: relativistic quantum matter, nonlinear gravity, and the coupling taken to zero", 2)
body("Section 8.21 closed the integration but admitted three caveats, and they were not small ones: the matter was a "
"classical, non-relativistic field; the gravity was linearised; and the coupling was dialled far above its physical "
"value. A single 'improved' run that changed all three at once would be unfalsifiable -- any success could be "
"attributed to any of the changes, and any failure to any other. The three axes are therefore attacked separately, so "
"that each upgrade is independently measured.")
body("First, the matter is made relativistic and quantum. The Schrödinger field is replaced by a genuine Dirac field: "
"a four-component spinor obeying a first-order equation with α and β matrices, carrying spin and antiparticle "
"components. And it is a many-fermion quantum state rather than a classical field -- a set of mutually orthonormal "
"occupied modes evolved as a Slater determinant. The matter Hamiltonian is the Dirac operator in a perturbed spatial "
"frame,")
add_eq("H_{m} = ∫ψ^{†} [ -i α_{i} ( δ_{ij} - (g/2) h_{ij} ) ∂_{j} + m β ] ψ   (hermitised)",
       "8.22a")
body("whose conjugate source for h is the relativistic momentum flux, arising from the same single Hamiltonian, so "
"the energy exchange stays derived rather than inserted. Because the gravitational coupling is a one-body operator, "
"determinant evolution is exact for the matter: there is no mean-field error in the matter sector at all, and Pauli "
"antisymmetry is preserved identically. It balances as before -- the matter energy falls by what the field energy "
"rises, to agreement of 8.6×10^-9, with the total conserved to 2×10^-10 -- and the occupied modes stay orthonormal to "
"5×10^-12, so what is being evolved remains a legitimate antisymmetrised fermionic state and not a drifting classical "
"field. A spherically symmetric control radiates a factor of 2×10^8 less, sitting at the integrator's own noise floor: "
"the monopole prohibition holds for relativistic quantum matter too.")
body("Second, the gravity is made nonlinear. The field acquires the derivative self-coupling that is the structural "
"signature of general relativity -- the statement that gravity gravitates:")
add_eq("H_{f} = ∫(1/2)( π^{2} + |∇h|^{2} ) + (λ/2) h_{kl} ∂_{k} h_{ij} ∂_{l} h_{ij}", "8.22b")
body("Two things are then measured, and the second is decisive. Energy remains conserved to 1.4×10^-11, which "
"establishes that the self-interaction is a genuine Hamiltonian term and not a force bolted onto the equations of "
"motion after the fact. And superposition fails: evolving two wave packets together no longer equals evolving each "
"alone and adding. In the linearised theory that residual is machine zero, measured at 5.7×10^-16; with the vertex "
"switched on it rises to 2.1×10^-3, tracking the product of the coupling and the amplitude exactly as a cubic term must. "
"A field whose waves scatter off one another is not a linear field.")
body("Third, the coupling. Physical gravitational coupling cannot be simulated directly and no honest report should "
"suggest otherwise: the ratio of radiated to rest energy for anything resolvable is around 10^-40, some thirty orders "
"below double precision. What can be established is the property that makes extrapolation legitimate rather than "
"hopeful -- that the transfer is exactly second order in the coupling. The radiated energy divided by the square of "
"the coupling is flat to six significant figures across four decades. Scaling to the physical value therefore changes "
"the magnitude and nothing else. The floor is reported rather than hidden: the closure test remains verifiable only "
"while the transfer exceeds the integrator's own drift, here down to a coupling of about 10^-3, and below that the "
"budget is extrapolated by the measured square law rather than checked. With all three upgrades switched on together, "
"the self-interaction carries 9.5 percent of the field energy and shifts the radiated energy by 2.0 percent -- it is "
"active rather than negligible -- and the budget still closes to five significant figures.")
result("Result 8.22 — the three caveats, addressed one at a time.", "The matter is now relativistic and quantum: a "
"Dirac spinor field carried as a many-fermion Slater determinant, whose evolution is exact for the matter because the "
"gravitational coupling is one-body, with Pauli antisymmetry surviving to ~10^-12 and the budget closing as it did for "
"the classical field. The gravity is nonlinear: it carries the derivative self-coupling of general relativity, "
"conserves energy (so it is Hamiltonian, not patched), and superposition fails -- the sharpest available proof that "
"the field is no longer the linear one. The coupling dependence is exactly g^2 to six significant figures over four "
"decades, so extrapolation to physical strength is arithmetic; it is still an extrapolation and is labelled one. "
"still open, and weakened by none of the above: the geometry is classical. This is semiclassical gravity -- quantum "
"matter, classical field -- so the measurement problem is untouched, and the cubic vertex is the structural "
"nonlinearity rather than the resummed Einstein-Hilbert series. There is no black hole here and none is claimed.")

heading("8.23  Einstein's quadrupole luminosity formula: the gravity arc's first hard number", 2)
body("Every gravitational result to this point is structural. A sign (the induced stiffness is positive), a rank (the "
"projector admits two polarizations), a scaling (the transfer goes as the square of the coupling), a prohibition (the "
"monopole channel is shut), a balance (the budget closes). Those are strong results, and several of them could have "
"come out the other way. But none of them can be contradicted by a known closed-form answer. The quadrupole "
"luminosity formula can. It is a number, it has been measured on real binary pulsars, and a theory either reproduces "
"it or fails.")
body("Section 8.20 could not test it, and said so: the law assumes a source small compared with the wavelength, that "
"source was not, and the claim was dropped rather than misreported. What follows returns to it with the compactness "
"controlled. The prediction is parameter-free -- derived from this report's own field normalisation, not fitted. "
"From the radiative Hamiltonian used throughout Sections 8.20 to 8.22 the retarded solution and the energy flux give")
add_eq("L = g^{2} / (160 π) × Q⃛_{ij} Q⃛_{ij},    "
       "Q_{ij} = ∫ρ ( x_{i} x_{j} - (1/3) r^{2} δ_{ij} ) d^{3}x", "8.23a")
body("and the two analytic steps behind that coefficient are audited before anything is simulated. Substituting the "
"circular-binary value of the quadrupole contraction into the general-relativistic form reproduces the textbook "
"binary luminosity exactly, and the transverse-traceless angular average is confirmed numerically to two parts in a "
"hundred thousand. So equation 8.23a is general relativity's law transcribed into this model's normalisation, with no "
"free constant left in it.")
body("The identity the whole formula rests on is checked first: that the integral of the stress over the source equals "
"half the second time derivative of the mass quadrupole. It holds to five parts in a hundred million for a genuinely "
"rotating extended body -- provided the body carries its full stress, ram pressure together with the centripetal "
"binding stress that actually holds a spinning object together. Omit that binding term and the identity fails "
"outright, which is the classic way this calculation goes wrong. It is also the reason the mass quadrupole, rather "
"than the stress, is what sets the radiation.")
body("Then the measurement. The field is evolved on the grid, driven by a compact rotating quadrupole whose third "
"derivative is known in closed form, and the radiated power is read off as the secular slope of the field energy. "
"This works cleanly for a structural reason: a rigidly rotating quadrupole has a near-zone field that simply rotates, "
"so its energy is constant, and the field energy is exactly a constant plus the radiated power times time. Nothing is "
"fitted but that line, and the linearity is 0.99998. The measured power divided by the predicted power comes out "
"0.9924, 0.9983 and 0.9996 in three independent configurations, stable to a part in a thousand across a fourfold "
"range of timestep. The two scaling laws follow: the power goes as the sixth power of the orbital frequency, fitted "
"exponent 6.007 -- precisely the law Section 8.20 had to drop -- and as the square of the quadrupole amplitude, "
"fitted exponent 2.000.")
body("Finally, the earlier failure is quantified rather than excused. For a Gaussian source the exact leading-multipole "
"power carries a computable form factor, and sweeping the source size tracks that form factor to four decimal places "
"across a twelvefold suppression. At the compactness of the Section 8.20 source the power is down by about a factor "
"of thirteen. The earlier result was not a model failure and dropping the claim there was correct.")
result("Result 8.23 — general relativity's radiation law, reproduced to sub-percent, with nothing fitted.", "The "
"measured radiated power agrees with the quadrupole luminosity formula to 0.9924, 0.9983 and 0.9996 of prediction, "
"with the coefficient derived from this model's own field normalisation and cross-checked against the "
"general-relativistic binary law before simulating. The sixth-power frequency law is recovered (exponent 6.007) and "
"the amplitude-squared law is exact (2.000). The supporting identity relating the stress integral to the mass "
"quadrupole holds to 5×10^-8, and the compactness dependence follows the exact form factor, quantifying why Section 8.20 "
"could not perform this test. This is the first result in the gravity programme that a known closed-form answer could "
"have contradicted -- everything preceding it was structural -- and it is general relativity's number. honest Scope: "
"the source is prescribed rather than self-gravitating. That is deliberate and is not the earlier limitation "
"returning: prescription was never the obstacle, compactness was, and back-reaction is a higher-order effect that "
"would contaminate a leading-order measurement. What is established is the radiation law.")

heading("8.24  Orbital decay against Peters-Mathews: a binary that spirals in", 2)
body("Section 8.23 reproduced the quadrupole luminosity formula -- a rate, the instantaneous power of a "
"prescribed source. The sharper question is what that rate does to a bound system, and there the "
"comparison is with a second and independent closed form: the Peters-Mathews orbital-decay law of "
"1964, which is the result that was actually confirmed on the Hulse-Taylor binary pulsar and earned "
"the 1993 Nobel Prize. A binary that radiates must shrink, and it must shrink at a specific rate.")
body("What gives this test teeth is a constraint that removes the last place a fudge could hide: "
"general relativity has only one Newton constant. The force that binds the orbit and the coupling "
"that sets the radiated power are not independent knobs. In this model the binding sector obeys "
"∇²Φ = g_N ρ, giving G_N = g_N/4π, while the radiative sector was fixed in Section 8.23 at "
"L = g^2/(160 π) Q⃛.Q⃛, giving G_rad = g^2/32π. Requiring that these be the same constant "
"locks")
add_eq("g_N = g^{2} / 8", "8.24a")
body("with nothing left to adjust. The orbital frequency, the binding energy and the radiated power "
"then all descend from one G, exactly as in general relativity. For an equal-mass circular binary the "
"consequences are closed-form,")
add_eq("L = (64/5) G^{4} m^{5} / a^{5},    da/dt = -(128/5) G^{3} m^{3} / a^{3},    "
       "a(t)^{4} = a_{0}^{4} - (512/5) G^{3} m^{3} t", "8.24b")
body("and the chain from the model's own normalisation to these expressions is audited before any field "
"is evolved: the luminosity computed as (G/5) times the quadrupole contraction agrees with the Peters "
"form exactly, and Newtonian energy balance reproduces the Peters decay rate to sixteen digits. The "
"values below are therefore general relativity's, transcribed, with no free parameter.")
body("The radiative field is then evolved on the grid, driven by a genuine Keplerian binary, and the "
"radiated power is read off as the secular slope of the field energy at four separations. The ratio "
"of measured to Peters luminosity is 0.9903, 0.9924, 0.9939 and 0.9973, and the fitted exponent of L "
"against separation is -4.983 where Peters requires exactly -5. That steep dependence is what makes "
"an inspiral run away: as the orbit shrinks the emission rises sharply, and the binary merges in "
"finite time. Converting the grid luminosity through energy balance gives the orbital decay rate, "
"which tracks Peters at the same fraction of a percent, and integrating it produces the "
"(t_c - t)^{1/4} chirp to coalescence at the Peters time.")
body("The residual is systematic rather than random: the ratio "
"moves toward unity as the orbital velocity falls, from 0.9903 at v/c = 0.160 to 0.9973 at v/c = "
"0.131. Peters is a leading-order result in the slow-motion expansion, so a small deficit that shrinks "
"with velocity is exactly the signature of the neglected higher post-Newtonian terms, not of numerical "
"error.")
result("Result 8.24 — the binary decays at the Peters-Mathews rate.", "With the binding and radiative "
"couplings locked to a single Newton constant by g_N = g^2/8, so that nothing can be tuned per "
"quantity, the grid-measured luminosity of a Keplerian binary matches the Peters value to 0.9903, "
"0.9924, 0.9939 and 0.9973 across separation, the exponent of L against separation comes out -4.983 "
"against the required -5, and the orbital decay rate da/dt tracks Peters at the same accuracy. The "
"residual shrinks as v/c falls, which is the signature of neglected higher post-Newtonian terms. This "
"is the model's second hard number and the observable that first proved gravitational radiation is "
"real. honest Scope: the radiation reaction is adiabatic -- the luminosity is measured on the grid and "
"fed back through energy balance, which is how binary inspirals are modelled in gravitational-wave "
"astronomy -- and is not a first-principles self-force. The rate is measured at fixed separation "
"deliberately: a field-energy budget over a moving orbit does not isolate radiated power, because the "
"near-zone standing energy itself changes as the orbit shrinks and contaminates the slope at the ten "
"to fifteen percent level. That confound is reported rather than absorbed. The orbit is Newtonian "
"point-mass with leading-quadrupole radiation at v/c of about 0.15, so percent-level agreement is the "
"correct expectation, not machine precision.")

heading("8.25  What the classical geometry costs: semiclassical gravity is inconsistent", 2)
body("Section 8.22 named the assumption that survives every other repair in this report -- the geometry is classical -- "
"and left it as an acknowledged limitation. That understates it. Every gravitational result from Section "
"8.18 onward sources a classical field on the expectation value of quantum matter, and that prescription is not "
"merely incomplete: it is inconsistent. This section measures the cost, which makes it an honest-negative result "
"about the framework the report has been using.")
body("Three failures, in increasing order of severity. First, sourcing gravity on the matter density makes the "
"evolution nonlinear in the wave function, so the superposition principle fails: the residual is machine zero with "
"gravity switched off and of order one-half with it on. The contrast with Section 8.22 is explicit: "
"the same measurement carries the opposite verdict. There, the failure of superposition was the point -- a "
"nonlinear field is what general relativity has. Here the nonlinear object is the wave function, and that is a defect "
"rather than a feature.")
body("Second, and more concretely absurd: a single particle gravitationally attracts itself. One particle, total "
"probability one, placed in a superposition of two locations has a density with two lumps, so the semiclassical field "
"pulls each lump toward the other. There is no second particle for it to respond to. Reading the force directly off "
"the field rather than inferring it from trajectories -- so that ordinary dispersion cannot confound the measurement "
"-- the attraction equals the Newtonian pull of a point mass of half the particle's own mass, to a ratio of 1.000 "
"once the lumps are well separated. This is not a small correction: it is the full gravitational force of a partner "
"that does not exist.")
body("Third, and sharpest, the case that requires no interpretation of quantum mechanics at all. Let a classical coin "
"flip decide whether a mass sits at one location or another. This is a proper mixture -- the mass really is at one of "
"them -- and every account of quantum theory agrees that the gravitational field is the field of the actual location. "
"Semiclassical gravity sources on the ensemble average and places the field at the midpoint, so it predicts that a "
"test mass released at the centre stays put. Measured here, the branch-wise force is the full Newtonian value while "
"the semiclassical force is zero to machine precision. The test mass must fall, and the theory says it will not. "
"Because the randomness is classical, no interpretation of quantum mechanics is available to rescue the prediction. "
"This is the Page-Geilker experiment of 1981, and it is the reason semiclassical gravity cannot be the final account.")
body("The model's own structure supplies the direction of the repair. In this project h is not a fundamental classical "
"field: it is a collective mode of A quantum condensate, and collective modes of quantum media are quantised -- "
"phonons are. The model's own logic therefore says that h should be quantised, and that the semiclassical treatment "
"used throughout Section 8 is an approximation adopted for tractability rather than a claim about nature. Quantising "
"a single radiative mode and coupling it to the same two-branch matter state gives the exactly solvable spin-boson "
"problem: matter and geometry become entangled, the matter decoheres, and the numerics reproduce the closed-form "
"answer to sixteen digits. The semiclassical treatment of the identical setup gives coherence exactly one, forever -- "
"a classical field sourced on a vanishing expectation value never moves, so it can never entangle with anything and "
"can never decohere anything. The self-attraction also disappears, because each branch now sources its own field "
"instead of one classical field being forced to serve both.")
result("Result 8.25 — the classical geometry is demonstrably wrong, not merely approximate.", "Semiclassical gravity "
"breaks the superposition principle (residual 2.8×10^-15 with gravity off, 4.6×10^-1 with it on), makes a single particle "
"attract itself with the full Newtonian force of a partner that does not exist (ratio 1.000 to the point-mass value), "
"and gets wrong the Page-Geilker case in which the randomness is classical and the correct answer needs no "
"interpretation (branch-wise force 0.776, semiclassical force 0.000). The model's own structure indicates the "
"replacement: h is a collective mode of a quantum medium and so should be quantised, and quantising one mode produces "
"matter-geometry entanglement and gravitational decoherence, matching the exact solution to ~10^-16, while removing "
"the self-attraction. honest limit: this quantises linearised gravity -- perturbative quantum gravity as an effective "
"theory, which is the easy and long-known part. It is not a theory of quantum geometry, it is silent on the "
"nonperturbative problem, and it does not solve the measurement problem: it relocates the branch structure into "
"matter-field entanglement without selecting an outcome. Single-mode coherence recurs at t = 2 π, as it must for one "
"mode; irreversible decoherence requires a continuum. The results of Sections 8.18-8.23 concern regimes without "
"macroscopic superposition, where the mean-field treatment is a controlled approximation -- but the framework itself "
"cannot be the final story, and this is the measurement of why.")

heading("8.26  Auditing the largest claim: what \"infrared fixed point\" can carry", 2)
body("The largest claim in this report is that general relativity is reached as an infrared fixed point, and it was "
"also the least supported. The phrase is a renormalisation-group statement, and until now no "
"renormalisation-group analysis had been performed anywhere in the project: no operator had been classified as "
"relevant or irrelevant, and no statement had been made about what the infrared forgets. The phrase was carrying the "
"authority of a calculation that had not been done. This section performs it, and then narrows the claim to what the "
"calculation supports.")
body("Operationally a fixed point asserts two things. First, universality: long-distance observables must lose their "
"dependence on the ultraviolet couplings. Second, an operator sorting: deformations must separate into irrelevant "
"ones the infrared forgets and relevant ones that destroy it. Both are computable in closed form here. Adding a "
"graviton mass to the deconfined propagator and factorising κ(q^{2}+A)(q^{2}+B) gives")
add_eq("G(r) = [ exp(-√B r) - exp(-√A r) ] / ( 4π κ (A - B) r ),     A + B = μ/κ,   AB = μ m^{2}/κ", "8.26a")
body("which returns the Section 8.11 form at m = 0, and which involves no box, no lattice and no periodic images, so "
"the tool cannot contaminate the quantity being tested. Varying the higher-derivative coefficient κ over four "
"decades at fixed Einstein coefficient, and probing each at equal depth into its own far field, the far-field law is "
"Newton to ten figures with force exponent -2 to five: κ sets only where the crossover sits, never what lies beyond "
"it. The residual dependence falls as exp(-r/ℓ), so the infrared forgets the ultraviolet exponentially rather than "
"as a power. That is universality, measured.")
body("The sorting is equally sharp, and because relevance is a statement about the limit rather than about any one "
"radius, each deformation is tracked outward. Higher-derivative structure decays as exp(-r/ℓ) and has vanished by a "
"hundred crossover lengths however large κ is made: irrelevant. A graviton mass ten thousand times smaller than the "
"Einstein scale instead grows with distance until it removes the inverse-square tail altogether: relevant. Since the "
"claim therefore rests entirely on that one relevant deformation being absent, the measured tail is converted into "
"an exclusion bound — any induced graviton mass above roughly 3e-5 of the Einstein scale would have shifted the "
"force exponent detectably, so the graviton's Compton wavelength exceeds the crossover scale by about four orders.")
result("Result 8.26 — the fixed-point claim, measured and narrowed.", "What is earned: the infrared forgets the "
"ultraviolet exponentially (the far-field law is unchanged to ten figures across 10000× in the higher-derivative "
"coefficient), and the operators sort as a fixed point requires — higher-derivative structure irrelevant, a graviton "
"mass relevant, with masslessness bounded to ~3e-5 of the Einstein scale. The linearised Einstein term is the "
"attractor of the long-distance theory, and that is a genuine renormalisation-group statement rather than a phrase. "
"What is not earned, and what the wording elsewhere has been corrected to reflect: the fixed point is empirical, not "
"protected. Diffeomorphism invariance is not an exact lattice symmetry (Section 8.12), so nothing forbids an induced "
"graviton mass; masslessness here is measured, not guaranteed, which is a weaker object than the photon's "
"symmetry-protected masslessness. The analysis is of the linearised propagator and reaches neither the nonlinear "
"Einstein equations nor a measurement of γ = 1, which remains an argument from Weinberg because the direct check is "
"regulator-limited. The report accordingly claims a linearised infrared attractor, not general relativity as a fixed "
"point. Section 8.27 then measures that lack of protection directly, and finds it: the tetrad "
"sector acquires an O(1) term at q = 0 while the photon's stays at 1e-10.")

heading("8.27  Is a graviton mass induced? The assumption the arc rests on", 2)
body("Section 8.26 identified a graviton mass as the one relevant deformation, so the infrared attractor stands or "
"falls on its absence — but that section bounded only a mass it had itself inserted into the propagator. The "
"propagator used everywhere else in this report, κ∇⁴ - μ∇², has no mass term written in it at all. Masslessness has "
"therefore been an assumption of the operator form throughout, never a measurement. The induced-gravity calculation "
"makes the omission explicit: Section 8.12 extracts μ as the q² coefficient of Π(q) - Π(0) and discards Π(0) as a "
"contact term. That discarded number is exactly the mass candidate, and it had never been examined.")
body("The question is posed so that no perturbative bookkeeping can conceal the answer. Rather than assembling a "
"bubble and a seagull and hoping the set is complete, the ground-state energy of the filled fermion sea is computed "
"as a function of a constant deformation. A constant deformation is precisely the q → 0 limit, and the sea energy "
"contains every order at once — bubble, seagull, and everything above — so nothing can be left out. Two deformations "
"are compared in the same regulator. A constant gauge field, k → k + A, is the photon's mass term, and gauge "
"invariance requires it to vanish: on a torus a constant A merely re-samples a complete period. A constant traceless "
"cone anisotropy, v = (1+ε, 1-ε, 1), is the h₊ polarisation of the tetrad metric of Section 8.5.")
body("The contrast is unambiguous. The photon's curvature is -1.5e-10 and improves steadily under refinement "
"(1.9e-7 at N = 24 falling to 5.6e-10 at N = 48): zero, and zero because a symmetry forbids it. The tetrad "
"graviton's is -0.27 — order unity, eight orders above the control, and stable both across the fermion gap and "
"under refinement. The sign is worth stating carefully, because it is not the naive expectation: it is negative, so a "
"uniform shear lowers the sea energy rather than costing it. The fermion loop does not merely give the cone a mass, "
"it destabilises the symmetric cone, and the medium's own shear rigidity must overcome that for the undeformed cone "
"to be stable at all. Either sign settles the question at issue: the q = 0 term is order unity and nothing forbids it.")
body("This is consistent with the report rather than damaging to it, and it corroborates an earlier result by a "
"completely different route. A tetrad mode carrying an order-unity term at q = 0 is not a massless long-range "
"mediator — which is precisely why Section 8.9 measured NO long-range force from the tetrad's inverse-square field, "
"and why the elastic route was declared dead on Eshelby-Crum grounds. Two unrelated calculations agree, and the "
"negative sign explains further why that route could never have been rescued by tuning: the fermion contribution "
"pushes the wrong way.")
result("Result 8.27 — masslessness cannot be inherited from the tetrad.", "Under a constant deformation, which is "
"the q → 0 limit computed to all orders at once, the photon's q = 0 term is -1.5e-10 and falls under refinement "
"while the tetrad graviton's is -0.27: order unity, eight orders larger, stable across gap and grid. The photon is "
"massless because an exact lattice symmetry forbids the term; the tetrad mode has no such protection and acquires "
"one, with a negative sign that destabilises the symmetric cone rather than simply making it massive. This is the "
"concrete content of Section 8.26's finding that the infrared attractor is empirical rather than protected, and it "
"independently corroborates the death of the elastic route in Section 8.9. What it does not show: the gravity this "
"report claims is the DECONFINED CURVATURE sector of Sections 8.11-8.12, not the tetrad, and its masslessness rests "
"on the biharmonic structure of the curvature field rather than on any symmetry protecting the cone shape. That "
"sector's q = 0 term is not measured here. The assumption is therefore narrowed, not removed — masslessness cannot "
"be inherited from the tetrad, so the curvature sector must supply it on its own. Section 8.28 then shows that this "
"is not a separate assumption at all: the curvature sector's q = 0 term IS the cosmological constant, and it "
"vanishes by the mechanism of Section 8.13.")

heading("8.28  The curvature sector's q = 0 term is the cosmological constant", 2)
body("Section 8.27 left the deconfined curvature sector's q = 0 term as the sharpest load-bearing assumption in "
"the gravitational arc. It is not a separate assumption. The q = 0 term of the induced graviton action is the "
"cosmological constant, as Section 8.13 already noted in passing — the induced Pi(0,0) and the vacuum stress ⟨T^{ij}⟩ "
"are the same object — and the algebra connecting them is elementary. A vacuum energy enters the action as √g Λ, and "
"expanding the metric determinant about flat space, g = δ + h, gives √g = 1 + ½ tr h + ⅛(tr h)² - ¼ tr(h²) + O(h³). "
"The piece quadratic in h therefore carries no derivatives at all, which is precisely what a mass term is:")
add_eq("S = -(Λ/8) [ 2 h_{ij} h_{ij} - (tr h)^{2} ]", "8.28a")
body("For a transverse-traceless perturbation this reduces to -(Λ/4) h_{ij} h_{ij}, a mass for the propagating "
"spin-2 modes themselves rather than merely a trace or conformal term — a point verified numerically here rather "
"than asserted, the quadratic expansion agreeing with the closed form to 8e-10 with the transverse-traceless cases "
"coming out at exactly -0.5. So m² is proportional to Λ, and whether the curvature graviton acquires a mass is the "
"cosmological-constant question in different words.")
body("Taken bare, the answer is as fatal as the tetrad's: the filled-sea energy density is 1.62 per site, order "
"unity in lattice units. What removes it is the result of Section 8.13. A self-sustained condensate vacuum "
"gravitates not its bare energy density but its grand potential, ρ_Λ = ε - μn = -P, and P vanishes identically at "
"self-sustained equilibrium. Measured across bare energies from 1 to 10^{122}, that cancellation holds to the "
"floating-point precision available — the largest cases cancel exactly in double precision, so their residual is "
"bounded by rather than equal to zero. Since m² is proportional to Λ, a relative residual of 1e-16 in Λ is 1e-8 in "
"the mass. A control confirms this is not a trivial identity: a rigid vacuum whose density cannot self-adjust "
"retains a ratio of 0.75 at every scale, so the equilibrium condition is doing real work.")
result("Result 8.28 — the last assumption is the cosmological-constant result.", "The curvature sector's q = 0 "
"term is the cosmological constant: expanding √g gives the derivative-free quadratic term -(Λ/8)[2h_{ij}h_{ij} - "
"(tr h)²], whose transverse-traceless part is nonzero, so m² is proportional to Λ for the propagating spin-2 modes. "
"Bare it is order unity, as fatal as the tetrad's; the self-sustained condensate vacuum removes it, because the "
"gravitating quantity is the grand potential -P, which cancels to the available precision over 122 decades against a "
"rigid-vacuum control that retains 0.75. The residual induced mass is ~2e-8 of the bare scale, roughly a "
"thousandfold below the ~3e-5 exclusion bound of Section 8.26, so the linearised infrared attractor survives its own "
"decisive test. Honest ceiling: this is protection by an equilibrium condition, not by a symmetry. It is exact at "
"equilibrium, but a self-sustained vacuum is a dynamical state rather than a redundancy of description, so it is "
"weaker than the photon's protection, which no dynamics can spoil. It is also not independent evidence — it is the "
"same Volovik mechanism as Section 8.13, so the two stand or fall together. What is genuinely gained is that the arc "
"has one fewer open assumption than it appeared to.")

heading("8.29  Is the induced graviton action diffeomorphism invariant?", 2)
body("Every statement about γ = 1 in this report has rested on Weinberg's theorem, and that theorem has two "
"hypotheses, not one: a massless spin-2 field, and a quadratic action invariant under the linearised "
"diffeomorphism h_{ij} → h_{ij} + ∂_{i}ξ_{j} + ∂_{j}ξ_{i}, coupled to a conserved source. Sections 8.27 and 8.28 "
"addressed the first. The second has always been assumed. Two earlier attempts to measure it failed and, between "
"them, concluded that it could not be measured: Section 8.12's Ward identity was broken by a hard-cutoff surface "
"term severe enough to spoil even the photon validation, after which the lattice analysis identified two "
"obstructions — the identity is inhomogeneous because the induced vacuum stress ⟨T^{ij}⟩ is nonzero, and "
"diffeomorphism invariance is not a lattice symmetry.")
body("The first obstruction has since dissolved. That ⟨T^{ij}⟩ is the same object as the q = 0 graviton mass, which "
"Section 8.28 identified as the cosmological constant and cancelled through the self-sustained vacuum, so the "
"identity can legitimately be made homogeneous and the question reopened. It is reopened here with the instrument "
"of Sections 8.27 and 8.28 rather than a perturbative bubble: the ground-state energy of the filled sea under a "
"finite-wavelength background deformation, which contains every order and every seagull at once. The point is not "
"pedantic — the earlier graviton bubble carried no seagull at all while the photon bubble beside it did, so the two "
"were never comparable.")
body("The instrument is calibrated twice before use. At q = 0 it reproduces Section 8.27's tetrad mass to ten "
"digits by an independent code path. At finite q the photon, whose U(1) invariance is exact on the lattice under "
"the Peierls substitution, shows a pure-gauge response of 2.2e-11 — the round-off floor of the derivative stencil — "
"against a transverse response of 6.4e-4, a ratio of 3e-8. That is the calibration the earlier attempt lacked.")
body("The graviton then fails. Its pure-gauge response is nonzero, and although most of that is the q = 0 mass "
"riding along, removing the mass leaves a two-derivative violation that does not go away. The decisive feature is "
"that the ratio of the violating response to the invariant one is flat in q, at 1.07 and 4.06 for the two "
"independent gauge modes. A diffeomorphism-violating operator holding a fixed fraction of the invariant one is "
"marginal rather than irrelevant, and by the criterion established in Section 8.26 it therefore never flows away: "
"diffeomorphism invariance is not emerging in the infrared. Rotational invariance fails in the same way. The two "
"transverse-traceless polarisations for q along z are related by a 45-degree rotation and must be degenerate; they "
"split by 12.4 per cent, converged in q and in the transverse grid alike. Fitting the four two-derivative "
"invariants against the unique linearised Einstein-Hilbert combination (1, -2, 2, -1) returns (1, 8.18, -0.56, "
"-0.12) — not a small deformation of Einstein-Hilbert, and the second coefficient does not even have the right sign.")
result("Result 8.29 — Einstein structure cannot be inherited from the tetrad.", "Measured rather than inferred, "
"with an instrument calibrated against both a known q = 0 result and an exact finite-q photon Ward identity. The "
"induced tetrad action does not annihilate the pure-gauge modes, is not rotationally invariant, and is not "
"Einstein-Hilbert. The residual diffeomorphism violation is marginal rather than irrelevant, so it survives the "
"infrared limit that Section 8.26 requires it to fail. This is a negative on a route already known to be dead: the "
"tetrad is not this model's gravity, as Sections 8.9 and 8.27 established from different directions, and the "
"deconfined curvature sector is untouched by the measurement. What changes is that the door Section 8.27 left ajar "
"is now shut in both directions — the tetrad supplies neither masslessness nor Einstein structure — and γ = 1 rests "
"entirely on Weinberg applied to the curvature sector. Measuring that directly is the outstanding problem of the "
"gravitational arc, and it can no longer be shortcut through the cone.")

heading("8.30  Can a projection rescue it? The induced form has no gauge null space", 2)
body("Section 8.29 leaves one hope standing. Its measurement is of the tetrad sector, whose degrees of freedom are "
"the compatible strain of the Dirac cone, whereas the gravity the model actually claims is the deconfined curvature "
"sector, which propagates only incompatible strain. If the physical curvature modes were cleanly separated from the "
"pure-gauge modes that carry the violation, the failure might live entirely in a subspace the curvature propagator "
"projects away, and γ = 1 could survive for the gravity that matters. That hope is a precise statement — it asks "
"whether the induced quadratic form is block-diagonal between physical and gauge modes — and it can be measured.")
body("It is measured here as the full six-by-six quadratic form on symmetric h_{ij} at fixed q, every diagonal and "
"every cross term, in an orthonormal basis split into three physical modes (the two transverse-traceless "
"polarisations and the transverse trace) and three gauge modes (the two spin-1 shears and the longitudinal mode). "
"Linearised diffeomorphism invariance makes an unambiguous prediction, stated before the result: the three gauge "
"directions must be exact null vectors, so the form must have a rank-three kernel spanned by them. Einstein-Hilbert's "
"own form on the same basis has exactly that kernel.")
body("The induced form has none. Its six eigenvalues are all nonzero, and the gauge directions are not its flat "
"directions but its stiffest, four to eight times larger than the physical modes. The block that decides the hope is "
"the physical-gauge mixing, and it is 42 per cent of the physical block itself, concentrated in the spin-0 sector — "
"the transverse trace coupled to the longitudinal gauge mode — which is exactly the sector that sets the spatial "
"curvature Ψ, hence γ, for a spherical source. A projection cannot remove a violation that is coupled into the modes "
"it retains. The spin-2 doublet fails independently, through the 12.6 per cent rotational-anisotropy split of Section "
"8.29, which no projection onto incompatible strain can reach at all because spin-2 cannot mix with the spin-1 gauge "
"modes.")
result("Result 8.30 — the projection loophole is closed.", "The induced tetrad form has no gauge null space: six "
"nonzero eigenvalues where Einstein-Hilbert has three zeros, with the gauge directions the stiffest modes rather than "
"the flat ones. It is not block-diagonal between physical and gauge subspaces — the mixing is 42 per cent of the "
"physical block, in the spin-0 sector where γ lives — so no projection onto the incompatible/curvature subspace can "
"recover an Einstein-Hilbert form. Section 8.29's negative is therefore structural rather than a near miss: the "
"induced tetrad action is not a deformation of Einstein-Hilbert in any subspace, and γ = 1 rests entirely on the "
"deconfined curvature sector measured directly.")

heading("8.31  The nonlinear self-coupling is fixed by the bootstrap, not free", 2)
body("Section 8.22 gave the gravitational field the cubic self-coupling that makes gravity gravitate and swept its "
"strength λ as a free parameter. It is not free. Deser's bootstrap requires a spin-2 field coupled to matter's stress "
"tensor to couple to its own stress tensor at the same strength, since the matter stress alone ceases to be conserved "
"once the field reacts back; iterating the requirement rebuilds the Einstein-Hilbert action, and its first step fixes "
"the cubic vertex. In this project's Hamiltonian the constraint is visible directly, because matter and the field "
"enter the same field equation as two source terms of identical form — matter as (g/2) h_{ab} S_{ab} and the field "
"as (λ/2) h_{ab} ∂_a h_{ij} ∂_b h_{ij}, whose contraction is the field's own stress tensor. Their ratio λ/g is a "
"Nordtvedt parameter, the amount by which gravitational binding energy gravitates relative to ordinary energy, and "
"the strong equivalence principle sets it to one.")
body("The identification is verified independently, by the response of the field energy to a coordinate deformation "
"— the constant-deformation method of Sections 8.27 and 8.28, which never refers to the contraction — agreeing on "
"all six stress components, shears included, to 1e-10. Against it, Section 8.22's five values λ = 0, 0.4, 0.8, 1.6, "
"200 at g = 6 are λ/g = 0, 0.067, 0.13, 0.27 and 33.3. None is one. The headline self-interaction figure of that "
"section, a two per cent shift in the radiated energy, was taken at λ/g = 33.3, a theory in which gravitational "
"binding energy gravitates thirty-three times too strongly; at the bootstrap value λ = g the same shift is six "
"hundredths of a per cent, and the energy budget still closes to a few parts in 10^{11}.")
result("Result 8.31 — the self-coupling was never free.", "Deser's bootstrap fixes the cubic vertex once the "
"quadratic term is Fierz-Pauli: in this Hamiltonian matter sources the field at g/2 and the field sources itself at "
"λ/2, so the strong equivalence principle reads λ = g and λ/g is a Nordtvedt parameter. The identification is "
"verified independently to 1e-10 on every stress component. Section 8.22's values are λ/g = 0, 0.067, 0.13, 0.27 and "
"33.3, none of them physical, and its headline self-interaction was measured at λ/g = 33.3; at λ = g the effect is "
"far smaller and the budget still closes. Nothing in the integration was wrong — the parameter was. The bootstrap "
"fixes the cubic term for the postulated Fierz-Pauli field of Sections 8.20-8.24, which reproduced the quadrupole "
"luminosity law and the Peters-Mathews inspiral; it does not connect that field to the induced action of Section "
"8.29, which is not Fierz-Pauli, and that gap remains the central open problem.")

heading("8.32  Does mass source curvature? The smooth coupling behind γ = 1, measured", 2)
body("The gap named at the end of Section 8.31 — that the model's actual gravity, the deconfined curvature sector, "
"has never been shown to reach γ = 1 — is the central open problem, and this section takes the first direct measurement "
"of it. Section 8.10 had already reduced γ to a single number: γ = κ/(4πG), where 4πG is the coupling with which a "
"static mass sources the Newtonian time potential Φ, and κ is the strength with which the same mass sources the "
"spatial curvature Ψ. The first half is measured and healthy; the question is the second. A static point mass is a "
"localized energy density, so it sources gravity only through its T^{00} component, and any spatial curvature it "
"produces must come from the induced coupling of that energy density to the spatial stress, Π^{00,ij}(q). γ = 1 needs "
"that coupling present at the Newtonian strength; γ = 0 is the value if the energy density sources no spatial stress "
"at all.")
body("Measured on the gapped Dirac cone with the static interband bubble of Section 8.12, calibrated against the "
"healthy induced photon, the coupling vanishes identically. Π^{00,ij}(q) is zero for every component, the trace "
"included, to machine precision and at every momentum, mass and cutoff, while the Newtonian Π^{00,00} is everywhere "
"nonzero. The reason is a selection rule rather than a small number: the energy-density vertex is proportional to the "
"identity, a scalar, and a scalar cannot source the spin-2 spatial stress. That the bubble is not simply blind is "
"confirmed by replacing the scalar mass with a genuine spin-2 source, which couples to the same stress at order unity. "
"Through the smooth induced loop a mass sources Φ and not Ψ, so the smooth-sector γ is zero.")
result("Result 8.32 — every smooth mechanism gives γ = 0; γ = 1 rests on emergent Weinberg alone.", "A static mass "
"presents a scalar energy density to gravity, and its induced coupling to the spatial stress vanishes identically — "
"Π^{00,ij} = 0 for every component including the trace, to machine precision across mass, cutoff and momentum — while "
"a genuine spin-2 source couples to the same stress at order unity, so the zero is a selection rule. This is the "
"fourth independent statement that the smooth route to γ = 1 is closed, joining the non-dynamical induced spatial "
"graviton (Section 8.12), the non-Einstein induced tetrad action (Section 8.29) and its absent gauge null space "
"(Section 8.30); pure elasticity gave the same zero from the other side (Section 8.10, a mass relaxing to a "
"compatible displacement). Every smooth mechanism the model has returns γ = 0. What is not touched is the topological "
"channel: genuine curvature here is incompatible strain, disclination density, which a smooth stress-stress bubble "
"cannot see, and whether a mass sources curvature through that channel is the remaining and harder measurement. As it "
"stands γ = 1 rests entirely on the emergent-Weinberg argument — a massless spin-2 on a conserved infrared stress "
"tensor is forced to be Einstein — which, unlike the model's emergent Lorentz invariance, has no direct in-model "
"confirmation and now has direct evidence against it in the smooth sector. This is the sharpest honest statement of "
"where the gravitational arc actually stands.")

heading("8.33  The topological channel: does mass carry net curvature charge? Gauss-Bonnet", 2)
body("Section 8.32's selection rule is specifically about spin — it forbids the scalar energy density from sourcing "
"the spin-2 part of the spatial metric. Curvature itself is a scalar, the Ricci scalar, and Section 8.9 located a "
"genuine scalar curvature degree of freedom in the medium: the third bond fluctuation per site, beyond the two "
"displacement modes. A scalar mass sourcing a scalar curvature is not forbidden by any spin rule, so this topological "
"channel is the one route left that could give γ = 1 directly, through the model's own gravity mechanism. That "
"mechanism is compression: field energy changes the medium's preferred local density, an eigenstrain θ*(x) ∼ ρ(x), "
"the non-uniform thermal expansion that buckles a heated plate. Unlike a body force, which relaxes to a compatible "
"displacement, an inhomogeneous eigenstrain leaves real Gaussian curvature η = ∂²θ* that light, riding the bond "
"metric, must see.")
body("It does, but of the wrong kind. The curvature is nonzero locally, yet its net charge vanishes identically: "
"∫η = ∫∂²θ* is a total derivative, zero to machine precision for every profile. A smooth localized compression makes "
"a curvature dipole — a dome in the compressed core, a saddle in the surrounding ring — with no net deficit angle, "
"which is the Gauss-Bonnet theorem: smooth deformations carry no topological curvature charge. Zero net charge means "
"no long-range bending, and the measured deflection from the compression curvature is zero at every impact parameter, "
"beside a genuine curvature charge η ∼ ρ that bends light and stays long-range. γ = 1 needs the light-bending charge "
"to be the mass itself — a disclination density proportional to ρ — and compression supplies none.")
result("Result 8.33 — every accessible direct channel gives γ = 0, each for its own reason.", "The model's gravity "
"is compression, a mass entering as an eigenstrain, and while that does curve the medium locally (η = ∂²θ* ≠ 0) its "
"net curvature charge is exactly zero by Gauss-Bonnet, so it bends no light at range. This closes the last accessible "
"direct channel: the smooth induced loop gave zero by a spin selection rule (Section 8.32), the elastic body force by "
"strain compatibility (Section 8.10), and compression by Gauss-Bonnet. Each sources zero net curvature charge, which "
"is carried only by a topological disclination, and mass nucleates none. So the model's gravity as directly realized "
"couples to mass as a scalar, through the Newtonian T^{00} alone — Nordström, γ = 0, not Einstein. The dimensional "
"caveat is stated plainly because the result is a negative: Gauss-Bonnet is a 2D theorem and these measurements are "
"2D and 2+1D, the settings where the incompatible-curvature picture is defined, whereas in 3+1D a smooth mass sources "
"smooth Ricci curvature with no topological defect, which is γ = 1. Whether the 3+1D medium evades the 2D obstruction "
"the way general relativity does is a further calculation not done here. What stands is that γ = 1 has no direct "
"realization in any channel the model exposes at the accessible scale, and rests entirely on the emergent-Weinberg "
"argument — the same infrared-emergent status as the graviton Ward identity — which, unlike the directly confirmed "
"emergent Lorentz invariance, has no positive in-model evidence and now a consistent set of direct results against it. "
"Newtonian gravity is real and healthy; the Einstein completion is argued, not demonstrated.")

heading("8.34  The 3+1D answer: the medium does not evade the obstruction, and why", 2)
body("Section 8.33 left one honest opening: Gauss-Bonnet is a two-dimensional theorem, and in 3+1D a smooth mass "
"sources smooth Ricci curvature with no topological defect, which is γ = 1, so whether the medium reaches the Einstein "
"value in the dimension that matters was not yet settled. This section settles it, in genuine 3D by direct ray "
"tracing, and the answer is that the medium does not reach γ = 1 — for a reason deeper than Gauss-Bonnet and "
"independent of dimension. In general relativity the spatial potential obeys a Poisson equation, ∂²Ψ = 4πGρ, so Ψ is "
"the long-range Newtonian potential, the same 1/r that carries the time potential Φ, and γ = Ψ/Φ = 1. The medium does "
"not solve a Poisson equation for its spatial metric. Its gravity is compression: a mass is an eigenstrain θ* ∼ ρ, and "
"the 3D incompatibility of an isotropic eigenstrain is δ_{ij}∂²θ* − ∂_i∂_j θ*, exactly the linearized Einstein tensor "
"of h_{ij} = 2θ*δ_{ij} — verified numerically to 10^{-16} — so the spatial metric follows the compression "
"algebraically, Ψ = θ* = ρ. That is local, dying with the mass rather than with the potential.")
body("Ray-traced in 3D, the consequence is unambiguous. A general-relativity control with Ψ set to the potential "
"holds γ = 1 at every impact parameter, confirming the method; the medium's γ, with Ψ = ρ, falls to zero with "
"distance — measured to a few parts in 10^{4} at the far field and stable under refinement, box size and source "
"width — because its spatial bending is short-range while the Newtonian bending is long-range. The 2D zero-charge "
"result of Section 8.33 was a symptom of this single fact: the medium sets the spatial metric equal to the local "
"compression, where general relativity sets it equal to the potential.")
result("Result 8.34 — γ = 0 in 3+1D too; the spatial metric follows the density, not the potential.", "The direct "
"search is now finished in the dimension that matters. In 3D, as in 2D, a mass entering as a compression eigenstrain "
"gives a spatial metric Ψ = θ* = ρ that follows the mass density locally, while the Newtonian Φ is the long-range "
"potential, so γ = Ψ/Φ falls to zero at range — measured by ray tracing against a general-relativity control that "
"holds γ = 1 at every impact parameter. The reason is dimension-independent: general relativity sources the spatial "
"metric through a Poisson equation and the medium sets it algebraically by compression, so Gauss-Bonnet was a symptom, "
"not the cause. Every channel in every accessible setting now gives γ = 0 — a spin selection rule (Section 8.32), "
"strain compatibility (Section 8.10), 2D Gauss-Bonnet (Section 8.33), and 3D local compression (here). γ = 1 requires "
"mass to source the propagating graviton's spatial polarizations as a Poisson-sourced potential, which is precisely "
"the coupling Section 8.32 measured to vanish: the graviton is massless and healthy (Section 8.12), but mass does not "
"couple to its spatial modes, only to the scalar compression. The premise of the emergent-Weinberg argument — that "
"mass couples to the full conserved stress tensor — is what fails, and it fails directly, in 3D. The bottom line for "
"the gravitational arc is now precise: Newtonian gravity is real, healthy and quantitative — the quadrupole law and "
"the Peters-Mathews inspiral stand — while the Einstein completion, γ = 1 and the light-bending factor of two, is not "
"realized by any direct mechanism in any dimension the model exposes, and survives only as an emergent-Weinberg "
"argument whose premise these measurements contradict. It is argued, not demonstrated.")

heading("8.35  The last route: can the medium's Poisson ratio make the graviton Fierz-Pauli?", 2)
body("Sections 8.32-8.34 closed every way a mass could SOURCE spatial curvature, but one route remained at the level "
"of the PROPAGATOR: γ = 1 also follows if the graviton's kinetic term is Fierz-Pauli, so that a mass sourcing the "
"time potential produces spatial curvature through the propagator's trace structure. Sections 8.29-8.30 showed the "
"fermion-induced tetrad action is not Fierz-Pauli, but the deconfined curvature sector's kinetic term is the medium's "
"own biharmonic elasticity, whose trace structure is set by its elastic constants — its Poisson ratio. This is the "
"classic elasticity-as-gravity question: a linear-elastic medium yields an effective linearised gravity whose "
"parameters depend on its Poisson ratio, and for one special value the graviton could in principle be Fierz-Pauli. "
"That was the last opening, and this section shuts it.")
body("Measured by straining the Lennard-Jones medium, the condensate is a central-force Cauchy solid with Poisson "
"ratio ν ≈ 1/3, but the value turns out not to matter. The medium's response to a mass splits into an elastic "
"relaxation — a displacement field chosen by the elastic constants — and the eigenstrain the mass imposes. Only the "
"relaxation depends on the Poisson ratio, and a relaxation is by construction a displacement, so its strain is "
"compatible and its curvature is identically zero: solving the 3D relaxation across the whole physical range of "
"Poisson ratios, from ν = −0.9 to ν = 0.49, the linearised Ricci scalar of the relaxed strain is machine zero at "
"every value. The light-bending curvature comes entirely from the incompatible part of the eigenstrain, which is "
"fixed by the mass and independent of the Poisson ratio, so γ ray-traced across the whole range is the same at every "
"value — the zero of Section 8.34. Tuning the medium's elasticity moves only a gauge degree of freedom, which drops "
"out of the gauge-invariant curvature, so it cannot carry the graviton's trace structure toward Fierz-Pauli.")
result("Result 8.35 — no Poisson ratio makes the graviton Fierz-Pauli; every route is now closed.", "The medium's "
"response to a mass is an elastic relaxation plus the imposed eigenstrain, and only the relaxation depends on the "
"Poisson ratio. A relaxation is a displacement, so its strain is compatible and its curvature is identically zero at "
"every Poisson ratio (verified from ν = −0.9 to 0.49, machine zero throughout), and it bends no light; the "
"light-bending curvature is the incompatible eigenstrain part, independent of ν, so γ ray-traced across the whole "
"range is the same value zero. The condensate's own ν ≈ 1/3 (a central-force Cauchy solid, measured by straining the "
"Lennard-Jones medium) is therefore irrelevant. This closes the propagator route the way Sections 8.29-8.30 closed "
"the induced-tetrad one: the trace structure that would give γ = 1 cannot be reached by tuning the elastic constants, "
"because they multiply only the compatible sector. With this, γ = 1 is reachable neither by the source — a mass "
"sources no curvature, by a spin rule, Gauss-Bonnet and local compression, in every channel and dimension — nor by "
"the propagator. The graviton is massless and healthy, but a static mass couples to none of its spatial polarisations "
"by any mechanism the medium provides. The gravitational arc rests where the measurements put it: Newtonian gravity "
"real, healthy and quantitative; the Einstein completion argued through emergent Weinberg, its premise directly "
"contradicted, and now shown not to be recoverable by any property of the medium the model can adjust.")

heading("8.36  The negative-space sublattice: does a second lattice supply the missing term?", 2)
body("Sections 8.29-8.35 all rested on one premise: the medium is a single Bravais lattice with central forces. "
"That premise forces the Cauchy relations and, more deeply, makes the light-seen strain the gradient of a single "
"displacement field, which Saint-Venant compatibility then makes flat. A medium of two interpenetrating lattices — "
"fcc plus its negative space, the bcc-like interstitial — breaks it: it carries a relative, or optical, mode "
"w = u_A − u_B that is not the gradient of the acoustic displacement. This is textbook — two interpenetrating fcc "
"lattices form diamond, which violates the Cauchy relations badly through exactly this internal-strain (Kleinman) "
"mode — so it is the one construction that attacks the real assumption. Measured in the clean 2D proxy, the "
"honeycomb of two triangular sublattices.")
result("Result 8.36 — the second sublattice adds a displacement, not an incompatible degree of freedom.", "The "
"optical mode is gapped, its q = 0 frequency scaling with the inter-sublattice coupling, so it is adiabatically "
"slaved to the strain — a second displacement, not an independent long-range field. It does genuinely break the "
"Cauchy relations (relaxed C₁₂ − C₆₆ = +0.29, the diamond departure) while leaving the compression modulus invariant "
"to fourteen digits and softening only shear: the internal mode couples to shear, never to compression, and an "
"isotropic strain sources zero internal shift (machine zero, by the three-fold symmetry). The full two-sublattice "
"response to a mass — macro displacement plus slaved optical shift — has incompatibility (linearised Ricci) machine "
"zero, against a live eigenstrain probe at 0.125: because the optical mode is slaved it is itself a displacement, so "
"the total strain is a symmetric gradient, compatible, flat. The negative-space sublattice cannot escape the "
"compatibility trap of Sections 8.34-8.35; γ = 0 survives the two-sublattice generalisation.")

heading("8.37  The whole γ arc in one number: the graviton propagator's trace coefficient", 2)
body("The entire Einstein question collapses to a single coefficient. Linearised gravity sends a static source to a "
"metric through the graviton propagator, whose only free piece for a massless spin-2 is the coefficient λ of the "
"trace term. For the pressureless sources that actually bend light this gives γ = λ/(1 − λ): the factor of two is the "
"trace term and nothing else, and pressure is a red herring — a source at the Sun's p/ρ ~ 10⁻⁵ moves γ by ~10⁻⁵, "
"nowhere near the factor of two. Fierz-Pauli, and general relativity, sit at λ = 1/2. Both of the model's measured "
"routes put it at λ = 0: the source coupling ⟨T00, T_ij⟩ = 0 of Section 8.32 (no trace term), and the induced "
"tetrad action's coefficients (1, +8.18, −0.56, −0.12) of Section 8.29 against Einstein-Hilbert (1, −2, 2, −1), off "
"by order one with the second sign wrong.")
result("Result 8.37 — γ = 1 and emergent Lorentz have opposite renormalisation-group fates; that is the whole answer.", "γ = 1 "
"and emergent Lorentz invariance are the same kind of statement — a continuum symmetry the lattice breaks — decided "
"by the same renormalisation-group question, and they get opposite answers. Lorentz's cone anisotropy is irrelevant, "
"scaling as (k/k_Planck)² → 0, so it flows away and the symmetry becomes exact in the infrared; this is why "
"relativity is real here. The graviton's diffeomorphism violation is marginal — Section 8.29's ratios 1.07 and 4.06 "
"hold flat as q → 0, and the rotational anisotropy converges to 12.4 per cent — so it holds a fixed fraction of the "
"invariant term and never flows away, leaving γ pinned off 1 at every scale. The model earns emergent special "
"relativity, fermions, most of quantum mechanics, electromagnetism and a real Newtonian gravity, but not the "
"Einstein factor of two, because the one coefficient that would deliver it is marginal, not irrelevant: no continuum "
"limit of this lattice recovers it. γ = 1 would require a background-independent construction, where the metric is "
"the field rather than a displacement on a fixed lattice — a different model, not a refinement of this one.")

heading("8.38  Does the photon's width bend light?", 2)
body("Every γ measurement so far ray-traced light as a zero-width geodesic, so 'light is a thin ray' is a genuine "
"buried premise, worth testing rather than assuming. The factor of two, however, is not a property of the photon but "
"of the geometry it moves through. Deflection is (1 + γ) × Newton, the '1' the time warp g₀₀ carried by a scalar "
"refractive index n = 1 − Φ, the 'γ' the spatial curvature Ψ that adds a second −Ψ to the index. Light samples both, "
"and samples them equally, for one reason — it moves at c: on a null path dx = c dt the space and time parts of the "
"metric enter with equal weight. That null-ness, not any transverse width, is the origin of the factor of two, and "
"the model already has null light while sourcing no Ψ.")
result("Result 8.38 — width multiplies whatever γ the geometry carries, and never creates it.", "A finite-width "
"photon — its centroid deflection exact by Ehrenfest's theorem, hence diffraction-independent — through the model's "
"index bends by the ray amount, implied γ climbing only from 0.0004 to 0.15 as the width grows to a third of the "
"impact parameter, an O((w/b)²) tidal correction. Through a general-relativistic index the same beam bends by twice "
"the Newtonian amount at every width. So the width average sees the factor of two when the geometry has it and not "
"when it does not: width multiplies the γ the geometry carries but never manufactures it. The model's γ = 0 is "
"robust to giving the photon any width, because a mass sources no spatial curvature for a probe of any size to feel. "
"Einstein's factor of two is about light's speed, not its width; the model has the speed and not the curvature.")

heading("8.39  The Lorentz-violation prediction, confronted with data", 2)
body("The model's one falsifiable empirical claim (Section 8.8) is a quadratic, mass-dimension-6 Lorentz violation, "
"subluminal and species-universal, with an effective scale E_QG ≈ 2.5 × 10¹⁹ GeV. Confronting it properly corrects an "
"over-generous earlier reading. For a quadratic (n = 2) violation the effect grows as E², so the strongest bounds "
"come from ultra-high-energy astrophysics and sit near the Planck scale — UHECR and the GZK threshold near "
"10¹⁸ GeV, PeV photons near 10¹⁹ — not the 10¹¹ GeV that a naive comparison against the UHECR energy suggests. The "
"model's effect is about 10⁻³ of current UHECR sensitivity: within roughly one to two orders of falsification, the "
"closest its predictions ever come, and a real target for next-generation ultra-high-energy observatories.")
result("Result 8.39 — consistent, partly confirmed, and near the frontier.", "The model has already passed a test it "
"could have failed. Its one universal cone forces c_photon = c_electron = c_graviton = c_neutrino at leading order, "
"and GW170817 confirmed it — gravitational waves and light within about a second over 130 million light years, "
"|Δc|/c < 3 × 10⁻¹⁵ — backed by neutrino multimessenger timing. Many Lorentz-violation scenarios predict a "
"species-dependent maximal speed and are constrained by exactly these observations; the one-structure construction "
"forbids the difference, and the data agree. The model's distinctive crystallographic anisotropy is a genuine "
"preferred-frame effect but doubly suppressed and likely averaged away in a poly-domained medium, so one-cone "
"universality, not the cubic pattern, is the practical discriminator. Its kill conditions need no Planck-energy "
"access — a confirmed linear (n = 1) photon dispersion, a species-dependent speed, or a superluminal photon would "
"each falsify it — and all are consistent with present data. The real experimental home is ultra-high-energy "
"astrophysics, not laboratory cavities.")

heading("8.40  The cosmological constant: the fine-tuning catastrophe dissolved, the observed value open", 2)
body("Integrating out the medium's modes gives a vacuum energy of order the node scale — with the node spacing fixed "
"at the Planck length, of order M_Planck⁴, about 10¹²² times the observed dark-energy density. Taken at face value "
"this is the worst prediction in physics. But the model's vacuum is a self-bound condensate, and what gravitates "
"cosmologically is not the bare energy density but the pressure: a self-sustained medium sits at its energy minimum, "
"where the pressure P = 0, so the cosmological constant is zero regardless of how large the energy density is. This "
"is Volovik's condensed-matter mechanism, and it is now measured rather than assumed. Its load-bearing claim — that "
"emergent gravity couples to the vacuum stress −P and not to the energy density — is verified by the induced-gravity "
"bubble: the energy density sources the Newtonian potential h₀₀ (⟨T00 T00⟩ ≠ 0) but sources no spatial curvature "
"(⟨T00, T_ij⟩ = 0, a symmetry zero), while a genuine spatial stress does. So the vacuum's enormous energy gravitates "
"only as an unobservable uniform h₀₀, and the cosmological constant is its pressure, which self-tunes to zero.")
body("Strikingly, that verifying fact is the same one that closes the light-bending arc. The selection rule "
"⟨T00, T_ij⟩ = 0 means energy density sources no spatial curvature, and it has two faces: a static mass sources no "
"spatial curvature, so γ = 0 (no factor of two, Sections 8.32-8.37), and the uniform vacuum sources no spatial "
"curvature, so Λ ≈ 0 (its energy does not drive expansion). The model does not reach Einstein gravity for the very "
"reason it has no cosmological-constant problem; the factor of two and the 10¹²² are one fact, and a theory that "
"bent light like general relativity would carry the catastrophe this one avoids. The zero is robust. On the real "
"Lennard-Jones condensate a large energy density coexists with machine-zero pressure at equilibrium, restored to "
"zero after a phase transition that shifts the vacuum; and it survives Weinberg's no-go, because absorbing a uniform "
"vacuum-energy change is a pure dilatation of the medium — all energies and lengths rescale together — leaving every "
"dimensionless observable, the emergent sound speeds and cone ratios, invariant to machine precision across vacua "
"spanning eight orders of energy. The self-adjusting field, the medium's overall scale, is a flat direction the "
"vacuum energy couples to and observable physics does not. And equilibrium is a dynamical attractor: any initial "
"vacuum energy decays to it within a few Planck times, so there is no initial cosmological constant to fine-tune.")
result("Result 8.40 — the 10¹²² fine-tuning is dissolved; the observed nonzero value is not predicted.", "What the "
"model settles is the equilibrium cosmological constant: exactly zero, self-tuned by the condition P = 0 that also "
"defines equilibrium, robust to phase transitions and to Weinberg's no-go, and reached from any initial condition by "
"an attractor in a Planck time. The 10¹²² fine-tuning catastrophe is genuinely removed — not by cancelling constants "
"but because the equilibrium vacuum gravitates zero by thermodynamics — and it is unified with the light-bending "
"result as one measured selection rule. What the model does not settle is the small observed nonzero value, ~10⁻¹²² "
"in Planck units. The natural off-equilibrium source, a cosmic expansion lag holding the medium a fraction ~Hτ off "
"equilibrium, gives a dynamical dark energy ρ_Λ ~ M_Planck³H ~ 10⁻⁶¹ — the geometric mean of the problem, still 61 "
"orders too large, and with the wrong equation of state (w ~ −½ in the matter era). A node-creation that tracked the "
"expansion could in principle give ρ_Λ ~ ρ_crit, but deriving the tracking efficiency shows the medium's node "
"creation is deficit-driven — it responds to the local chemical potential, the tension P = Bδ, not to the global "
"Hubble rate — so the efficiency is zero and the residual is the linear 10⁻⁶¹, while perfect growth gives zero; "
"neither reaches 10⁻¹²². The catastrophe is removed; the observed dark-energy density is relocated to why the vacuum "
"sits so extraordinarily near equilibrium, and remains open.")

heading("8.41  The electroweak Higgs mechanism: SU(2)×U(1) broken to one massless photon", 2)
body("Section 8.14 induced non-Abelian Yang-Mills, and the screening arc measured the Abelian Higgs mechanism — a "
"gauged condensate giving its photon a Meissner mass. The Standard Model's electroweak sector is their non-Abelian "
"marriage: a Higgs doublet condensate breaks SU(2)_L × U(1)_Y down to a single U(1)_EM, so that three weak bosons "
"acquire a mass while the photon stays exactly massless. This runs in the model's own terms. The condensate — the "
"same amplitude mode that gaps for gravity (§8.10) and self-tunes the cosmological constant (§8.40) — is placed in a "
"doublet of SU(2)_L with hypercharge Y = ½, and its equilibrium amplitude is a vacuum expectation value ⟨Φ⟩ = "
"(0, v/√2). The induced gauge kinetic term |D_μ Φ|² then costs a mass for whichever gauge directions move the vacuum, "
"mixing the four fields (W¹, W², W³, B) into a mass matrix whose spectrum is the electroweak prediction.")
result("Result 8.41 — the condensate reproduces the electroweak breaking pattern and mass relations.", "Diagonalising "
"the induced 4×4 mass matrix gives the spectrum {0, m_W, m_W, m_Z} — exactly one massless boson and three massive, for "
"any couplings, with m_W = gv/2 and m_Z = √(g²+g'²) v/2 to the digit and the photon at machine zero. This is the "
"SU(2)×U(1) → U(1) signature: not zero massless bosons and not two. The measured W/Z mass ratio is exactly the cosine "
"of the weak mixing angle, m_W/m_Z = cos θ_W = g/√(g²+g'²) (sin²θ_W = 0.223 at the physical point), so the mixing "
"angle is not a separate input but the geometry of the W³–B mixing the condensate induces. And the surviving massless "
"photon is the electric-charge direction Q = T₃ + Y, kept massless because the vacuum is Q-neutral (Q⟨Φ⟩ = 0) while "
"the three orthogonal generators move it and gain mass — the single photon forced by the doublet's representation and "
"its hypercharge. As in Section 8.14 this is the mechanism, not a derivation: the group, the doublet representation "
"and Y = ½ are inputs, and what is shown is that the model's condensate breaks them the way the Standard Model does. "
"Why the condensate carries Y = ½ — the assignment that leaves electromagnetism unbroken, fixed together with the "
"fermion charges by anomaly cancellation — is the next question, not settled here.")

heading("8.42  Anomaly cancellation fixes the hypercharges: the Standard Model's charges, derived", 2)
body("Sections 8.14 and 8.41 realized the Standard Model's gauge dynamics and its electroweak breaking, but took the "
"fermion hypercharges — and so the electric charges — as inputs. This section derives them. The lever is anomaly "
"cancellation: a chiral gauge symmetry is quantum-mechanically consistent only if its anomalies cancel, and the "
"model's emergent gauge symmetries are exact, realized as Wilson-link lattice gauge invariances (Section 8.14) with an "
"exact photon Ward identity (Section 8.16's regulator). An exact chiral gauge symmetry cannot be anomalous — one "
"cannot keep gauge invariance exact on a lattice while carrying a gauge anomaly — so here anomaly cancellation is not "
"optional but forced, and it constrains the hypercharges. Given only the observed representations of one generation — "
"the quark doublet (3, 2), the singlets (3̄, 1) and (3̄, 1), the lepton doublet (1, 2) and the singlet (1, 1), with "
"the five hypercharges left unknown — the four conditions [SU(3)²U(1)], [SU(2)²U(1)], [grav²U(1)] and the cubic "
"[U(1)³] apply.")
result("Result 8.42 — the hypercharges are fixed to the Standard-Model values; charge quantisation is derived.", "The "
"three linear conditions leave a two-parameter family (Y₄ = −3Y₁, Y₅ = 6Y₁, Y₂ + Y₃ = −2Y₁), and the cubic collapses "
"it: with r = Y₂/Y₁ it reads r² + 2r − 8 = 0, so r = −4 or r = +2, the two branches being the same theory with the two "
"quark singlets relabelled. Up to that relabelling and one overall scale, the hypercharges are unique — the "
"Standard-Model assignment 1/6, −2/3, 1/3, −1/2, 1 (each anomaly coefficient machine zero) — derived rather than "
"input. The electric charges Q = T₃ + Y then come out quantised and Standard-Model-valued: quarks at +2/3 and −1/3, "
"leptons at 0 and −1, so the proton (uud) and electron charges cancel to 10⁻¹⁶. Charge quantisation — the exact "
"equality of the proton and electron charge — is thus a consistency requirement, not a coincidence: it is the unique "
"anomaly-free assignment, forced by the model's exact gauge invariance. This is the sharpest Standard-Model prediction "
"in the report. Its scope is honest: anomaly cancellation fixes the hypercharges given the representations and the "
"one-generation content, which remain inputs, as do the gauge group and the number of generations; the one continuous "
"freedom, the U(1) hypercharge, is what is newly forced. And it presumes a genuine exact 4D gauge symmetry — the "
"domain-wall construction of Section 8.15 can park an anomaly in the bulk, so a truly four-dimensional symmetry with "
"no accessible bulk is where the constraint bites.")

heading("8.43  Where the gauge group and the generations come from: discrete band topology", 2)
body("Section 8.42 fixed the hypercharges but left the deepest inputs open: the gauge group SU(3)×SU(2)×U(1), the "
"fermion representations, and the number three of generations. Deriving those specific values from a specific lattice "
"is a genuine open problem and is not attempted here. What this section establishes is the category of the answer, "
"which reframes the question. In an emergent-fermion medium these are not continuous parameters to be fine-tuned but "
"discrete data — the symmetry and topology of the band structure — so they are quantised by their nature. Two pieces "
"make this concrete. First, the low-energy fermions live at the isolated Dirac points of the lattice, whose number is "
"fixed by fermion doubling and not tunable: the honeycomb has exactly two (the valleys K and K', verified as the two "
"inequivalent zeros of the structure factor), so its low-energy theory carries a two-fold flavour multiplet, and the "
"symmetry rotating that degenerate multiplet — gauged by the fermion loop of Section 8.14 — is the emergent gauge "
"group. A richer lattice gives more Dirac points and a larger group; the group is the symmetry of the multiplet and "
"the multiplet size is a discrete lattice output. Second, the number of chiral generations is a topological index: "
"the Chern number of a Wilson-Dirac band, computed from its Berry curvature, is an integer (0, ±1 across the mass "
"regimes here) that jumps only when the gap closes, and by bulk-boundary correspondence (Section 8.15) the number of "
"chiral families on a domain wall equals the jump in that integer across it.")
result("Result 8.43 — the gauge group and generation count are discrete band-structure data, not continuous tunings.", "The "
"honeycomb's two inequivalent Dirac points give a two-fold flavour multiplet whose rotation symmetry is the seed of a "
"non-Abelian gauge group; the Wilson-Dirac Chern number comes out 0, −1, +1, 0 across the mass regimes, so the number "
"of chiral generations on a domain wall is a quantised topological integer, changing only at a gap closing. This "
"answers 'why a specific compact group and a small integer of generations' — because these are the symmetry and "
"topology of the emergent fermion band structure, discrete and robust, rather than arbitrary continuous inputs. It "
"does NOT derive SU(3)×SU(2)×U(1), the representations, or the number three: those are properties of the physical "
"medium's actual lattice, which is not fixed here. But taken with Section 8.42, the character of the remaining input "
"has changed — with the hypercharges forced by anomaly cancellation and the group and generation count established as "
"discrete band data, what is left to specify is which single lattice the medium realises, a discrete structural "
"question rather than a list of continuous fine-tunings.")

heading("8.44  Which lattice: the fermion lattice requires angular rigidity", 2)
body("Section 8.43 reduced the remaining Standard-Model input to a single discrete choice — the lattice the medium "
"realises — because the gauge group and the generation count are band-structure data of that lattice. This section "
"characterises that choice. The medium self-assembles under a central pair force (the Lennard-Jones interaction of the "
"integration sections), whose ground state is the close-packed triangular lattice in two dimensions and the "
"face-centred cubic lattice in three; its coordination is six and it carries no Dirac point, so a purely central medium "
"hosts bosons but no fermions. Whether it can instead select the bipartite honeycomb — the coordination-three lattice "
"whose two inequivalent Dirac points give the emergent fermions of Section 8.2 — is a question about the interaction, "
"and three measurements answer it. First, a zero-temperature lattice sum ranks the candidates: under Lennard-Jones the "
"triangular lattice is lowest and the honeycomb highest, and adding a second isotropic length scale (a repulsive "
"shoulder) can move the ground state to the square lattice but never to the honeycomb, across the whole scan — a "
"close-packed lattice escapes any single isotropic barrier by dilating until its second-neighbour shell clears it, so "
"no number of central length scales forces coordination three. Second, an angular three-body term of Stillinger-Weber "
"form λ Σ (cos θ + ½)², minimised at 120°, is identically zero on the honeycomb (all bonds meet at 120°) and large on "
"the triangular lattice (60° bonds); as λ grows the ground state passes from triangular through square to honeycomb, "
"the honeycomb winning above λ* ≈ 0.57 of the nearest-neighbour bond energy, an order-unity amount of directional "
"bonding. Third, the same force — analytic, checked against finite differences to one part in 10^10 — is annealed "
"dynamically: a heated honeycomb seed is preserved under the angular force (coordination three, angles 120°) but "
"densifies to coordination six under the central force alone, and a disordered droplet close-packs under the central "
"force while forming honeycomb domains under the angular one. The collapse has a simple cause: coordination three is "
"below the two-dimensional Maxwell rigidity threshold of four, so the honeycomb is mechanically floppy under central "
"forces and rigid only once angular bonds are added. The bond-orientational order ψ₆ is near one for both the "
"honeycomb and the triangular lattice, so coordination, not ψ₆, is the order parameter that separates them.")
result("Result 8.44 — the fermion lattice requires order-unity angular rigidity, a property the medium's bonding either "
"has or lacks.", "A central pair potential selects the close-packed triangular or face-centred-cubic lattice "
"(coordination six, no Dirac point) and, with any number of isotropic length scales, at most the square lattice; the "
"bipartite honeycomb that carries the emergent fermions is never the ground state of a central interaction and is "
"mechanically floppy in one, its coordination three lying below the Maxwell threshold four. A 120° three-body term "
"makes the honeycomb both the energetic ground state, above λ* ≈ 0.57 of the bond energy, and the dynamically "
"self-assembled and stable lattice (the force gated against finite differences to 10^-10; a heated seed held at "
"coordination three, against densification to six without it). So the discrete lattice choice that Section 8.43 "
"isolated as the last Standard-Model input maps onto one physical property of the medium: whether its bonding carries "
"order-unity angular rigidity. This is the same non-central stiffness the emergent-Lorentz construction already "
"required (Section 8.1, where bare central forces give a longitudinal speed above the transverse and only shear-bearing "
"bonds give a single round cone), so two independent parts of the program demand the same departure from a purely "
"central medium. It does not derive which lattice the medium realises; it converts that input from an unstructured "
"choice into a specific question about the interaction.")

heading("8.45  One stiffness, three jobs: the angular sector selects the lattice and carries the cone", 2)
body("Section 8.44 found that the fermion lattice needs an order-unity angular stiffness, and Section 8.1 found that a "
"purely central medium has two acoustic cones (a longitudinal speed above the transverse) and needs a non-central "
"stiffness to give one. This section asks whether those are the same stiffness by using the same force model — "
"Lennard-Jones plus λ Σ (cos θ + ½)² — and reading the medium's acoustic cone off its elastic constants as a function "
"of the same λ. The elastic constants are computed by homogeneous strain of a periodic cell with the atoms relaxed "
"internally, giving the two-dimensional bulk modulus K from a dilation and the shear modulus μ from a shear, whence "
"c_L = √((K+μ)/ρ) and c_T = √(μ/ρ). Two controls fix the method: the triangular lattice under the central force alone "
"returns c_L/c_T = √3 and Poisson ratio 1/3 (the central-force Cauchy solid, matching the medium's measured 1.7330 and "
"0.330), and the honeycomb under the central force alone returns μ = 0 — it is floppy, its coordination three lying "
"below the two-dimensional Maxwell threshold of four, so a central-force honeycomb has no transverse cone at all. The "
"honeycomb cone is elastically isotropic (the shear modulus is the same in two orientations to one part in 10^7), so "
"the cone is well defined. Sweeping λ then separates the two sectors cleanly: the bulk modulus K is independent of λ to "
"machine precision (compression is central), while the shear modulus μ — and therefore the entire transverse cone c_T "
"— is sourced entirely by the angular term, rising from zero at λ = 0 roughly linearly, so c_T grows as √λ. At the "
"selection threshold λ* ≈ 0.57 the honeycomb already carries an order-unity transverse cone (c_T/c_L ≈ 0.33). A fully "
"round cone c_L = c_T is the μ ≫ K limit, reached as λ grows past λ*, which is the equal-stretch-and-bend (vector-Hooke) "
"point of Section 8.1.")
result("Result 8.45 — one order-unity angular stiffness does the work of three: it selects the honeycomb, makes it "
"mechanically stable, and is the sole source of its transverse cone.", "In the same force model as Section 8.44, a "
"purely central medium (λ = 0) leaves the honeycomb floppy — zero shear modulus, no transverse cone, not the ground "
"state — while the triangular control gives the central-force values c_L/c_T = √3 and ν = 1/3. Turning on an angular "
"stiffness of order the bond energy does three things at once, none of which a central force does: it selects the "
"honeycomb (λ* ≈ 0.57, Section 8.44), it rigidifies it, and it supplies its entire transverse acoustic cone. The split "
"is exact: the bulk modulus is λ-independent to machine precision (compression is central), and the shear modulus, "
"hence c_T, is sourced wholly by the angular term (zero at λ = 0, order unity by λ*). Full cone universality c_L = c_T "
"is the same knob turned further, into the μ ≫ K regime that reproduces Section 8.1's single round cone. So the "
"non-central stiffness that Section 8.44 identified as the requirement for emergent fermions is the same one Section "
"8.1 required for the medium's single (phonon-sector) cone; the two independent barriers turn on the same property of "
"the medium, and the mechanical sector — lattice selection, stability, and the acoustic cone — lives in the angular "
"stiffness while compression stays central. The scope is the mechanical cone: the fermion cone v_F is set by the "
"electronic hopping (v_F = (3/2) t a) and is an independent scale, so locking it to the boson cone — cross-statistics "
"universality — remains the composite-boson question of Section 8.1, which this stiffness does not by itself settle.")

heading("8.46  Two cones, one of them physical: the mechanical sound does not spoil the fermion cone", 2)
body("Section 8.45 gave the honeycomb a mechanical acoustic cone c_T. The matter that lives on the honeycomb, though, "
"are the Dirac fermions of Section 8.2, whose cone v_F = (3/2) t a is set by the electronic hopping t — a different "
"scale. This asks whether the two are one cone and, if not, whether the mismatch spoils the emergent Lorentz "
"invariance. Three measurements answer it. First, the cones answer to independent knobs: the hopping t moves v_F with "
"the mechanics fixed, while the angular stiffness λ moves c_T with the hopping fixed, so they coincide only at a tuned "
"hopping (t ≈ 1.11 here) that no symmetry enforces — the cross-statistics two-cone problem of Section 8.5, now between "
"the mechanical phonon and the fermion. Second, the physical emergent boson rides v_F, not c_T: the lower edge of the "
"interband particle-hole continuum, ω_min(q) = min_k t(|f(k+q)| + |f(k)|), goes to v_F |q| as q → 0, so the composite "
"photon and graviton of Sections 8.5 and 8.7 are on the fermion cone. Third, and decisively, the mismatch is benign. A "
"phonon is a bond-length modulation — a strain-dependent hopping, an off-diagonal Bloch term — so it preserves the "
"sublattice (chiral) symmetry and cannot make a mass; it only moves the Dirac point, a pseudo-gauge field, keeping the "
"cone gapless. At realistic (small) amplitude every phonon pattern leaves the gap at zero, and only an order-unity "
"distortion annihilates the cone in a Lifshitz transition (a uniform stretch never does; a shear doublet at u ≈ 1, a "
"single bond at u ≈ 1.5). The contrast is an on-site sublattice mass, which gaps the cone at once — exactly the term a "
"bond phonon is forbidden from producing.")
result("Result 8.46 — the medium has two cones, but only the fermionic one is physical, and the mechanical one cannot "
"spoil it.", "The mechanical acoustic cone c_T (Section 8.45) and the fermion cone v_F (the electronic hopping) are "
"independent knobs that coincide only by tuning. But the emergent relativistic matter, and the composite photon and "
"graviton, ride v_F — the interband particle-hole collective mode goes to v_F |q| — and a phonon, being an off-diagonal "
"bond perturbation, acts on the fermions only as a pseudo-gauge field: it moves the Dirac point without gapping it, and "
"cannot make the on-site mass that would. The cone's survival is topological, holding until an order-unity Lifshitz "
"distortion. So the two-cone mismatch is harmless: the fermion cone, which is the emergent Lorentz cone, is insulated "
"from the medium's own sound, and c_T is a decoupled sub-quantum spectator — the Volovik picture, in which the "
"underlying atoms' sound and the emergent light live on different cones without conflict. The photon and graviton "
"kinetic terms are induced by the fermion loop (Sections 8.5, 8.7), so they inherit v_F rather than the bare "
"elasticity. Scope: the protection is shown for a static (frozen) phonon acting as a gauge field; the dynamical phonon "
"self-energy on the cone is a further computation, but the cone's survival is topological rather than perturbative.")

heading("8.47  The Dirac cone survives phonon fluctuations to all orders: chiral protection", 2)
body("Section 8.46 showed a static phonon acts as a pseudo-gauge field and cannot gap the Dirac cone. This closes the "
"dynamical version, and by a symmetry that holds to all orders. A phonon is a bond-length modulation — an off-diagonal "
"hopping term — and the honeycomb Dirac Hamiltonian carries the sublattice (chiral) symmetry S H S = −H with S the "
"sublattice sign. A bond term is chiral-odd (it transforms like H), so it cannot generate the chiral-even mass that "
"would open a gap, at any order, static or dynamical; the only gapping term is an on-site sublattice mass, which a bond "
"phonon does not produce. The parities are exact numerically (the bond perturbation gives S δH S + δH = 0, the "
"staggered mass gives S M S − M = 0, both to machine precision), and the consequence is verified on the spectrum: "
"random bond disorder — a frozen-phonon ensemble, the fluctuation content a dynamical phonon supplies — keeps the "
"honeycomb gapless at zero energy up to order-unity amplitude, while a staggered mass of 0.1 opens a hard gap. So the "
"fermion cone survives phonon fluctuations, not by cancellation but by an exact symmetry.")
body("What this does not settle is whether the two cones actually merge. The mechanical cone c_T and the fermion cone "
"v_F coexist without conflict, but whether the velocity mismatch flows to zero in the infrared — v_F = c_T, a single "
"cone rather than mere coexistence — is a two-velocity renormalisation-group question, and that calculation is "
"regulator-limited. With two velocities there is no single Lorentz-covariant cutoff, so a hard momentum/frequency "
"cutoff breaks the very symmetry it is meant to measure: the one-loop flow fails its own Lorentz gate — the anomalous "
"velocity dimensions γ_v (fermion self-energy) and γ_c (boson polarisation), which must be equal at v = c, differ by an "
"amount comparable to the signal, with a large non-logarithmic residual — exactly as the induced-graviton Ward identity "
"of Section 8.29 failed under a hard cutoff. A symmetry-preserving scheme (dimensional regularisation, or the lattice "
"Brillouin-zone regulator that rescued the lattice Ward identity) is needed to settle whether the cones merge. What is "
"robust is the gaplessness, which rests on the exact chiral symmetry, not on a loop.")
result("Result 8.47 — the Dirac cone survives phonon fluctuations to all orders, by chiral symmetry; whether the two "
"cones merge is regulator-limited.", "A phonon is a bond (off-diagonal) term, chiral-odd under the sublattice symmetry "
"S H S = −H, so it cannot generate the chiral-even mass that gaps the cone — at any order. Verified: the bond "
"perturbation and the staggered mass carry exactly opposite chiral parity (machine zero), and random bond disorder "
"keeps the spectrum gapless up to order-unity amplitude while a staggered mass gaps it at once. This closes the "
"dynamical form of the two-cone seam — the mechanical sound cannot spoil the fermion cone, statically or dynamically, "
"because the protection is an exact symmetry rather than a cancellation. The remaining refinement — whether v_F and c_T "
"merge into one velocity in the infrared — is a two-velocity RG that is regulator-limited: no Lorentz-covariant cutoff "
"exists for two velocities, so the one-loop flow fails its own γ_v = γ_c gate at v = c, as in Section 8.29, and a "
"symmetry-preserving scheme is required to settle it. The two-cone mismatch is thus established as harmless; whether it "
"vanishes is open.")

heading("8.48  The generation number as a measured band-structure invariant", 2)
body("Section 8.43 reframed the number of fermion families as a Chern index rather than a continuous parameter, but it "
"only computed a Chern number of one, for a Wilson band, and it asserted rather than measured the bulk-boundary count "
"that ties that index to a family of chiral fermions. This section makes the statement a measurement, on the medium's "
"own two-band structure. The emergent low-energy fermion has two sublattices, H(k) = d(k)·σ, and the generation content "
"is the Chern number of the occupied band — which is exactly the number of times the inter-sublattice coupling winds "
"around the Brillouin zone. Taking a coupling that winds n times, (sin k_x + i sin k_y)^n, gapped by a Wilson mass, the "
"family number is read out two independent ways that must agree. In the bulk, the Fukui–Hatsugai–Suzuki lattice flux — "
"a gauge-invariant integer on any grid — gives Chern number C = −n for n = 1, 2, 3. On the boundary, a cylinder "
"(periodic in one direction, open in the other) carries exactly n in-gap chiral branches crossing zero energy on each "
"edge, counted directly from the ribbon spectrum. Bulk and boundary agree: the generation number is a measured integer, "
"the winding of one structural function, quantized by topology rather than tuned.")
body("The two edges carry their n branches with opposite chirality, so the closed lattice nets to zero — the "
"Nielsen–Ninomiya theorem, here measured (the fitted branch slopes give +n on one edge and −n on the other). This has a "
"concrete structural consequence: n net-chiral generations cannot be a property of a bulk two-dimensional medium, they "
"must live on a defect — an edge or domain wall — with the opposite chirality residing elsewhere as a compensating "
"mirror sector, consistent with the domain-wall mechanism by which the model produces a chiral fermion at all. What this "
"does not do is derive the number three. The winding is a free integer; nothing in the medium's currently-known "
"structure pins it. The result converts 'why three generations' into 'the inter-sublattice coupling winds three times' — "
"a measured geometric restatement of the same input, with the count now a band-structure invariant rather than an "
"arbitrary continuous parameter.")
body("The winding also reconciles the two mechanisms that Section 8.43 had left as separate statements. A winding-n "
"point is a multi-Weyl point — the inter-sublattice coupling has a degree-n zero — and such points are protected only by "
"crystalline symmetry, hence fine-tuned. A generic perturbation (a uniform inter-sublattice coupling suffices) splits "
"the degree-n zero into exactly n unit-winding Dirac points, with the total winding conserved at n; this is measured by "
"the contour winding of the coupling, which reads n around the multi-Weyl point and one around each fragment, summing to "
"n. The robust content of 'n generations' is therefore n ordinary Dirac points — a count of fermion doublers — rather "
"than a fragile n-fold degeneracy, and the single multi-Weyl point is only the symmetric limit. That identifies Section "
"8.43's flavour count (the number of Dirac points) and its generation index (the Chern number) as one and the same "
"conserved integer: the Chern number is the summed chirality, i.e. the summed winding, of the Dirac points, so counting "
"them and computing the topological index return the same number.")
result("Result 8.48 — the generation number is a measured topological invariant of the emergent band, not a tuned "
"parameter.", "The number of fermion families equals the Chern number of the two-band medium fermion — the winding of "
"its inter-sublattice coupling around the Brillouin zone — and is read out two agreeing ways for winding n = 1, 2, 3: "
"the bulk Fukui–Hatsugai–Suzuki lattice flux gives C = −n exactly, and a cylinder carries exactly n chiral edge "
"branches, measured from the ribbon spectrum rather than asserted (upgrading §8.43). The two edges carry opposite "
"chirality (fitted slopes ±n), so by Nielsen–Ninomiya the closed lattice nets to zero: the n net-chiral generations "
"must live on a defect (edge/domain wall) with a compensating mirror sector. The winding-n point is a fine-tuned "
"multi-Weyl degeneracy that a generic perturbation fragments into n unit-winding Dirac points (total winding conserved), "
"so the robust family number is a count of ordinary Dirac points — which unifies §8.43's flavour count and Chern index "
"as one conserved integer. Honest scope: this does not derive three — "
"the winding is a free integer, so 'why three' becomes the measured, geometric statement 'the coupling winds three "
"times', the family count now a quantized band invariant rather than a continuous input.")
body("A construction that measures a winding it was built to have invites an authenticity check: does the topological "
"machinery compute real invariants, or does it reflect the input back? The guard is to run the identical code on "
"independent models it was not built for. The same Fukui–Hatsugai–Suzuki flux reproduces the Haldane model's analytic "
"topological phase boundary |M| = 3√3 t₂ sin φ (the number 3√3 = 5.196, recovered for three values of φ) with its ±1 "
"Chern lobes, and the same honeycomb tight-binding reproduces graphene's bandwidth 6t and Fermi velocity v_F = (3/2) t a "
"to five digits — all un-tuned. So the §8.48 integers are the code computing real topology, not a construction "
"reflecting itself. This validates the tooling's correctness, and only that; it is not a claim of contact with "
"experiment.")

heading("8.49  The last assumption behind γ = 0: does the model's mass nucleate net disclination charge?", 2)
body("The gravitational arc closed every channel to γ = 1, but §8.33's topological channel did so in the "
"continuum: it modelled the mass as a smooth compression eigenstrain θ*(x) ∼ ρ(x) and showed its net curvature "
"charge ∫η = ∫∇²θ* vanishes by Gauss–Bonnet. That argument silently assumes the model's mass is only a change of "
"bond length, never a change of bond connectivity. Net curvature charge is carried by exactly one object — a "
"topological disclination, a site whose coordination departs from six, a genuine deficit angle — and a disclination "
"is a change of connectivity, not of length. So one premise was never tested directly: whether concentrated energy, "
"placed in the actual discrete medium the way the model's gravity places it, NUCLEATES a disclination with net "
"charge proportional to the mass. That is the only microscopic route left to γ = 1, and this section measures it on "
"the lattice. The measurement is made clean by reading coordination through a cutoff that scales with the local "
"spacing, so it is covariant under compression: it counts a topological re-bonding, never a mere shortening of "
"bonds — the discrete separation of the length change §8.33 covers from the connectivity change it assumed away.")
body("Four gates settle it. A positive control confirms the measurement is not blind: inject genuine net "
"disclination density (a cone embedding — delete a sixty-degree wedge and stretch the remainder to fill the plane) "
"and the net interior charge reads a clear +45, with the charge enclosed by a loop growing with radius rather than "
"saturating — a real deficit-angle charge is unscreened and long-range, exactly what γ = 1 needs. The total "
"coordination charge over the disc is a topological invariant (Σ(6−z) = 150 = 6χ, the boundary), unchanged to the "
"integer by any amount of compression, so net charge cannot be sourced locally — an interior disclination must be "
"paid for elsewhere. Then the measurement itself: driving the medium with the model's mass, a compression well of "
"depth proportional to the mass against a clamped rim, the core spacing falls measurably (1.12 → 1.04, a real ten "
"per cent compression) while the net interior disclination charge stays exactly zero at every mass — indeed no "
"defect nucleates at all, not even a neutral dislocation pair; the response is purely elastic. A striking "
"corollary surfaced in the construction: a single isolated disclination cannot even be built in the flat bulk "
"without it shattering into a defect scar or demanding a compensating charge — the conservation law asserting "
"itself in the very attempt. The mass, measured directly, is the smooth compression eigenstrain of §8.33, and it "
"nucleates none of the charge that would bend light at range.")
result("Result 8.49 — the model's mass nucleates no net disclination charge; γ = 0 holds at the microscopic level "
"too.", "Reading coordination on the lattice with a compression-covariant cutoff separates a change of bond length "
"(compression, the model's gravity) from a change of bond connectivity (a disclination, the only carrier of net "
"curvature charge). The measurement sees genuine net charge when it is present — an injected disclination density "
"reads +45 and is long-range, its enclosed charge growing with radius (positive control) — and it is topologically "
"conserved (Σ(6−z) invariant under compression, Gauss–Bonnet). But the model's own mass, a compression well that "
"demonstrably compresses the medium, nucleates zero net disclination charge at every amplitude, and in the tested "
"range nucleates no defect at all: the response is purely elastic. This is §8.33's ∫η = 0 measured on the lattice "
"rather than assumed in the continuum — a seventh independent γ = 0, and the first at the microscopic / topological "
"level, closing the arc's last standing assumption that the model's mass is pure compression. Scope, stated because "
"the result is a negative: the medium is the 2D triangular proxy where the disclination picture of curvature is "
"defined (as in §8.33, test_disclination_force, test_fracton_gravity); a hypothetical gravity mechanism that "
"changed the medium's connectivity rather than its bond lengths is outside the model's compression mechanism as "
"defined, and is the one thing this measurement cannot exclude by construction. Within the mechanism the model "
"actually has, γ = 1 has now been sought and found absent in every channel, continuum and discrete.")

heading("8.50  The measurement problem, as far as physics reaches: einselection by the medium", 2)
body("Section 8.6 took quantum mechanics apart and left one residue it would not claim: the measurement problem's "
"hard core, the selection of a single definite outcome, kept as a postulate. That residue has structure worth "
"resolving, because the measurement problem is really two questions and only one of them is metaphysics. The first "
"— why measurements have a preferred basis (position, not its superpositions) and why the interference between "
"branches vanishes — is ordinary physics: decoherence, the environment monitoring the system, einselecting the "
"basis that survives. The second — why one branch is realized rather than all of them — is the genuine hard core. "
"This section measures exactly how far the first reaches, with the environment being nothing added: the condensate's "
"own phonons.")
body("The bath is the medium's measured spectrum. Relaxing the node medium and diagonalizing its dynamical matrix "
"(the Hessian of the Lennard-Jones energy) gives the condensate's phonon modes; a which-path superposition of a "
"system particle at two positions couples to the local displacement of the medium at the particle's site — the same "
"'the medium responds to where the energy is' coupling that produces compression-gravity. The independent-boson "
"(pure-dephasing) evolution is then exact, and four things follow. The off-diagonal coherence decays to zero while "
"the populations stay frozen to machine precision — a superposition becomes a mixture (a coherence of 1.000 falls "
"to 0.01 while the branch weights hold at 0.5). The basis that goes diagonal is position, the operator the bath "
"couples to: a position eigenstate has nothing to dephase and is robust (its decoherence function is identically "
"zero), while its superpositions die — the medium einselects the pointer basis rather than having one assumed. The "
"rate scales as the square of the which-path separation — a fitted log-log slope of 2.0000 — so a mesoscopic "
"superposition decoheres astronomically faster than a microscopic one, which is the quantitative origin of the "
"quantum-classical boundary. And the surviving mixture carries the Born weights on its diagonal, unchanged.")
result("Result 8.50 — the medium einselects pointer states and Born-weighted branches; the single outcome stays the "
"only postulate.", "Decoherence is the part of the measurement problem that is physics, and the model supplies it "
"from its own structure: coupling a which-path superposition to the condensate's measured phonon bath — through the "
"same local coupling that makes compression-gravity — destroys the coherence between branches (1.000 → 0.01) while "
"preserving their populations, einselects the position basis the bath monitors (a position eigenstate does not "
"dephase; its superpositions do), and does so at a rate ∝ (separation)² (log-log slope 2.0000) that makes anything "
"mesoscopic instantly classical. The classical mixture left behind carries the Born weights of §8.6 on its diagonal. "
"The honest boundary, reached now from the other side: this is an improper mixture — the reduced system purity falls "
"to Σ|c_i|² -weighted Tr(ρ²) = 0.50, the equal-branch floor, while the global system-plus-bath state stays pure "
"(the evolution is unitary), so no single branch is selected. Einselection and Born weights are mechanism; which "
"outcome occurs is untouched, exactly the hard core §8.6 marked — now shown to be genuinely all that remains, not a "
"placeholder for missing physics. Scope: linear (pure-dephasing) system-bath coupling, for which the model is exact; "
"the coupling strength sets only the decoherence timescale, not any of the four structural results.")

heading("8.51  Sharpening the one prediction: how robust is the Lorentz-violation coefficient?", 2)
body("The consolidation of Sections 8.11–8.16 left the model with a single live falsifiable prediction — the "
"quadratic, subluminal, species-universal Lorentz violation of Sections 8.8 and 8.39, with an effective scale "
"E_QG,2 = E_Planck / √ζ set by the boost coefficient ζ ≈ 0.245 read off the emergent dispersion. That coefficient "
"came from one microscopic choice, an fcc nearest-neighbour graph, and a prediction that hinges on an arbitrary "
"choice is a fit. So the honest question is which parts of the prediction are firm and which are soft: does ζ "
"survive changing the lattice, the neighbour range, the direction? Extracting ζ the same way across simple-cubic, "
"body-centred and face-centred lattices settles it, and the answer is sharper than expected. The coefficient is "
"essentially lattice-independent — ζ = 0.250, 0.245, 0.243 for sc, bcc, fcc — so E_QG,2 sits at 2.0–2.03 times the "
"Planck energy regardless of the microscopic lattice, not a fitted number but a structural one. Widening the "
"coupling from nearest-neighbour to the second shell moves it more (ζ: 0.245 → 0.37, still subluminal, E_QG within "
"~20 per cent), so the neighbour range is the largest sensitivity and even it is mild. The violation is subluminal "
"(ζ > 0) on every lattice tested, quadratic by the structure of the k⁴ term, and the crystallographic anisotropy is "
"a subleading fraction of ζ (≈ 0.28 on fcc) that averages down in a poly-domained medium.")
result("Result 8.51 — the prediction's falsifiable content is robust; only its coefficient's last digit and its "
"anisotropy pattern are soft.", "Stress-testing the Lorentz-violation coefficient across microscopic choices "
"separates the prediction into firm and soft. Firm, and lattice-independent: the violation is quadratic (n = 2, not "
"the linear form of many quantum-gravity scenarios), subluminal (ζ > 0 on simple-cubic, bcc and fcc alike), "
"species-universal (one cone, Section 8.5), and set at an effective scale of a few times the Planck energy — "
"E_QG,2 = 2.4–2.5 × 10^19 GeV across lattices, moving by only ~20 per cent even when the coupling range is doubled. "
"Soft, and model-dependent: the exact coefficient 0.245 and the crystallographic anisotropy pattern, which shift "
"with the lattice and wash out under domain averaging. This matters for falsifiability, because it is exactly the "
"firm part that nature would test — n = 2 versus n = 1, subluminal versus superluminal, one cone versus a "
"species-dependent speed, and a near-Planckian scale — while the fit-like number carries no weight. The one "
"prediction the model has left is therefore a structural claim, not a tuned one, which is the strongest form an "
"unfalsified prediction can take.")

heading("8.52  The first real-data confrontation: GW170817 and the one universal cone", 2)
body("Every confrontation with experiment to this point has either reproduced a textbook number (the graphene and "
"Haldane anchors of Section 8.48) or cited published bounds (Sections 8.16, 8.39). This section takes an actual "
"measurement and confronts the model's single firm, structural prediction against it: one universal light cone, so "
"photons and gravitons travel at the same speed at leading order (Sections 8.4–8.5). The measurement is the "
"multi-messenger detection of the binary-neutron-star merger GW170817 and its gamma-ray burst GRB 170817A. The "
"burst arrived 1.74 ± 0.05 s after the gravitational-wave merger, from a source at a conservative distance of at "
"least 26 Mpc; assuming the gamma rays were emitted within the standard 0–10 s intrinsic window after merger, the "
"observed timing constrains the fractional difference in propagation speed to lie between −3×10⁻¹⁵ and +7×10⁻¹⁶ of c.")
body("Reconstructing that bound from the raw timing reproduces the published interval to the figures given "
"([−3.1×10⁻¹⁵, +6.5×10⁻¹⁶]), which validates the confrontation. The model's leading-order prediction is that "
"photons and gravitons share one cone exactly, (v_gw − v_γ)/c = 0, and zero sits inside the measured interval — the "
"firm structural prediction is consistent with a real measurement, a genuine pass rather than a reproduction. The "
"model's own residual Lorentz violation is far below the reach of this event: its n = 2 dispersion gives "
"v(E)/c − 1 = −ζ (E/E_Planck)², so at the burst's ~185 keV photons and ~100 Hz gravitons the predicted "
"photon–graviton speed difference is about 6×10⁻⁴⁷, some thirty-one orders of magnitude inside the bound. GW170817 "
"therefore tests the universality half of the prediction — one cone, no species-dependent speed — and passes it, "
"while the dispersion coefficient stays untouched, because the quadratic suppression makes the effect negligible "
"until the ultra-high-energy frontier (Section 8.39). The falsifiers this event can bear on are all currently "
"passed: no leading-order species-dependent speed (predicted zero, bounded below 10⁻¹⁵), a subluminal rather than "
"superluminal signal, and a quadratic rather than linear energy dependence.")
result("Result 8.52 — the one-cone prediction passes its first confrontation with real data.", "The model's single "
"firm falsifiable prediction — one universal cone, so light and gravity propagate at the same speed at leading "
"order — is confronted with the GW170817 / GRB 170817A timing and passes: the measured |v_gw − v_γ|/c, bounded "
"between −3×10⁻¹⁵ and +7×10⁻¹⁶ (reproduced here from the 1.74 s delay over ≥26 Mpc), contains the model's exact "
"zero. This is the suite's first contact with an actual measurement rather than a textbook value, and it moves the "
"program's weakest point — external validation — off zero. Its honest reach is narrow and stated plainly: the event "
"tests the universality of the cone, not the Lorentz-violation coefficient, whose Planck-suppressed n = 2 signal is "
"about 6×10⁻⁴⁷ at these energies — thirty-one orders below the bound. Confirming the coefficient, the part of the "
"prediction that is genuinely new physics, requires the ultra-high-energy frontier — and Section 8.53 works through "
"it and finds, honestly, that even there the coefficient is out of reach, because the very properties that pass "
"GW170817 (subluminal, one cone) evade the bounds that would test it. What GW170817 establishes is that the firm, "
"structural half of the prediction survives contact with the sky.")

heading("8.53  Working through the ultra-high-energy frontier: the coefficient evades its own tests", 2)
body("GW170817 confronted the universality half of the prediction and passed. The coefficient half — the n = 2 "
"dispersion v(E)/c = 1 − ζ(E/E_Planck)², ζ ≈ 0.245 — needs the ultra-high-energy frontier, where a quadratic effect "
"finally grows, and this section works through the three datasets that define that frontier: the Pierre Auger "
"cosmic-ray spectrum and its GZK suppression at ~4×10¹⁹ eV; LHAASO's gamma rays up to 1.4 PeV, the highest photons "
"ever seen; and IceCube's ultra-high-energy neutrinos, including the 6.3 PeV Glashow-resonance event and the 290 TeV "
"neutrino coincident with a gamma-ray flare of the blazar TXS 0506+056. The confrontation is consistent with all "
"three, but the honest content is why, and it is not the reassuring answer.")
body("Two properties the model relies on to pass GW170817 are exactly what make it evade the strongest "
"ultra-high-energy bounds. First, it is subluminal (ζ > 0): the sharpest clean photon bounds at these energies are "
"photon decay, γ → e⁺e⁻, and vacuum Cherenkov, which are kinematically allowed only for superluminal photons, so a "
"subluminal photon does not decay and LHAASO's headline Lorentz-violation limits — superluminal-decay limits — do "
"not constrain the model at all. Second, it is one cone: in the Coleman–Glashow / Jacobson–Liberati–Mattingly "
"analysis, threshold reactions such as the GZK photopion process are anomalous only through the difference of the "
"species' Lorentz-violating coefficients, and for a single universal cone that difference is zero, so the leading "
"threshold shift cancels and the species-dependent GZK bound (E_QG,2 ≳ 10¹⁸ GeV) does not apply as stated. What the "
"frontier does test is universality again, now at high energy: the TXS 0506+056 neutrino–gamma coincidence bounds "
"the neutrino–photon speed difference to about 10⁻¹² at 290 TeV, and the model predicts about 10⁻²⁸ — a one-cone "
"pass at 10⁵ GeV, complementing GW170817 at the MeV scale. Its own dispersion at every frontier energy is minute "
"(~10⁻¹⁷ for the highest cosmic rays, ~10⁻²⁷ for PeV photons).")
result("Result 8.53 — the frontier tests the model's universality, not its coefficient, and tempers §8.39.", "Worked "
"through against real data: the model is consistent with the Auger GZK suppression, the LHAASO PeV gamma rays, and "
"the IceCube ultra-high-energy neutrinos. But the reason is a genuine and slightly sobering one. The two features "
"that carried the GW170817 pass — a subluminal sign and one universal cone — are precisely what disarm the frontier "
"tests: photon-decay bounds require superluminal propagation, which the model does not have, and the species-"
"dependent threshold (GZK) bounds require a difference between species' cones, which one cone sets to zero. So the "
"coefficient ζ — the genuinely-new-physics part of the prediction — is not reached by any current UHE dataset; the "
"frontier tests the universality (which passes again, now at 10⁵ GeV via TXS 0506+056, as well as at the MeV scale "
"via GW170817). This tempers §8.39: its reading that the model sits ~1.4 orders below the frontier and is reachable "
"by next-generation observatories compared E_QG against species-dependent bounds that a universal, subluminal model "
"evades. The honest status is that the model's one live prediction is confirmed in its structural half (one cone, "
"twice) and, in its quantitative half (the coefficient), currently beyond reach — for a reason internal to the "
"prediction itself, not merely a matter of sensitivity. A dedicated universal-LV threshold analysis, or a genuinely "
"new observable, would be needed to reach ζ. Correction: one gate of this section was wrong, and Section 8.54 does "
"exactly that threshold analysis and reverses the GZK conclusion — the n = 2 threshold, unlike a velocity shift, "
"does not cancel under universality, so the coefficient is not evaded but is in fact in ~1 order-of-magnitude tension "
"with the observed GZK cutoff.")

heading("8.54  Correcting §8.53: the GZK threshold does not cancel, and the coefficient is in tension", 2)
body("Section 8.53 concluded that the model's one-cone universality lets it evade the GZK bound, by a Coleman–Glashow "
"cancellation. That was an error, and it is worth correcting precisely because it reverses the most consequential "
"line in the whole external-validation arc. Coleman–Glashow cancellation is an n = 0 statement: a universal maximal "
"velocity, shared by all species, is unobservable because it is a global rescaling of the speed of light. The "
"model's Lorentz violation is n = 2 — a momentum-dependent p⁴ dispersion — and a universal n = 2 coefficient is not "
"a rescaling and does not cancel in thresholds. Vacuum Čerenkov and photon decay do still cancel, because those are "
"genuine relative-speed effects: a proton cannot out-run a photon with which it shares a cone, so those bounds "
"(including LHAASO's) are correctly evaded. But the GZK photopion threshold is not a relative-speed effect, and it "
"does not cancel.")
body("Done properly, the threshold condition for p + γ → p + π with a universal dispersion E = p + m²/2p − ξp³/2M² "
"carries a Lorentz-violating term (3/2)ξ x(1−x) p³/M², where x is the final proton's momentum fraction. Keeping the "
"proton and pion coefficients separate, this term is ξ_p(1−x³) − ξ_π(1−x)³, and at the GZK kinematic optimum "
"x ≈ 0.87 the proton piece (1 − x³ ≈ 0.34) dwarfs the pion piece ((1−x)³ ≈ 0.002): the shift is set by the absolute "
"proton coefficient and survives universality, verified numerically (the universal and proton-only threshold "
"obstructions agree to a per cent, while the pion-only one is negligible). The sign matters and is fixed: the model "
"is subluminal (ξ > 0), which makes the term positive, raises the photon energy the reaction needs, and suppresses "
"GZK — protons become stable and the cutoff is erased. At the model's coefficient ξ ≈ ζ/2 ≈ 0.12 the effect is not "
"subtle: the reaction would require cosmic-microwave-background photons about thirty times hotter than they are, so "
"it never proceeds, and the model predicts no GZK cutoff at all — in direct conflict with the suppression that Auger "
"and Telescope Array observe near 4×10¹⁹ eV. The same statement in the standard parametrization: the model's "
"|η₄| ≈ 0.12 exceeds the literature dimension-six proton bound |η₄| ≲ 10⁻² (Jacobson–Liberati–Mattingly) by about "
"an order of magnitude.")
result("Result 8.54 — the model's one prediction is in ~1-order tension with the GZK cutoff, and on the edge of "
"falsification.", "This corrects the GZK gate of Section 8.53 and, with it, the external-validation verdict. The n = 2 "
"GZK threshold does not cancel under one-cone universality — that cancellation is the n = 0 (velocity) theorem, and "
"an n = 2 dispersion is not a velocity. A direct threshold calculation shows the model's universal, subluminal "
"coefficient would suppress the photopion reaction so strongly (requiring CMB photons ~30× hotter than they are) "
"that no GZK cutoff would form, contradicting the observed suppression; equivalently |η₄| ≈ 0.12 against the "
"dimension-six proton bound |η₄| ≲ 10⁻². So the model's single live prediction, comfortably safe against the "
"universality tests (GW170817, TXS 0506+056), is in roughly one order-of-magnitude tension with the ultra-high-"
"energy cosmic-ray spectrum. Whether that is a clean exclusion or a survivable tension turns on one number the model "
"does not compute: the composite proton inherits not the fundamental ξ but ξ times a parton sum Σ z_i³ ≈ 0.1 — of "
"the very same order as the tension. So the honest status is sharp and uncomfortable: the model's one prediction now "
"sits on the edge of falsification by existing data, decided by a QCD compositeness factor it has not supplied. This "
"is the program's first genuine tension with a real measurement — external validation cutting against the model — "
"and it is the most important item the report now carries. Section 8.55 then computes the compositeness factor "
"that §8.54 left open, and it resolves the tension to marginal: the model sits right at the GZK frontier, neither "
"cleanly excluded nor cleanly safe.")

heading("8.55  The deciding calculation: proton compositeness pulls the coefficient to the GZK frontier", 2)
body("Section 8.54's tension rested on one assumption it flagged and did not evaluate: that the composite proton "
"inherits the fundamental universal coefficient ξ. It does not, and computing the suppression is the single "
"calculation that decides whether the model's one prediction is excluded or survives. In the parton picture a "
"proton of momentum p is a set of partons carrying momentum fractions z_i (Σ z_i = 1), each with the universal "
"fundamental dispersion, so its energy is E_p = p + m_p²/2p − ξ(Σ z_i³) p³/2M² and the proton's effective "
"coefficient is ξ_eff = ξ·⟨Σ z_i³⟩. There is a clean identity for the suppression factor: with D(z) the parton "
"number density and P(z) = z D(z) the momentum density (∫P = 1, the momentum sum rule), ⟨Σ z_i³⟩ = ∫z³ D(z) dz = "
"∫z² P(z) dz = ⟨z²⟩_P, the mean of z² under the proton's own momentum distribution. Because the proton's momentum "
"is shared among many partons at low-to-moderate z, and z² weights them down, this is a small number.")
body("Evaluated on realistic parton distributions — valence carrying ~39 per cent of the momentum, gluons ~46, "
"sea ~15, with the standard shapes — the momentum sum rule checks (∫P = 1, momentum-weighted ⟨x⟩_P ≈ 0.2) and the "
"suppression factor comes out Σ z³ = ⟨x²⟩_P ≈ 0.065, valence- and gluon-dominated at moderate x with the soft sea "
"contributing under a tenth. This cuts the fundamental ξ ≈ 0.12 to an effective proton coefficient "
"ξ_eff ≈ 8×10⁻³ — right against the dimension-six GZK bound of ~10⁻². The ~12-fold bare tension of §8.54 is pulled "
"to about 0.8 of the bound, and a sweep over harder and softer parton shapes keeps ξ_eff between 0.6 and 0.9 of "
"the bound throughout. Compositeness does most of the work the tension needed, and lands the model not clear of the "
"constraint but exactly on it.")
result("Result 8.55 — compositeness resolves the tension to marginal; the prediction sits on the GZK frontier.", "The "
"deciding calculation is done. The composite proton does not inherit the fundamental coefficient: its effective "
"n = 2 coefficient is suppressed by Σ z³ = ⟨z²⟩ under the proton's momentum density, a small number (≈ 0.065 on "
"realistic parton distributions, and 0.05–0.07 across a shape sweep) because the momentum is spread over many "
"low-to-moderate-z partons and z² weights them down. That cuts the fundamental ξ ≈ 0.12 to ξ_eff ≈ 8×10⁻³, turning "
"§8.54's ~12× GZK tension into a ~0.8× one. So the resolution is neither exclusion nor safety but the frontier "
"itself: the model's one live prediction survives every test it has faced — the universality coincidences (GW170817, "
"TXS 0506+056) and now the GZK threshold — and is pinned to the edge of current ultra-high-energy sensitivity. "
"Whether it is finally inside or outside the bound is not resolvable at present precision; it turns on a "
"factor-of-two-level proton Lorentz-violation moment (a lattice or global-fit calculation) and a firm GZK "
"Lorentz-violation bound. This is the most falsifiable place a prediction can sit, and it makes the model's one "
"empirical claim a concrete target for the next round of ultra-high-energy cosmic-ray data. Section 8.56 puts real "
"error bars on the factor and confirms it: anchored to measured parton moments, ξ_eff spans roughly half to one and "
"a half times the bound, so the model straddles the frontier rather than clearing it.")

heading("8.56  Real error bars on the compositeness factor: the model straddles the GZK frontier", 2)
body("Section 8.55 computed the suppression factor Σ z³ = ⟨x²⟩_P with Beta-function toy parton distributions and got "
"a single value, ≈ 0.065, landing ξ_eff at about 0.8 of the GZK bound. A toy shape deserves real error bars, and "
"anchoring the factor to measured data both sharpens it and — honestly — widens it, because two systematics enter "
"that the toy hid. The parton momentum fractions are taken from the world global fits (gluon ~0.42, u+ū ~0.34, "
"d+d̄ ~0.19, s+s̄ ~0.035, summing to one by the momentum sum rule), and the hardness of the valence distribution — "
"which controls the moment — is pinned to the measured isovector second moment. There the data themselves disagree: "
"physical-point lattice QCD gives ⟨x²⟩_{u-d} = 0.083(14), while global fits give about 0.055, a roughly "
"two-fold tension that is itself part of the error bar, and one the model inherits directly because Σ z³ scales "
"with it.")
body("Carried through, parton distributions consistent with the global-fit moment give Σ z³ = ⟨x²⟩_P ≈ 0.042–0.064, "
"and scaling to the higher lattice moment pushes the top of the range toward ≈ 0.12, so ξ_eff spans about 0.5 to "
"1.5 times the GZK bound: the global-fit end sits just inside, the lattice end just outside. Real parton data, in "
"other words, confirm §8.55's picture with quantified uncertainty and decline to rescue the model — it straddles "
"the exclusion boundary rather than clearing it. One caveat runs the other way and is the single largest remaining "
"uncertainty: the moment Σ z³ = ∫x³ D(x,Q²) dx decreases as the scale Q² rises and the soft-parton content grows, "
"and the fundamental Lorentz-violating operator's natural scale is the deep ultraviolet, so the physically relevant "
"factor is at or below the ~2 GeV reference value used here — the reference-scale band is an upper estimate, "
"trending safer. Quantifying that shift is a real calculation, the operator's anomalous dimension, not done here.")
result("Result 8.56 — real parton moments put the model's coefficient at 0.5–1.5× the GZK bound; it straddles the "
"frontier.", "Replacing §8.55's toy parton distributions with measured inputs — the world momentum fractions and "
"the physical-point lattice second moment ⟨x²⟩_{u-d} = 0.083(14), against the global-fit ~0.055 — gives the "
"compositeness factor Σ z³ = ⟨x²⟩_P ≈ 0.04–0.06 on global-fit-consistent distributions, rising toward ≈ 0.12 if the "
"higher lattice moment is taken at face value. So the effective proton coefficient is ξ_eff ≈ 0.5–1.5 times the "
"dimension-six GZK bound: just inside on global fits, just outside on the lattice moment. The toy 0.8× of §8.55 "
"becomes a data-anchored band that straddles the exclusion boundary, and the real moments do not move the model "
"clear of it. Two numbers would decide it — the lattice-versus-global-fit moment tension, and a firm GZK "
"Lorentz-violation bound — with a third, the Lorentz-violating operator's ultraviolet scale, trending the answer "
"safer but uncomputed. The bottom line is unchanged and now carries error bars: the model's one empirical claim is "
"pinned to the ultra-high-energy frontier, the program's most falsifiable and most exposed result, and a concrete "
"target for the next round of cosmic-ray data and of lattice or global-fit moment calculations.")

result("Result 8 — scorecard.","The barriers usually fatal to a 'space is a medium' theory now carry concrete "
"in-model demonstrations: emergent Lorentz invariance, emergent fermions (a Dirac cone plus a single chiral "
"fermion on a domain wall), a proper relativistic QFT on quantization, and an emergent photon. More striking than "
"the individual results is that one principle — everything from one structure — surmounts several at once. Even quantum "
"mechanics is largely condensate mechanics (§8.6): the Schrödinger wave, ℏ, the Born rule and de Broglie's "
"λ = h/p emerge; §8.50 then einselects the pointer basis and Born-weighted branches from the medium's own phonon "
"bath (decoherence rate ∝ separation², an improper mixture), leaving only the single-outcome selection — the "
"measurement problem's hard core — as a postulate. Cast against experiment, the "
"model yields a specific, falsifiable Lorentz-violation signature (§8.8), its first genuine prediction. Gravity "
"took the longest and required a retraction: §8.9 records the failure of every elastic route (the curvature sector is "
"unshieldable but repulsive and growing; the tetrad is long-range in field but shieldable and force-free by "
"Eshelby-Crum), and §8.10 then solves it by applying the project's own range principle to gravity's mediator for the "
"first time — the condensate's unprotected, hence gapped, amplitude mode, which is exactly why gravity always looked "
"screened, and which at criticality yields Newton's law, 1/r^2, universally attractive. Sections 8.11-8.12 then carry "
"gravity from that scalar force to the tensor theory: the confining curvature sector deconfines into a Newtonian "
"graviton once the Sakharov loop supplies a positive Einstein term (and that sign is measured, μ > 0, by calibration "
"against the model's own healthy photon), and the radiative spin-2 graviton is dynamical and healthy in 3+1D. The "
"Einstein normalization γ = 1 follows from Weinberg only conditionally — for a mass that couples to the graviton's "
"spatial modes — and §8.32-8.37 measure that premise to fail in every smooth channel, so the realized theory is "
"Nordström (γ = 0) with a healthy but matter-decoupled spin-2 graviton alongside it. Section 8.26 then tests the fixed-point language "
"itself and narrows it: what is measured is that the linearised Einstein term is the infrared attractor, "
"empirically rather than by protection. Section 8.13 dissolves the cosmological-constant fine-tuning (the self-sustained condensate "
"vacuum gravitates its grand potential -P, which vanishes at equilibrium for any zero-point energy), and Section 8.14 "
"shows the photon's induction mechanism scales to non-Abelian yang-mills (SU(2), SU(3)), while Section 8.15 shows "
"chirality is consistently realizable alongside it, the anomaly being a quantized identity settled by Callan-Harvey "
"inflow -- now confirmed dynamically (Section 8.17), in the program's first running simulation of two emergent sectors "
"together. Section 8.16 then puts the model against data: the Lorentz-violation signature is "
"safe but not presently falsifiable, the testable prediction is the short-range gravitational γ, and "
"the 10^122 criticality tuning is retracted. Sections 8.18-8.19 then close two more: gravitational back-reaction now runs as a conserving, convergent simulation that binds matter into a virial-satisfying soliton, and the magnitude of G is fixed at O(1) × the Planck area by the physical lattice cutoff. Section 8.20 then shows the field radiates like spin-2 -- propagating at c with two polarizations, and with the monopole channel shut to machine precision, which scalar gravity would have left open -- and Section 8.21 closes the integration, coupling matter, the radiative field and their energy exchange in one evolution so that a source radiates and thereby decays, with the budget balancing. Section 8.22 then takes that closure past the toy on all three of its admitted caveats -- separately, so each is measured rather than bundled: the matter becomes a relativistic quantum Dirac field carried as a many-fermion Slater determinant, the gravity becomes nonlinear (proved by the failure of superposition, machine zero without the vertex and finite with it), and the coupling dependence is measured to be exactly second order over four decades, so extrapolation to physical strength is arithmetic rather than hope. Section 8.23 then supplies the arc's first hard number: general relativity's quadrupole luminosity law, reproduced to sub-percent with a coefficient derived from the model's own normalisation and audited against the GR binary formula before simulating (measured/predicted 0.9924-0.9996, frequency law Ω^6.007, amplitude law M2^2.000). Everything preceding it in the gravity arc was structural and could not be contradicted by a closed-form answer; this could, and was not. Section 8.24 supplies the second: the Peters-Mathews orbital-decay law, with the binding and radiative couplings locked to one Newton constant (g_N = g^2/8) so nothing is tunable per quantity -- the grid luminosity of a Keplerian binary tracks Peters to 0.990-0.997 across separation, the exponent comes out a^-4.983 against the required -5, and the residual shrinks as v/c falls, the signature of neglected higher post-Newtonian terms. Section 8.25 then turns the same scrutiny on the report's own framework, with a negative result: the classical geometry assumed throughout 8.18-8.23 is not merely approximate but inconsistent -- it breaks superposition, it makes a single particle attract itself with the full force of a nonexistent partner, and it fails the Page-Geilker case where the randomness is classical and no interpretation can rescue it. The model's own structure indicates the repair, since h is a collective mode of a quantum medium and so should be quantised. What remains open is the derivation of the specific Standard-Model group, representations, hypercharges and chiral "
"content (all still inputs), the observed value of the cosmological constant, the measurement problem's hard "
"core, and -- sharpened rather than removed by Section 8.22 -- the fact that the geometry is classical throughout: "
"this is semiclassical gravity, and nothing here bears on quantising geometry. Four self-corrections are now on record — the within-sector Lorentz result, the retracted gravity route, the "
"refuted critical-nucleus prediction, and the retracted 10^122 tuning.")
table(["Barrier","Status","Key result"],
 [["Emergent Lorentz","achieved","one round universal cone, violations ~ (E/E_Planck)^2; cross-statistics universality holds once the boson is a fermion composite (§8.5)"],
  ["Fermions","achieved","Dirac cone on honeycomb; single chiral fermion on a domain wall (evades Nielsen-Ninomiya)"],
  ["Quantum mechanics","largely emergent","quantizes to a relativistic QFT; and from the condensate directly (§8.6): the Schrödinger wave + ℏ as a material property, the Born rule as a Valentini attractor (the unique equilibrium of a Nelson diffusion at the medium's own ℏ/2m — though its drift is read from |ψ|, not yet derived from the sub-quantum medium), and de Broglie v=∇(S)/m for a particle's own wave. Only guidance by a separate pilot wave + definite outcomes (measurement) stay a postulate"],
  ["Emergent photon","achieved","the Dirac-node position: a fluctuation of the medium's own bonds, on the fermion cone"],
  ["Long-range gravity","achieved as scalar gravity (§8.10)","the mediator is the condensate's amplitude mode: unprotected, hence gapped — which is why gravity always looked screened. It couples monopolarly to positive-definite energy, and scalar exchange between like charges attracts. Measured: λ*m_A = 1.00, and E = -C exp(-R/λ)/R exactly, so at criticality Newton's law, 1/r^2, universally attractive. The elastic route is provably dead (§8.9: Bitter-Crum + Eshelby-Crum)"],
  ["Newtonian gravity + healthy spin-2 graviton","the graviton measured; the Einstein limit argued, not realized (§8.11-8.12, §8.26, §8.32-8.37)","the confining curvature sector deconfines into a Newtonian graviton once the induced Einstein term μ>0 (measured, by calibration against the healthy photon), and the radiative spin-2 graviton is dynamical and healthy in 3+1D. γ=1 follows from Weinberg only conditionally — for a mass that couples to the graviton's spatial modes — and §8.32-8.37 measure that premise to vanish in every smooth channel, so the realized theory is Nordström (γ=0). The Einstein normalization is emergent (diffeomorphism invariance is not a lattice symmetry), not lattice-exact; the magnitude of G stays cutoff-dependent; and §8.49 measures even the topological disclination channel to give zero (the mass compresses the medium but nucleates no net curvature charge), so γ=1 is closed in every channel, continuum and discrete"],
  ["Spin-2 graviton (dynamical)","achieved in 3+1D (§8.12)","the transverse-traceless graviton is non-dynamical in 2+1D (0 polarizations) but dynamical in 3+1D (2 polarizations): the induced TT kinetic term is nonzero, the two polarizations are degenerate (helicity 2), and the mode is healthy (same sign as the transverse photon)"],
  ["Empirical prediction","one genuine signature, near-Planckian (§8.8, §8.39)","the surviving falsifiable prediction is the n=2, species-universal, subluminal Lorentz violation; confronted with UHE astrophysics (§8.39) it sits ~1.4 orders below the frontier, not 16, and next-generation UHE observatories can reach it. The gravitational 'scale-dependent γ' once offered as the reachable prediction (§8.11) is RETRACTED: §8.32–8.49 measure γ = 0 at every scale, so γ does not climb to 1 and there is no GR-at-long-range to deviate from. Short-range gravity still bounds the scalar amplitude mode's gap (m_A >~ 4 meV, near the dark-energy scale), but that constrains the mode's range, not a light-bending signature. §8.51 sharpens the LV prediction: its coefficient is lattice-independent (E_QG,2 ≈ 2x Planck across sc/bcc/fcc), so 'n=2, subluminal, one-cone, near-Planckian' is firm structural content, not a tuned number"],
  ["Gravitational back-reaction","runs as a conserving simulation (§8.18)","matter sources the potential and the potential moves matter, solved together: energy conserved to ~10^-9 and norm to ~10^-14, a self-gravitating bound soliton forms (a gravity-off control disperses), the relaxed soliton satisfies the virial identity 2T+W=0, and both converge under mesh refinement. Scope: non-relativistic, scalar/Newtonian, classical matter; the radiative spin-2 sector is not evolved"],
  ["Magnitude of G","fixed at the Planck area (§8.19)","the Sakharov cutoff-ambiguity does not apply because the cutoff is physical (a0 = l_Planck): over the full Brillouin zone the induced stiffness is O(1) in lattice units, so G = O(1) a0^2, with G ~ a0^2/N_f. Gravity is weak because a0 is Planckian -- no hierarchy, no tuning. The O(1) coefficient stays scheme-sensitive"],
  ["Radiative back-reaction","integration closed for classical matter (§8.21)","matter, a dynamical radiative field and their energy exchange run in one evolution from one Hamiltonian: the matter energy falls by exactly what the field energy rises (2.0451×10^-2 vs 2.0452×10^-2, total conserved to 2×10^-6), a spherical control radiates 4×10^5 times less, and E_rad/g^2 is flat. Radiation reaction is derived, not inserted. Scope: classical, non-relativistic matter and linearised gravity, at exaggerated coupling"],
  ["Relativistic quantum matter + nonlinear gravity","the three caveats addressed separately (§8.22)","the matter becomes a Dirac field carried as a many-fermion Slater determinant -- evolution exact for the matter since the coupling is one-body, Pauli holding to 5×10^-12, budget closing to 8.6×10^-9, spherical control 2×10^8 times smaller; the field gains the derivative self-coupling of general relativity, conserving energy (Hamiltonian, not patched) while superposition fails (5.7×10^-16 at zero coupling, 2.1×10^-3 with the vertex on); and E_rad/g^2 is flat to six figures over four decades, so extrapolation to physical coupling is arithmetic. still open: the geometry is classical (semiclassical gravity), the cubic vertex is the structural nonlinearity not the resummed Einstein-Hilbert series, and physical coupling is extrapolated, never simulated"],
  ["Quadrupole luminosity law","reproduced to sub-percent, nothing fitted (§8.23)","the gravity arc's first result that a known closed-form answer could have contradicted -- all earlier ones were structural (a sign, a rank, a scaling, a prohibition, a balance). L = g^2/(160 π) Q⃛.Q⃛ is derived from the model's own field normalisation and checked against GR's binary law before simulating; measured/predicted = 0.9924/0.9983/0.9996, the frequency law comes out Ω^6.007 (exactly 6 required, and the very claim §8.20 had to drop) and the amplitude law M2^2.000. The supporting identity ∫T_ij = (1/2) Ï_ij holds to 5×10^-8 given the centripetal binding stress. Scope: the source is prescribed -- compactness, not prescription, was the obstacle"],
  ["Orbital decay (Peters-Mathews)","binary inspiral reproduced, second hard number (§8.24)","the consequence of the radiation law for a bound system, against a second independent closed form -- the law confirmed on the Hulse-Taylor pulsar. GR has one Newton constant, so g_N = g^2/8 locks the binding and radiative sectors and nothing is tunable per quantity. Grid luminosity of a Keplerian binary vs Peters: 0.9903/0.9924/0.9939/0.9973; exponent L ~ a^-4.983 (required -5); da/dt tracks Peters at the same accuracy; integrating gives the (t_c-t)^{1/4} chirp. The residual shrinks as v/c falls (0.160 → 0.131) -- the signature of neglected higher PN terms, not error. Scope: adiabatic reaction (grid L through energy balance), not a self-force; measured at fixed separation because a moving-orbit energy budget conflates radiation with changing near-zone energy"],
  ["Classical geometry","semiclassical gravity shown inconsistent (§8.25)","an honest-negative on the framework used from 8.18 on: sourcing a classical field on <T> breaks the superposition principle (2.8×10^-15 off vs 4.6×10^-1 on), makes a single particle attract itself with the full Newtonian force of a partner that does not exist (ratio 1.000), and gets wrong the Page-Geilker case where the randomness is classical and no interpretation can rescue it (branch-wise 0.776 vs semiclassical 0.000). The model's own structure indicates the repair -- h is a collective mode of a quantum medium, so it should be quantised -- and quantising one mode gives matter-geometry entanglement plus decoherence, matching the exact answer to 10^-16. Limit: linearised quantum gravity, the easy part; the measurement problem is untouched"],
  ["Infrared fixed point (the claim itself)","measured, and narrowed to an attractor (§8.26)","the report's largest claim, audited. earned: the far field forgets the ultraviolet coefficient exponentially (Newton to ten figures across 10000× in κ), and the operators sort as a fixed point requires -- higher-derivative structure irrelevant, a graviton mass relevant, with any induced mass excluded above ~3e-5 of the Einstein scale. not earned: the fixed point is empirical not protected (no exact lattice symmetry forbids the relevant deformation), the analysis is linearised, and γ=1 stays argued. Wording corrected from 'general relativity as a fixed point' to 'the linearised Einstein term as an infrared attractor'"],
  ["Induced graviton mass","measured: the tetrad is not protected (§8.27)","the assumption under the whole arc, tested. The propagator has no mass term written in it, and 8.12 discards Π(0) as a contact term -- that discarded number IS the mass candidate. Computed with no perturbative bookkeeping, as the sea energy under a constant deformation (the q→0 limit to all orders at once): the photon gives -1.5e-10, falling under refinement, because a symmetry forbids it; the tetrad cone shear gives -0.27, order unity, eight orders larger, stable across gap and grid -- and negative, so it destabilises the symmetric cone rather than merely making it massive. Corroborates the dead elastic route (8.9) by an unrelated route. Does not measure the deconfined curvature sector, so the assumption is narrowed, not removed"],
  ["Curvature-sector q=0 mass","it IS the cosmological constant, and it vanishes (§8.28)","the last load-bearing assumption, closed. Expanding √g Λ gives a derivative-free quadratic term -(Λ/8)[2h_ij h_ij - (tr h)^2] whose transverse-traceless part is nonzero, so m^2 is proportional to Λ for the propagating spin-2 modes (verified against the closed form to 8e-10). Bare it is order unity, as fatal as the tetrad's; the self-sustained condensate vacuum removes it, since what gravitates is the grand potential -P, cancelling to available precision over 122 decades against a rigid-vacuum control retaining 0.75. Residual mass ~2e-8 of the bare scale, ~1000x below 8.26's exclusion bound. Ceiling: protection by an equilibrium condition, not a symmetry, and not independent of 8.13"],
  ["Graviton diffeo invariance","measured: the tetrad action is not Einstein-Hilbert (§8.29)","Weinberg's theorem needs a diffeomorphism-invariant quadratic action as well as a massless spin-2, and that hypothesis had always been assumed. Reopened because 8.28 removed the inhomogeneous <T> obstruction, and measured nonperturbatively as the sea energy under a finite-wavelength deformation, so no seagull can be missing. Calibrated against 8.27 at q=0 to ten digits and against an exact finite-q photon Ward identity (pure gauge 2.2e-11 vs transverse 6.4e-4). The tetrad fails: the pure-gauge response survives removal of the mass term, its ratio to the invariant response is flat in q (marginal, not irrelevant, so it never flows away), rotational invariance breaks by 12.4 per cent, and the fitted coefficients (1, 8.18, -0.56, -0.12) miss Einstein-Hilbert (1, -2, 2, -1) with one sign wrong. Does not touch the deconfined curvature sector; gamma = 1 now rests on it alone"],
  ["Graviton gauge null space","projection cannot rescue it (§8.30)","the loophole 8.29 left. The model's gravity is the incompatible curvature sector, so one might hope the diffeo violation lives in a gauge subspace a projection discards. Measuring the full 6x6 induced quadratic form settles it: the form has NO gauge null space (six nonzero eigenvalues where Einstein-Hilbert has three zeros), the gauge directions are its STIFFEST modes, and the physical-gauge mixing is 42 per cent of the physical block -- concentrated in the spin-0 sector where gamma is defined. A projection cannot remove a violation coupled into the modes it keeps. The 8.29 negative is structural, not a near miss"],
  ["Nonlinear self-coupling","fixed by Deser's bootstrap, not free (§8.31)","8.22 swept the cubic coupling lambda as a free parameter. It is not one: matter and the field enter the same equation as g/2 h S and lambda/2 h dh dh, so lambda/g is a Nordtvedt parameter and the strong equivalence principle fixes lambda = g. Verified independently by the field energy's deformation response to 1e-10 on every stress component. 8.22's five values are lambda/g = 0, 0.067, 0.13, 0.27, 33.3 -- none physical; its headline self-interaction was taken at 33.3, where gravitational binding energy gravitates 33x too strongly. At lambda = g the effect is far smaller and the budget still closes. The parameter was wrong, not the integration"],
  ["Does mass source curvature?","no, in the smooth channel: gamma_smooth = 0 (§8.32)","gamma = kappa/(4 pi G) reduces the light-bending parameter to one number: does a mass source spatial curvature, and at what strength. A static mass is a scalar energy density (T00 ~ E*I), and its induced coupling to the spatial stress vanishes identically -- Pi^{00,ij} = 0 for every component including the trace, to machine precision across mass, cutoff and momentum -- while a genuine spin-2 source couples at O(1), so the zero is a selection rule. The fourth independent closure of the smooth route to gamma = 1 (with 8.12, 8.29, 8.30), and elasticity gave the same zero from the other side (8.10). Every smooth mechanism gives gamma = 0. The TOPOLOGICAL (disclination) channel where real curvature lives is measured separately -- zero in the continuum by Gauss-Bonnet (8.33) and zero on the lattice, where the mass nucleates no net disclination charge (8.49). gamma = 1 now rests entirely on emergent Weinberg, which -- unlike emergent Lorentz -- has no direct confirmation and now has direct evidence against it in every channel"],
  ["Topological channel (curvature charge)","also 0, by Gauss-Bonnet (§8.33)","the one route left after 8.32: curvature is a scalar, so a scalar mass CAN source it, through the model's own compression mechanism (mass = eigenstrain theta* ~ rho). It does curve the medium locally (eta = lap theta* =/= 0), but the net curvature charge integral(eta) = integral(lap theta*) is a total derivative, exactly zero -- a curvature dipole (dome core, saddle ring) with no deficit angle, bending no light at range, vs a genuine charge eta ~ rho that does. gamma = 1 needs mass to carry net curvature charge = nucleate a disclination density ~ rho, which the model does not. So every accessible direct channel gives gamma = 0: spin rule (8.32), compatibility (8.10), Gauss-Bonnet (8.33). The 2D-theorem caveat is resolved in 8.34, and 8.49 confirms it on the lattice -- concentrated energy compresses the medium but nucleates no net disclination charge"],
  ["The Poisson-ratio / propagator route","closed; no nu makes it Fierz-Pauli (§8.35)","the last opening: gamma=1 could come from the PROPAGATOR (a Fierz-Pauli graviton kinetic term turns the time potential into spatial curvature), and the curvature sector's kinetic term is the medium's biharmonic elasticity, set by its Poisson ratio -- not covered by 8.29-8.30 (the induced tetrad). Measured: the LJ condensate is a central-force Cauchy solid, nu ~ 1/3. But the medium's response splits into an elastic RELAXATION (nu-dependent, but a displacement, so COMPATIBLE = zero curvature = bends no light -- verified machine-zero Ricci from nu=-0.9 to 0.49) plus the eigenstrain (nu-independent). So the Poisson ratio tunes only a gauge mode; gamma ray-traced is the same value zero for every nu. The trace structure for gamma=1 cannot be reached by tuning elastic constants. Every route now closed: neither the source (8.32-8.34) nor the propagator (8.29-8.30, 8.35) gives gamma=1. Newtonian gravity real; the Einstein completion not recoverable by any property of the medium"],
  ["The 3+1D answer (does the medium evade it?)","no; gamma = 0 in 3D too (§8.34)","settles 8.33's caveat by direct 3D ray tracing. The medium does NOT reach gamma = 1, for a reason deeper than Gauss-Bonnet and independent of dimension: general relativity sources the spatial metric by a Poisson equation (lap Psi = 4 pi G rho, Psi = the long-range potential, gamma = 1), while the medium's compression sets Psi = theta* = rho ALGEBRAICALLY -- local, dying with the mass. The 3D incompatibility of an isotropic eigenstrain is the linearized Einstein tensor delta lap theta* - d_i d_j theta* (verified to 1e-16), so Psi = theta* = rho. Ray-traced: a GR control (Psi = potential) holds gamma = 1 at every impact parameter; the medium's gamma falls to zero at range (few parts in 1e4, stable under grid/box/width). The 2D zero-charge result was a symptom. Every channel in every dimension gives gamma = 0. gamma = 1 needs mass to source the graviton's SPATIAL polarizations as a Poisson potential -- exactly the coupling 8.32 measured to vanish. Newtonian gravity real and quantitative; the Einstein completion argued, not demonstrated"],
  ["The negative-space sublattice","adds a displacement, not curvature (§8.36)","a second interpenetrating (bcc-like) sublattice was the one construction attacking the single-lattice premise: its relative (optical) mode is not a gradient of the acoustic displacement, and two fcc = diamond, which violates Cauchy through it. Measured (honeycomb): the optical mode is gapped, so slaved -- a second displacement. It genuinely breaks Cauchy (relaxed C12-C66=+0.29) but leaves the bulk modulus invariant, couples only to shear (isotropic strain sources zero shift, machine zero), and the full response to a mass has curvature machine zero. It adds a displacement, not an incompatible degree of freedom; gamma=0 survives"],
  ["The whole gamma arc in one number","marginal, not irrelevant (§8.37)","the Einstein question reduces to the graviton propagator's trace coefficient lambda: for pressureless (light-bending) sources gamma=lambda/(1-lambda), so the factor of two IS the trace term and pressure is a red herring. Einstein sits at lambda=1/2; both measured routes put the model at lambda=0 (source <T00,Tij>=0; induced coefficients (1,+8.18,-0.56,-0.12) vs (1,-2,2,-1)). gamma=1 and emergent Lorentz are the same lattice-broken symmetry with opposite RG fates: Lorentz's anisotropy is irrelevant (~(k/kP)^2, flows away, emerges), the graviton's diffeo violation is marginal (flat ratio, converged 12.4%), so gamma stays off 1 at every scale. That asymmetry is why the model earns SR but not the factor of two"],
  ["Does the photon's width bend light?","no; width is not the missing term (§8.38)","the factor of two is the geometry, not the probe: light samples the spatial curvature Psi equally with the time warp because it moves at c (null path, dx=c dt), not because of any width. A finite-width photon (centroid exact by Ehrenfest) through the model's index bends by the ray amount, implied gamma 0.0004->0.15 across w/b=0.02->0.33 (O((w/b)^2) tidal correction); through a GR index it bends by 2x at every width. Width multiplies whatever gamma the geometry carries, never creates it. gamma=0 robust to any width -- a mass sources no Psi for a probe of any size"],
    ["Gravitational radiation","spin-2, monopole forbidden (§8.20)","the linearised TT field propagates at c with exactly 2 polarizations, and a spherically pulsating source radiates ~10^-14 of an equal quadrupole (ratio 3×10^-13, machine zero). Scalar gravity would radiate the monopole; spin-2 forbids it. Scope: linearised, prescribed source -- inspiral back-reaction and the quadrupole luminosity law are not tested"],
  ["Chirality + anomalies","consistent by inflow (§8.15)","a chiral gauge theory is inconsistent unless anomalies cancel; here bulk Chern number = chiral modes per wall = charge pumped per flux quantum = one integer (measured -1, +1/-1, Σ 0). Each wall is anomalous, the lattice is vector-like, the bulk supplies the inflow. not the SM's own 4D cancellation — the bulk does the work"],
  ["Cosmological constant","fine-tuning dissolved, and unified with gamma=0 (§8.13, §8.40)","the self-sustained condensate vacuum gravitates its grand potential -P, zero at equilibrium for any bare zero-point energy (measured across 122 orders, no tuning). The load-bearing claim -- gravity couples to -P, not the energy density -- is now MEASURED: energy density sources h00 but not spatial curvature (<T00,Tij>=0), the SAME selection rule that gives gamma=0. So gamma=0 and Lambda=0 are one fact; a GR-like light-bending theory would carry the 10^122. Robust: survives Weinberg's no-go (absorbing vacuum energy is a pure dilatation, dimensionless observables invariant to machine precision across 8 orders), and equilibrium is an attractor erasing initial conditions in a Planck time. Open: the observed nonzero value -- the expansion lag gives 10^-61 (wrong size and w~-1/2), tracking needs a deficit-driven efficiency the medium lacks; relocated to why the vacuum sits so near equilibrium"],
  ["Non-Abelian gauge fields","mechanism achieved (§8.14)","the fermion loop induces genuine Yang-Mills for SU(2) and SU(3): a uniform non-commuting field costs ~A^4 = Tr[A,A]^2 (self-interaction), a commuting one is pure gauge. Universal coupling from exact lattice gauge invariance"],
  ["Electroweak breaking","mechanism achieved (§8.41)","the condensate as a doublet with hypercharge 1/2 breaks the induced SU(2)xU(1) to one U(1): the gauge spectrum is {0, m_W, m_W, m_Z} -- exactly one massless photon and three massive weak bosons, m_W=gv/2, m_Z=sqrt(g^2+g'^2)v/2, for any couplings. m_W/m_Z = cos(theta_W) (sin^2 = 0.223 at the physical point); the surviving photon is Q=T3+Y, massless because the vacuum is Q-neutral. The breaking pattern and mass relations are reproduced, not fitted. The group, rep and Y=1/2 are inputs (as in 8.14); why Y=1/2 -- fixed with the fermion charges by anomaly cancellation -- is open"],
  ["Hypercharges / charge quantisation","derived from anomaly cancellation (§8.42)","the model's exact emergent gauge invariance forbids a gauge anomaly, so anomaly cancellation is mandatory. Given only the one-generation representations, the four conditions ([SU3]^2U1, [SU2]^2U1, grav^2 U1, U1^3) fix the five hypercharges uniquely (up to scale and the u<->d relabelling) to the SM values 1/6,-2/3,1/3,-1/2,1. Hence Q=T3+Y is quantised: quarks +2/3,-1/3, leptons 0,-1, proton+electron charge = 0 to 1e-16. Charge quantisation as a consistency requirement. Scope: fixes hypercharges GIVEN reps and content"],
  ["SM group / generations","recast as discrete band topology (§8.43)","not derived, but the CATEGORY is fixed: the emergent gauge group is the symmetry of the lattice's Dirac-point multiplet (honeycomb = 2 valleys), and the generation count is a Chern index (integer, 0/+-1 computed, jumps only at gap closings), so a domain wall carries |ChernJump| chiral families. Why a compact group and a small integer of generations is answered (symmetry+topology of the band structure, quantized); why exactly SU(3)xSU(2)xU(1) and three is a discrete property of the physical medium's lattice, still input. The Yukawa/flavour structure also remains open"],
  ["Which lattice (the last SM input)","reduced to one property of the interaction: angular rigidity (§8.44)","the group and generation count are band data of the medium's lattice (§8.43), so the lattice is the remaining input. A central pair force gives the close-packed triangular/fcc lattice (coordination 6, no Dirac point); with any number of isotropic length scales at most the square lattice -- the bipartite honeycomb is never the ground state and is mechanically floppy (coordination 3 < the 2D Maxwell threshold 4). A 120-degree three-body term makes the honeycomb the ground state above λ* ~ 0.57 of the bond energy, and it then self-assembles and stays rigid (analytic force gated to 1e-10; a heated seed held at coordination 3 vs densifying to 6 without it). So the fermion lattice needs O(1) angular rigidity -- the same non-central stiffness emergent Lorentz required (§8.1). Does not derive which lattice; converts the input into one property of the interaction"],
  ["Angular stiffness does triple duty","fermion lattice + stability + Lorentz cone from one knob (§8.45)","the non-centrality §8.44 needs for the honeycomb is the same one §8.1 needs for a single Lorentz cone. Same force model, elastic constants vs the same lambda: triangular control gives c_L/c_T=sqrt3 and nu=1/3 (matches the medium's 1.7330, 0.330); central-force honeycomb has mu=0 (floppy, coordination 3 < Maxwell 4, no transverse cone). Turning on O(1) angular stiffness selects the honeycomb (lambda*~0.57), rigidifies it, AND is the SOLE source of its transverse cone: bulk modulus K is lambda-independent to machine precision (compression is central), shear modulus mu and c_T are entirely angular-sourced (0 -> O(1), c_T~sqrt(lambda)). Full c_L=c_T is the mu>>K limit past lambda* (§8.1's vector-Hooke). Two independent barriers turn on one property of the medium"],
  ["Mechanical vs fermion cone","two cones, but the mismatch is benign (§8.46)","the angular stiffness gives a mechanical cone c_T (§8.45) while the Dirac fermions have their own cone v_F=(3/2)ta from the hopping -- independent knobs, equal only by tuning (t~1.11). But the emergent matter and the composite photon/graviton ride v_F (the interband particle-hole collective mode -> v_F|q|), and a phonon is an off-diagonal bond perturbation = a pseudo-gauge field: it moves the Dirac point without gapping it (chiral protection; only an O(1) Lifshitz distortion annihilates the cone, and an on-site sublattice mass -- which a phonon cannot make -- gaps it at once). So v_F, the emergent Lorentz cone, is insulated from the mechanical sound; c_T is a decoupled sub-quantum spectator (Volovik). Scope: static-phonon (gauge-field) protection; the dynamical self-energy is a further computation"],
  ["Dirac cone vs phonon fluctuations","survives to all orders; whether cones merge is regulator-limited (§8.47)","a phonon is a bond (off-diagonal) term, chiral-odd under S H S=-H, so it cannot make the chiral-even sublattice mass that gaps the cone -- at any order (bond and mass carry exactly opposite chiral parity, machine zero; random bond disorder stays gapless to O(1), a staggered mass gaps at once). Closes the DYNAMICAL two-cone seam: the mechanical sound can't spoil the fermion cone, statically or dynamically, by exact symmetry. OPEN (regulator-limited): whether v_F and c_T actually MERGE (full Lorentz) is a two-velocity RG with no covariant cutoff -- one loop fails its own gamma_v=gamma_c gate at v=c, like §8.29's Ward identity. So the mismatch is harmless; whether it vanishes is open"],
  ["Generation number","a measured band invariant, not a tuned parameter (§8.48)","the family number is the Chern number of the two-band medium fermion = the winding of its inter-sublattice coupling around the BZ. Read out two agreeing ways for winding n=1,2,3: the bulk Fukui-Hatsugai-Suzuki lattice flux gives C=-n exactly, and a cylinder carries exactly n chiral edge branches, MEASURED from the ribbon spectrum (upgrades §8.43's asserted bulk-boundary count to a measurement). The two edges carry opposite chirality (fitted slopes +-n), so by Nielsen-Ninomiya the closed lattice nets to zero: the n net-chiral generations must live on a defect (edge/domain wall) with a compensating mirror sector. The winding-n point is a fine-tuned multi-Weyl degeneracy that a generic perturbation fragments into n unit Dirac points (total winding conserved), so the robust count is n ordinary Dirac points -- unifying §8.43's flavour count [A] and Chern index [B] as one conserved integer. Does NOT derive three -- the winding is a free integer, so 'why three' becomes the measured geometric statement 'the coupling winds three times', the count now a quantized band invariant rather than a continuous input"]],
 cap="Table 5.  The fundamental-physics barriers and their status in the model.")

# ===== 9 Synthesis =====
heading("9  Synthesis", 1)
body("Read as a whole, the program builds one object — a medium carrying a field — and promotes it step by "
"step: a real field (particles), a complex field (charge and spin), a mobile medium (robust topology), a "
"density coupling (gravity), a larger target space (3D point charges), and a gauge field (electromagnetism). "
"At each step the same question recurs and receives the same kind of answer. A bare topological charge's "
"self-energy always grows with system size — logarithmically for a 2D vortex, linearly for a 3D line or a 3D "
"hedgehog — and a strong long-range force between charges appears only when the mediator is massless and "
"protected by a symmetry: a global Goldstone mode, or a gauge photon in the Coulomb phase. Gauge the same "
"symmetry into its broken phase and the force screens. From one medium and one field, the model thus "
"reproduces the qualitative menu of real forces and both gauge phases, with the reach of each governed by a "
"single principle.")

# ===== 10 Limitations =====
heading("10  Limitations and scope", 1)
body("This is a toy model, and the results are qualitative correspondences, not derivations of the Standard "
"Model or general relativity. The gravity sector is weak and non-self-binding on the nearly incompressible "
"medium; a definitive compressible-medium test awaits a properly conservative soft potential. The 3D "
"point-charge interaction of Route B is measured on a pinned, relaxed texture, and Route C's Coulomb phase is "
"the classical minimum of the Maxwell energy rather than a full quantum treatment. Range claims of the "
"long-range kind rest on system-size scaling rather than infinite volume. The correspondences are "
"qualitative — 'EM-like', 'nuclear-like' — and quantitative matching of coupling constants or mass ratios is "
"not attempted. What survives these caveats is the central, dimension-robust and repeatedly gate-checked "
"statement: within this medium, a force's range is set by whether a symmetry protects its mediator from a "
"mass term.")
body("The fundamental-physics program of Section 8 carries a further, honest qualification. Most of those results are "
"established at the level of the dispersion relation, the band structure, and the defect algebra rather than a single "
"running simulation. Sections 8.17 and 8.18 take the first steps past this. Section 8.17 runs chiral matter and a "
"gauge field together in time and recovers the anomaly as real charge transport; Section 8.18 runs gravity with "
"genuine back-reaction -- matter sourcing the potential and the potential moving matter, self-consistently, with "
"energy conserved to parts in a billion, a self-gravitating bound state formed, the virial identity satisfied, and "
"convergence under mesh refinement. Those are the standards a scientific simulation must meet, and in those sectors "
"the model now meets them. What is still not a single running simulation is the full target: emergent "
"Lorentz-invariant, chiral, quantum matter interacting through the emergent spin-2 gravity with radiative "
"back-reaction. That coupling is now done, in Section 8.21: matter, a dynamical radiative field and a self-consistent "
"exchange of energy between them run in one evolution from one Hamiltonian, and a source radiates and thereby decays "
"with the budget balancing to four significant figures. The integration gap this report kept naming is therefore "
"closed. It was closed first only for classical, non-relativistic matter coupled to linearised gravity at exaggerated "
"coupling, and Section 8.22 then addressed those three caveats one at a time: the matter is now a Dirac field carried "
"as a many-fermion Slater determinant, the field carries the derivative self-coupling of general relativity (proved "
"active by the failure of superposition), and the coupling dependence is measured to be exactly second order across "
"four decades, which makes the extrapolation to physical strength arithmetic. What remains genuinely open there is "
"narrower but not smaller: the geometry is classical. Every result in this report treats the gravitational field as a "
"classical object sourced by quantum matter -- semiclassical gravity -- and Section 8.25 establishes that this is not "
"a gap to be acknowledged but a prescription that is inconsistent: it breaks the superposition principle, it makes a "
"single particle attract itself with the full Newtonian force of a partner that does not exist, and it fails the "
"Page-Geilker case in which the randomness is classical and no interpretation of quantum mechanics can rescue it. The "
"model's own structure indicates the repair -- h is a collective mode of a quantum medium, so it should be quantised "
"as phonons are, and quantising one mode yields matter-geometry entanglement and gravitational decoherence. But that "
"is linearised quantum gravity, the easy and long-known part; it is not a theory of quantum geometry, and it does not "
"select an outcome, so the measurement problem below stands untouched. The results of Sections 8.18-8.23 concern "
"regimes without macroscopic superposition, where the mean-field treatment is a controlled approximation, and they "
"are not retracted -- but the framework carrying them cannot be the final story. Alongside this, the cubic vertex is "
"the structural nonlinearity rather than the resummed Einstein-Hilbert series: "
"there is no black hole in this model and none is claimed. Physical coupling is reached by extrapolation, not by "
"simulation, and cannot be reached by any direct simulation in double precision. Each barrier is met "
"individually, and two pairs now jointly; assembling them all into one consistent theory that also fixes the "
"Standard-Model content and the constants of nature is the work of fundamental physics itself, not of this toy. The value of Section 8 is to show that these barriers, usually treated as fatal to any "
"'space is a medium' program, are here concrete and — one at a time — surmountable. The quantum-mechanics result "
"of Section 8.6 carries its own explicit boundary: the wave, ℏ, the Born statistics and de Broglie's relation "
"emerge, but the guidance of a particle by a separate pilot wave and the selection of a single definite outcome are "
"not derived — the measurement problem is left where it stands for every interpretation of quantum theory.")

# ===== 11 Conclusion =====
heading("11  Conclusion", 1)
body("Starting from space as an active medium, the project produced persistent particles, gave them charge "
"and spin as topological and Noether quantum numbers, showed those numbers robust on a self-organizing "
"substrate, and derived gravity from the medium — finding it real but short-ranged. Pursuing why, it "
"identified screening as a mass term and a conservation law as the escape, then realized, from one field "
"promoted step by step, a long-range Goldstone force, a mass-screened force, a gauge-screened (Meissner) "
"force, and finally a genuine 1/r² electromagnetism between quantized monopoles in three dimensions. The "
"model did not fail to make gravity Newtonian so much as explain, in its own terms, which forces are "
"long-range and why. It then turned the same measuring discipline on the question of whether it could be "
"fundamental, and found the classic obstacles — Lorentz invariance, chiral fermions, quantization, and even "
"spin-2 gravity — not fatal but, one at a time, met: an emergent Lorentz-invariant cone, a single chiral "
"fermion on a domain wall, and a relativistic quantum field theory. Pressed further, these separate victories "
"collapsed into one "
"principle. A field bolted on beside the medium brings a second light cone; a field made of the medium inherits "
"the first. Carried to its end, the medium's own bond fluctuations, seen by its fermions, are the emergent "
"photon and the emergent graviton — so matter, light and gravity share a single cone not by tuning but by "
"construction. That graviton was then made dynamical: a luminal spin-2 wave whose Einstein-Hilbert kinetic term is "
"induced from the fermion loop rather than imposed. Turned on quantum mechanics "
"itself, the same discipline found most of it to be condensate mechanics: the Schrödinger wave and ℏ as a "
"material property, the Born rule as a stochastic attractor, and de Broglie's λ = h/p all emerge, with only "
"the guidance of a particle by a separate wave and the selection of a single outcome — the measurement problem's "
"hard core — left as a postulate. And, for the first time, the model made a prediction rather than a reproduction: "
"a Planck-suppressed, cross-species-universal, crystallographically-anisotropic quadratic Lorentz violation that "
"survives every present bound. Put against real numbers in Section 8.39, that signature is not the safe-by-sixteen-"
"orders curiosity it first seemed: because the violation is quadratic its strongest bounds are the near-Planckian "
"ones from ultra-high-energy astrophysics, and the model's predicted quantum-gravity scale sits only about 1.4 "
"orders above the current UHECR frontier — safe today, but within reach of next-generation ultra-high-energy "
"observatories. That is the model's one genuine, still-open falsifiable prediction. The gravitational alternative "
"once offered as the reachable one — a scale-dependent light-bending parameter below the amplitude mode's Compton "
"wavelength — is retracted: Sections 8.32–8.49 measure γ = 0 at every scale, so light-bending never climbs toward "
"the Einstein value and there is no scale-dependent deviation to look for.")
body("The same discipline that produced those results also destroyed one, and then rebuilt it. Pressed on why gravity "
"cannot be shielded, the model gave a clear answer -- screening is neutralization, not loss, and a "
"topological charge is unneutralizable, so quantization does in the medium what the absence of negative mass does in "
"nature. But applying the same standard of proof to the force, rather than to the field, brought the whole "
"gravitational programme down: two like curvature charges repel with a force that grows with separation, and the "
"tetrad graviton, for all its genuine long-range 1/r^2 field, exerts no long-range force at all. Every one of those "
"failures was an elastic calculation, and naming that gave the answer. A mass is a force dipole, so the elastic "
"sector cannot host gravity for any choice of moduli — while the principle the project had established for every "
"other force, that range is set by whether A symmetry protects the mediator from A mass term, had never once been "
"applied to gravity's own mediator. Applied at last, it identifies that mediator as the condensate's amplitude mode: "
"unprotected, therefore gapped, therefore Yukawa — which is exactly why every earlier measurement had found gravity "
"screened — and coupled monopolarly to positive-definite energy, so that its exchange between like charges is "
"universally attractive, with no sign inserted by hand. Its range is the inverse gap, measured; its potential is a "
"screening exponential times a 1/r Newtonian core, measured; and at criticality the exponential goes to unity and "
"Newton's law stands alone. Gravity, in this medium, is what the medium's own principle always said it would be.")
body("That scalar gravity was then carried the rest of the way to the tensor theory. The medium's curvature sector — "
"the graviton's true home — is confining on its own, a string tension rather than a force; but the same fermion loop "
"that induces the photon induces an Einstein term that deconfines it into a Newtonian graviton, and the sign that "
"decides whether this works was measured, positive, by holding the induced gravity against the model's own healthy "
"photon. In the physical four dimensions that graviton is a genuine dynamical spin-2 wave with two degenerate "
"helicity-2 polarizations — a firm output. Its Einstein normalization — light bending by the famous factor of two — "
"is another matter. Weinberg's theorem forces it only conditionally, for a mass that actually couples to the spin-2 "
"mode's spatial polarizations, and Sections 8.32-8.37 measure that coupling to vanish in every smooth channel the "
"model exposes; so the factor of two is argued, not realized. And unlike the model's emergent Lorentz invariance — "
"whose lattice breaking is irrelevant and flows to zero in the infrared — the diffeomorphism breaking that pins γ "
"off unity is marginal and does not flow away (Section 8.37), so the two are not, in fact, emergent in the same "
"sense. As realized the model is Nordström: a healthy scalar gravity carrying a matter-decoupled spin-2 graviton "
"alongside it; the Einstein completion is closed in every channel the model exposes, the last of them — whether a "
"mass nucleates net topological disclination charge — measured directly on the lattice in Section 8.49 and found "
"absent. The same condensate structure then "
"disarms the deepest quantitative disaster in physics: the vacuum's enormous zero-point energy does not gravitate, "
"because a self-sustained condensate gravitates its grand potential, which vanishes at equilibrium for any bare "
"value — the cosmological-constant fine-tuning dissolved, though not its observed residue derived. And the induction "
"mechanism that built the photon builds non-Abelian Yang-Mills as readily, for SU(2) and SU(3) alike, with chirality "
"riding alongside it consistently: the anomaly on a domain wall is a quantized integer, and the bulk supplies exactly "
"what each wall loses. The same confrontation with data that deflated the Lorentz prediction also retracted the "
"model's other great fine-tuning — the 10^122 approach to criticality, which had rested on the superseded scalar "
"reading of gravity — so that both of the program's 10^122 tunings have now fallen, each to a structural result "
"rather than a fitted parameter.")
body("What remains is honest and specific. The linearised Einstein term is reached as an infrared attractor, not as a "
"lattice-exact law, and Section 8.26 bounds what that claim can carry: universality and the operator "
"sorting are measured, but the fixed point is empirical rather than symmetry-protected, the nonlinear "
"Einstein equations are not reached — the magnitude of Newton's constant is cutoff-dependent, and a direct lattice-exact measurement "
"of γ = 1 is structurally unavailable because diffeomorphism invariance is emergent by construction. The observed "
"nonzero cosmological constant, the derivation of the specific Standard-Model group and its representations and "
"constants (the group remains an input, only its Yang-Mills dynamics are induced), the residual measurement problem, "
"and quantitative rather than qualitative correspondences are, in the spirit of the project, measurements waiting to "
"be made. A model that can be made to say something false, and then be caught doing it, and then be made to say "
"something true, is the only kind worth building.")

# --- references ---
heading("References", 1)
refs=[
 "J. D. Eshelby, The continuum theory of lattice defects, Solid State Physics 3, 79 (1956). (Bitter–Crum; centres of dilatation.)",
 "J. Goldstone, Field theories with «superconductor» solutions, Nuovo Cimento 19, 154 (1961).",
 "J. M. Kosterlitz and D. J. Thouless, Ordering, metastability and phase transitions in two-dimensional systems, J. Phys. C 6, 1181 (1973).",
 "A. A. Abrikosov, On the magnetic properties of superconductors of the second group, Sov. Phys. JETP 5, 1174 (1957).",
 "P. W. Anderson, Plasmons, gauge invariance, and mass, Phys. Rev. 130, 439 (1963). (The Anderson–Higgs mechanism.)",
 "G. 't Hooft, Magnetic monopoles in unified gauge theories, Nucl. Phys. B 79, 276 (1974); A. M. Polyakov, JETP Lett. 20, 194 (1974).",
 "T. A. DeGrand and D. Toussaint, Topological excitations and Monte Carlo simulation of Abelian gauge theory, Phys. Rev. D 22, 2478 (1980).",
 "J. C. Maxwell, On the calculation of the equilibrium and stiffness of frames, Phil. Mag. 27, 294 (1864). (Central-force rigidity.)",
 "G. E. Volovik, The Universe in a Helium Droplet, Oxford (2003). (Emergent relativity, gauge fields and gravity near a Fermi point.)",
 "H. B. Nielsen and M. Ninomiya, Absence of neutrinos on a lattice, Nucl. Phys. B 185, 20 (1981). (Fermion doubling.)",
 "D. B. Kaplan, A method for simulating chiral fermions on the lattice, Phys. Lett. B 288, 342 (1992); C. G. Callan and J. A. Harvey, Nucl. Phys. B 250, 427 (1985). (Domain-wall chiral fermions.)",
 "M. Pretko and L. Radzihovsky, Fracton-elasticity duality, Phys. Rev. Lett. 120, 195301 (2018).",
 "S. Weinberg and E. Witten, Limits on massless particles, Phys. Lett. B 96, 59 (1980).",
 "A. D. Sakharov, Vacuum quantum fluctuations in curved space and the theory of gravitation, Dokl. Akad. Nauk SSSR 177, 70 (1967). (Induced gravity.)",
 "F. Guinea, M. I. Katsnelson and A. K. Geim, Energy gaps and a zero-field quantum Hall effect in graphene by strain engineering, Nature Phys. 6, 30 (2010). (Strain-induced emergent gauge fields.)",
 "K. G. Wilson, Confinement of quarks, Phys. Rev. D 10, 2445 (1974). (Lattice gauge theory; exact lattice gauge invariance and the Wilson plaquette action.)",
 "S. L. Adler, Einstein gravity as a symmetry-breaking effect in quantum field theory, Rev. Mod. Phys. 54, 729 (1982). (Induced gravity; the sign of the induced Newton constant.)",
 "S. Weinberg, Photons and gravitons in S-matrix theory: derivation of charge conservation and equality of gravitational and inertial mass, Phys. Rev. 135, B1049 (1964). (A massless spin-2 coupled to a conserved stress tensor is Einstein; γ = 1.)",
 "B. P. Abbott et al. (LIGO Scientific, Virgo, Fermi-GBM, INTEGRAL), Gravitational Waves and Gamma-Rays from a Binary Neutron Star Merger: GW170817 and GRB 170817A, Astrophys. J. Lett. 848, L13 (2017). (The 1.74 s delay bounds |v_gw − v_em|/c between −3×10⁻¹⁵ and +7×10⁻¹⁶; used in §8.52.)",
 "T. Jacobson, S. Liberati, D. Mattingly, Lorentz violation at high energy: concepts, phenomena and astrophysical constraints, Annals Phys. 321, 150 (2006), arXiv:astro-ph/0505267. (Dimension-six n=2 threshold analysis and the proton bound |η₄| ≲ 10⁻²; used in §8.54.)",
 "Physical-point lattice-QCD determination of nucleon PDF moments, arXiv:2605.02808 (2026). (The isovector second moment ⟨x²⟩_{u-d} = 0.083(14) at μ = 2 GeV; used in §8.56 to bound the proton compositeness suppression.)",
 "R. Voss, Butler–Voss Condensate — project reference (CHEATSHEET) and source repository, 2026 (private).",
]
for i,rf in enumerate(refs,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3); p.paragraph_format.first_line_indent=Inches(-0.3)
    p.paragraph_format.space_after=Pt(4)
    nr=p.add_run(f"[{i}]  "); nr.bold=True; nr.font.size=Pt(9.5); nr.font.color.rgb=GREY
    p.add_run(rf).font.size=Pt(9.5)

apx=doc.add_paragraph(); ar=apx.add_run("Implementations (pure NumPy; private repository): simulation.py, "
"prototype_complex.py, h9_binding.py, prototype_mobile_nodes.py (H1–H10); integration_field_medium.py, "
"integration_phase2.py (§4); integration_phase3_variational.py, integration_phase3d.py, "
"density_response_3d_large.py (§5); screening_diagnosis.py, screening_gauss.py, screening_topocharge.py, "
"screening_gauged.py, screening_gauged_mobile.py (§6); screening_topocharge_3d.py, route_b_hedgehog.py, "
"route_c_monopole.py (§7); test_lorentz.py, test_lorentz_unified.py, test_lorentz_boost.py, test_dirac.py, "
"test_domain_wall.py, test_quantization.py, test_graviton.py, test_fracton_gravity.py, test_graviton_spin2.py, "
"test_cone_universality.py, test_cone_lock.py, test_induced_action.py, test_emergent_tetrad.py (§8.1-8.5); "
"test_emergent_qm.py, test_born_rule.py, test_double_solution.py (§8.6); test_graviton_dynamics.py, "
"test_induced_gravity.py (§8.7); test_lv_prediction.py (§8.8); test_shielding.py, test_tetrad_shielding.py, "
"test_disclination_force.py, test_tetrad_force.py (§8.9); test_critical_gravity.py (§8.10); test_two_gravities.py, "
"test_deconfinement.py (§8.11); test_induced_sign.py, test_spin2_dynamical.py, test_lattice_ward.py (§8.12); "
"test_cosmological_constant.py (§8.13); test_yang_mills.py (§8.14); test_anomaly_inflow.py (§8.15); "
"test_experimental_bounds.py (§8.16); "
"test_realtime_pump.py (§8.17); test_backreaction.py (§8.18); "
"test_newton_constant.py (§8.19); "
"test_gravitational_radiation.py (§8.20); "
"test_radiative_backreaction.py (§8.21); "
"test_relativistic_backreaction.py (§8.22); "
"test_quadrupole_luminosity.py (§8.23); "
"test_inspiral_peters.py (§8.24); "
"test_semiclassical_inconsistency.py (§8.25); "
"test_ir_fixed_point.py (§8.26); "
"test_graviton_mass.py (§8.27); "
"test_curvature_mass.py (§8.28); "
"test_graviton_transversality.py (§8.29); "
"test_graviton_nullspace.py (§8.30); "
"test_deser_bootstrap.py (§8.31); "
"test_gamma_source.py (§8.32); "
"test_gamma_topological.py (§8.33); "
"test_gamma_3d.py (§8.34); "
"test_gamma_elastic.py (§8.35); "
"test_gamma_sublattice.py (§8.36); "
"test_gamma_trace.py (§8.37); "
"test_photon_width.py (§8.38); "
"test_lv_confrontation.py (§8.39); "
"test_vacuum_gravitates.py, test_cc_weinberg.py, test_cc_offequilibrium.py, test_cc_attractor.py, "
"test_cc_tracking.py (§8.40); "
"test_electroweak.py (§8.41); "
"test_anomaly_hypercharge.py (§8.42); "
"test_sm_structure.py (§8.43); "
"test_lattice_selection.py (§8.44); "
"test_cone_unification.py (§8.45); "
"test_two_cones.py (§8.46); "
"test_chiral_protection.py (§8.47); "
"test_generations.py (§8.48); "
"test_validation_anchors.py (§8.48, un-tuned Haldane + graphene cross-checks); "
"test_gamma_nucleation.py (§8.49, disclination-nucleation channel); "
"test_decoherence.py (§8.50, einselection from the medium's phonon bath); "
"test_lv_robustness.py (§8.51, LV-coefficient robustness across lattices); "
"test_gw170817_onecone.py (§8.52, first real-data confrontation: GW170817 one-cone timing); "
"test_lv_uhe_reach.py (§8.53, UHE-frontier confrontation: Auger/LHAASO/IceCube); "
"test_lv_gzk_threshold.py (§8.54, universal-LV GZK threshold: the coefficient in tension); "
"test_lv_proton_compositeness.py (§8.55, the parton suppression that pulls it to the GZK frontier); "
"test_lv_proton_moment_data.py (§8.56, real PDF-moment error bars on the suppression).")
ar.font.size=Pt(8.5); ar.font.color.rgb=GREY; ar.italic=True; apx.paragraph_format.space_before=Pt(12)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
if os.path.exists(OLD):
    os.remove(OLD)
print("SAVED:", OUT)
print("removed old screening-only doc:", not os.path.exists(OLD))
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
